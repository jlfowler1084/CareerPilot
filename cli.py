"""CareerPilot — Main CLI entry point."""

import logging
import sys

import click
from rich.console import Console
from rich.panel import Panel
from rich.table import Table

if sys.platform == "win32":
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")
    sys.stderr.reconfigure(encoding="utf-8", errors="replace")

console = Console()

# Canonical transcript kind values — imported at module level for click.Choice
from src.transcripts.transcript_parser import CANONICAL_KINDS  # noqa: E402

# Categories that support interactive response
ACTIONABLE_CATEGORIES = {"recruiter_outreach", "interview_request", "offer"}

# Category display colors
CATEGORY_COLORS = {
    "recruiter_outreach": "green",
    "interview_request": "bright_green",
    "offer": "bright_cyan",
    "job_alert": "blue",
    "rejection": "red",
    "irrelevant": "dim",
}


@click.group(invoke_without_command=True)
@click.option("--debug", is_flag=True, help="Enable debug logging.")
@click.pass_context
def cli(ctx, debug):
    """CareerPilot — Personal career management platform."""
    level = logging.DEBUG if debug else logging.INFO
    logging.basicConfig(
        level=level,
        format="%(asctime)s %(name)s %(levelname)s %(message)s",
    )
    if ctx.invoked_subcommand is None:
        ctx.invoke(dashboard)


@cli.command()
@click.option("--days", default=7, help="Number of days to look back (default 7).")
def scan(days):
    """Scan Gmail for recruiter emails, classify, and draft responses."""
    from src.gmail.responder import RecruiterResponder
    from src.gmail.scanner import GmailScanner

    scanner = GmailScanner()
    try:
        scanner.authenticate()
    except FileNotFoundError as e:
        console.print(f"[red]{e}[/red]")
        return
    except Exception:
        console.print("[red]Gmail authentication failed. Check logs for details.[/red]")
        return

    results = scanner.scan_inbox(days_back=days)

    if not results:
        console.print("[yellow]No recruiter emails found.[/yellow]")
        return

    # Display results table
    table = Table(title=f"Gmail Scan Results ({len(results)} emails)")
    table.add_column("#", style="dim")
    table.add_column("Category", style="bold")
    table.add_column("From")
    table.add_column("Subject")
    table.add_column("Company")
    table.add_column("Role")
    table.add_column("Urgency")
    table.add_column("Summary")

    for i, r in enumerate(results, 1):
        color = CATEGORY_COLORS.get(r["category"], "white")
        table.add_row(
            str(i),
            f"[{color}]{r['category']}[/{color}]",
            r["sender"][:30],
            r["subject"][:40],
            r["company"],
            r["role"],
            r["urgency"],
            r["summary"][:50],
        )

    console.print(table)
    console.print()

    # Interactive response flow for actionable emails
    actionable = [r for r in results if r["category"] in ACTIONABLE_CATEGORIES]
    if not actionable:
        console.print("[dim]No actionable emails (recruiter_outreach/interview_request/offer) found.[/dim]")
        return

    # Create responder using scanner's authenticated service
    responder = RecruiterResponder(scanner._service)

    # Try to set up calendar for availability (non-blocking)
    cal_scheduler = None
    try:
        from src.calendar.scheduler import CalendarScheduler
        cal_scheduler = CalendarScheduler()
        cal_scheduler.authenticate()
        console.print("[dim]Calendar connected — availability will be included in responses.[/dim]")
    except Exception:
        console.print("[dim]Calendar not available — using placeholder times in responses.[/dim]")
        cal_scheduler = None

    console.print(
        f"\n[bold]{len(actionable)} actionable email(s) found.[/bold] "
        "Drafts are saved to Gmail — [bold]nothing is sent[/bold] without your approval.\n"
    )

    for r in actionable:
        console.rule(f"[bold]{r['subject'][:60]}[/bold]")
        console.print(f"  From: {r['sender']}")
        console.print(f"  Company: {r['company']}  |  Role: {r['role']}  |  Urgency: {r['urgency']}")
        console.print(f"  {r['summary']}")
        console.print()

        action = _prompt_action()
        if action == "q":
            console.print("[yellow]Quitting scan.[/yellow]")
            return
        if action == "s":
            continue

        mode_map = {"r": "interested", "d": "not_interested", "m": "more_info"}
        mode = mode_map[action]

        # Fetch full email for drafting
        email_data = scanner.get_email_details(r["message_id"])
        if not email_data:
            console.print("[red]Failed to fetch email details, skipping.[/red]")
            continue

        # Pull availability for interested mode
        availability_text = None
        avail_slots = []
        if mode == "interested" and cal_scheduler:
            try:
                avail_slots = cal_scheduler.get_availability(days_ahead=5)
                if avail_slots:
                    availability_text = cal_scheduler.format_slots(avail_slots, max_slots=3)
                    console.print(f"  [cyan]Available slots: {availability_text}[/cyan]")
                else:
                    console.print("  [yellow]No open slots found in the next 5 days.[/yellow]")
            except Exception:
                console.print("  [dim]Could not fetch availability.[/dim]")

        _handle_draft_flow(
            responder, email_data, mode, r["message_id"],
            availability_text=availability_text,
            cal_scheduler=cal_scheduler,
            company=r.get("company", ""),
        )
        console.print()


def _prompt_action():
    """Prompt user for action on an email. Returns r/d/m/s/q."""
    while True:
        choice = console.input(
            "[bold][r][/bold]espond interested  "
            "[bold][d][/bold]ecline  "
            "[bold][m][/bold]ore info  "
            "[bold][s][/bold]kip  "
            "[bold][q][/bold]uit: "
        ).strip().lower()
        if choice in ("r", "d", "m", "s", "q"):
            return choice
        console.print("[red]Invalid choice. Enter r, d, m, s, or q.[/red]")


def _handle_draft_flow(responder, email_data, mode, message_id,
                       availability_text=None, cal_scheduler=None, company=""):
    """Generate a draft, show it, and let user approve/edit/cancel."""
    draft_text = responder.draft_response(
        email_data, mode=mode, availability_text=availability_text,
    )

    if not draft_text:
        console.print("[red]Failed to generate draft. Check logs.[/red]")
        return

    while True:
        console.print()
        console.print(Panel(
            draft_text,
            title=f"Draft Reply ({mode})",
            border_style="cyan",
        ))

        choice = console.input(
            "[bold][a][/bold]pprove (save as draft)  "
            "[bold][e][/bold]dit (re-generate with feedback)  "
            "[bold][c][/bold]ancel: "
        ).strip().lower()

        if choice == "a":
            draft_id = responder.save_draft(message_id, draft_text)
            if draft_id:
                console.print(f"[green]Draft saved to Gmail (draft_id={draft_id}). NOT sent.[/green]")
            else:
                console.print("[red]Failed to save draft. Check logs.[/red]")

            # Offer calendar hold for interested responses
            if mode == "interested" and cal_scheduler:
                _offer_calendar_hold(cal_scheduler, company)

            return

        elif choice == "e":
            feedback = console.input("What should be different? ")
            augmented = dict(email_data)
            augmented["body"] = (
                email_data.get("body", "") +
                f"\n\n--- USER FEEDBACK ON PREVIOUS DRAFT ---\n"
                f"Previous draft:\n{draft_text}\n\n"
                f"Requested changes: {feedback}"
            )
            draft_text = responder.draft_response(
                augmented, mode=mode, availability_text=availability_text,
            )
            if not draft_text:
                console.print("[red]Failed to re-generate draft. Check logs.[/red]")
                return

        elif choice == "c":
            console.print("[yellow]Draft cancelled.[/yellow]")
            return

        else:
            console.print("[red]Invalid choice. Enter a, e, or c.[/red]")


def _offer_calendar_hold(cal_scheduler, company=""):
    """After approving an interested draft, offer to create calendar holds."""
    hold_choice = console.input(
        "Create calendar hold for suggested times? [bold][y][/bold]/[bold][n][/bold]: "
    ).strip().lower()

    if hold_choice != "y":
        return

    try:
        slots = cal_scheduler.get_availability(days_ahead=5)
        if not slots:
            console.print("[yellow]No available slots to hold.[/yellow]")
            return

        title = f"Interview — {company}" if company else "Interview Hold"
        # Create holds for up to 3 slots
        for slot in slots[:3]:
            event_id = cal_scheduler.create_hold(title, slot)
            if event_id:
                console.print(
                    f"  [green]Hold created: {slot.strftime('%A %B %#d at %#I:%M %p %Z')}[/green]"
                )
            else:
                console.print(
                    f"  [red]Failed to create hold for {slot.strftime('%A %B %#d')}[/red]"
                )
    except Exception:
        console.print("[red]Failed to create calendar holds. Check logs.[/red]")


@cli.command()
@click.option("--days", default=5, help="Number of days to look ahead (default 5).")
def calendar(days):
    """Show Google Calendar availability for the next 5 days."""
    from src.calendar.scheduler import CalendarScheduler

    scheduler = CalendarScheduler()
    try:
        scheduler.authenticate()
    except FileNotFoundError as e:
        console.print(f"[red]{e}[/red]")
        return
    except Exception:
        console.print("[red]Calendar authentication failed. Check logs for details.[/red]")
        return

    # Show existing events
    events = scheduler.get_events(days_ahead=days)
    if events:
        events_table = Table(title=f"Calendar Events (next {days} days)")
        events_table.add_column("Event", style="bold")
        events_table.add_column("Start")
        events_table.add_column("End")
        events_table.add_column("Status")

        for evt in events:
            status_color = "green" if evt["status"] == "tentative" else "white"
            events_table.add_row(
                evt["title"],
                evt["start"],
                evt["end"],
                f"[{status_color}]{evt['status']}[/{status_color}]",
            )
        console.print(events_table)
        console.print()

    # Show available slots
    slots = scheduler.get_availability(days_ahead=days)
    if slots:
        avail_table = Table(title=f"Available Interview Slots (next {days} days)")
        avail_table.add_column("Day", style="bold")
        avail_table.add_column("Time")
        avail_table.add_column("Status", style="green")

        for slot in slots:
            avail_table.add_row(
                slot.strftime("%A %B %#d"),
                slot.strftime("%#I:%M %p %Z"),
                "open",
            )
        console.print(avail_table)
        console.print(f"\n[dim]Formatted: {scheduler.format_slots(slots, max_slots=5)}[/dim]")
    else:
        console.print("[yellow]No available slots found in the next {days} days.[/yellow]")


@cli.group()
def journal():
    """Manage journal entries and insights."""
    pass


@journal.command("new")
@click.option("--type", "entry_type", type=click.Choice(["daily", "interview", "study", "project", "reflection"]),
              prompt="Entry type", help="Type of journal entry.")
@click.option("--mood", default=None, help="Optional mood (e.g. focused, frustrated).")
@click.option("--time", "time_spent", default=None, type=int, help="Minutes spent.")
def journal_new(entry_type, mood, time_spent):
    """Create a new journal entry."""
    from src.journal.entries import JournalManager

    console.print("[dim]Enter your journal entry (press Enter twice to finish):[/dim]")
    lines = []
    while True:
        line = console.input("")
        if line == "" and lines and lines[-1] == "":
            lines.pop()  # remove trailing blank
            break
        lines.append(line)

    content = "\n".join(lines).strip()
    if not content:
        console.print("[yellow]Empty entry, cancelled.[/yellow]")
        return

    manager = JournalManager()
    filename = manager.create_entry(entry_type, content, mood=mood, time_spent=time_spent)
    console.print(f"[green]Entry saved: {filename}[/green]")


@journal.command("list")
@click.option("--days", default=30, help="Number of days to look back.")
@click.option("--type", "entry_type", default=None, help="Filter by entry type.")
def journal_list(days, entry_type):
    """Show recent journal entries."""
    from src.journal.entries import JournalManager

    manager = JournalManager()
    entries = manager.list_entries(days_back=days, entry_type=entry_type)

    if not entries:
        console.print("[yellow]No entries found.[/yellow]")
        return

    table = Table(title=f"Journal Entries ({len(entries)} found)")
    table.add_column("Date", style="bold")
    table.add_column("Type")
    table.add_column("Tags")
    table.add_column("Mood")
    table.add_column("File", style="dim")

    for e in entries:
        tags_str = ", ".join(e["tags"]) if e["tags"] else ""
        table.add_row(e["date"], e["type"], tags_str, str(e["mood"]), e["filename"])

    console.print(table)


@journal.command("show")
@click.argument("filename")
def journal_show(filename):
    """Display a specific journal entry."""
    from rich.markdown import Markdown
    from src.journal.entries import JournalManager

    manager = JournalManager()
    entry = manager.get_entry(filename)

    if not entry:
        console.print(f"[red]Entry not found: {filename}[/red]")
        return

    header = f"[bold]{entry.get('date', '?')}[/bold] ({entry.get('type', '?')})"
    if entry.get("tags"):
        header += f"  tags: {', '.join(entry['tags'])}"
    if entry.get("mood"):
        header += f"  mood: {entry['mood']}"

    console.print(header)
    console.print()
    console.print(Markdown(entry.get("content", "")))


@journal.command("insights")
def journal_insights():
    """Run weekly summary on last 7 days of entries."""
    from src.journal.entries import JournalManager
    from src.journal.insights import InsightsEngine

    manager = JournalManager()
    entries_meta = manager.list_entries(days_back=7)

    if not entries_meta:
        console.print("[yellow]No entries in the last 7 days.[/yellow]")
        return

    # Load full content for each entry
    entries = []
    for meta in entries_meta:
        full = manager.get_entry(meta["filename"])
        if full:
            entries.append(full)

    engine = InsightsEngine()

    console.print("[bold]Weekly Summary[/bold]")
    console.print()
    summary = engine.weekly_summary(entries)
    if summary:
        console.print(Panel(summary, title="Weekly Insights", border_style="cyan"))
    else:
        console.print("[red]Failed to generate summary.[/red]")

    console.print()
    momentum = engine.momentum_check(entries)
    status_colors = {"strong": "green", "steady": "blue", "slipping": "yellow", "stalled": "red"}
    color = status_colors.get(momentum["status"], "white")
    console.print(f"Momentum: [{color}]{momentum['status'].upper()}[/{color}]")
    if momentum["explanation"]:
        console.print(f"  {momentum['explanation']}")


@journal.command("search")
@click.argument("keyword")
def journal_search(keyword):
    """Search journal entries for a keyword."""
    from src.journal.entries import JournalManager

    manager = JournalManager()
    results = manager.search_entries(keyword)

    if not results:
        console.print(f"[yellow]No entries matching '{keyword}'.[/yellow]")
        return

    table = Table(title=f"Search: '{keyword}' ({len(results)} matches)")
    table.add_column("File", style="bold")
    table.add_column("Snippet")

    for r in results:
        table.add_row(r["filename"], r["snippet"])

    console.print(table)


@cli.group(invoke_without_command=True)
@click.pass_context
def skills(ctx):
    """Show skill inventory with gap visualization."""
    if ctx.invoked_subcommand is not None:
        return

    from src.skills.tracker import SkillTracker

    tracker = SkillTracker()
    tracker.seed_defaults()
    skill_data = tracker.display_skills()
    tracker.close()

    if not skill_data:
        console.print("[yellow]No skills found.[/yellow]")
        return

    table = Table(title="Skill Inventory")
    table.add_column("Skill", style="bold")
    table.add_column("Category")
    table.add_column("Level", justify="center")
    table.add_column("Target", justify="center")
    table.add_column("Gap", justify="center")
    table.add_column("Progress")
    table.add_column("Last Practiced", style="dim")

    for s in skill_data:
        gap = s["gap"]
        if gap == 0:
            gap_str = "[green]0[/green]"
        elif gap == 1:
            gap_str = f"[yellow]{gap}[/yellow]"
        else:
            gap_str = f"[red]{gap}[/red]"

        table.add_row(
            s["name"], s["category"],
            str(s["current_level"]), str(s["target_level"]),
            gap_str, s["bar"],
            s["last_practiced"][:10] if s["last_practiced"] else "",
        )

    console.print(table)


@skills.command("update")
@click.argument("name")
@click.argument("level", type=int)
def skills_update(name, level):
    """Update a skill level (1-5)."""
    from src.skills.tracker import SkillTracker

    if not 1 <= level <= 5:
        console.print("[red]Level must be between 1 and 5.[/red]")
        return

    tracker = SkillTracker()
    tracker.seed_defaults()
    if tracker.update_skill(name, level):
        console.print(f"[green]Updated '{name}' to level {level}.[/green]")
    else:
        console.print(f"[red]Skill '{name}' not found.[/red]")
    tracker.close()


@skills.command("scan")
def skills_scan():
    """Parse all tracked applications and extract skills via AI."""
    from rich.progress import Progress

    from src.db import models
    from src.intel.skill_analyzer import SkillGapAnalyzer

    conn = models.get_connection()
    apps = conn.execute(
        "SELECT COUNT(*) FROM applications WHERE description IS NOT NULL AND description != ''"
    ).fetchone()[0]

    if apps == 0:
        console.print(
            "[yellow]No applications with stored job descriptions. "
            "Save JDs when tracking jobs to use this feature.[/yellow]"
        )
        conn.close()
        return

    console.print(f"[dim]Scanning {apps} applications for skills...[/dim]")

    analyzer = SkillGapAnalyzer()

    with Progress(console=console) as progress:
        task = progress.add_task("Extracting skills...", total=apps)

        def on_progress(current, total):
            progress.update(task, completed=current)

        result = analyzer.scan_applications(conn, progress_callback=on_progress)

    conn.close()

    console.print(
        f"\n[green]Scan complete:[/green] {result['apps_scanned']} applications, "
        f"{result['skills_found']} skill mentions extracted."
    )
    console.print("[dim]Run 'skills gaps' to see your skill gap analysis.[/dim]")


@skills.command("gaps")
def skills_gaps():
    """Show skill gaps ranked by market demand."""
    from datetime import datetime as dt

    from src.db import models

    conn = models.get_connection()
    demands = models.get_skill_demand(conn)
    conn.close()

    if not demands:
        console.print("[dim]No skill demand data yet. Run 'skills scan' first.[/dim]")
        return

    # Count totals for header
    total_skills = len(demands)
    total_apps = conn.execute(
        "SELECT COUNT(DISTINCT application_id) FROM skill_application_map"
    ).fetchone()[0] if False else 0  # avoid reopening conn

    # Group by match level
    gaps = [d for d in demands if d.get("match_level") == "gap"]
    partials = [d for d in demands if d.get("match_level") == "partial"]
    strongs = [d for d in demands if d.get("match_level") == "strong"]

    if gaps:
        console.print("\n[bold red]HIGH DEMAND GAPS (not on resume):[/bold red]")
        for d in gaps[:10]:
            req = f"{d['required_count']} required" if d["required_count"] else ""
            console.print(
                f"  [red]{d['skill_name']}[/red] "
                f"{'.' * max(1, 25 - len(d['skill_name']))} "
                f"{d['times_seen']} jobs ({req}) -- Gap"
            )

    if partials:
        console.print("\n[bold yellow]PARTIAL MATCHES (could strengthen):[/bold yellow]")
        for d in partials[:10]:
            req = f"{d['required_count']} required" if d["required_count"] else ""
            console.print(
                f"  [yellow]{d['skill_name']}[/yellow] "
                f"{'.' * max(1, 25 - len(d['skill_name']))} "
                f"{d['times_seen']} jobs ({req}) -- Partial"
            )

    if strongs:
        console.print("\n[bold green]STRONG MATCHES (competitive advantages):[/bold green]")
        for d in strongs[:10]:
            req = f"{d['required_count']} required" if d["required_count"] else ""
            console.print(
                f"  [green]{d['skill_name']}[/green] "
                f"{'.' * max(1, 25 - len(d['skill_name']))} "
                f"{d['times_seen']} jobs ({req}) -- Strong"
            )

    console.print()


@skills.command("plan")
def skills_plan():
    """Show or generate a study plan for top skill gaps."""
    import json as _json

    from src.db import models
    from src.intel.skill_analyzer import SkillGapAnalyzer

    conn = models.get_connection()
    plan = models.get_study_plan(conn)

    if not plan:
        # Try to generate one
        gaps = models.get_top_gaps(conn, limit=5)
        if not gaps:
            console.print(
                "[dim]No skill gaps found. Run 'skills scan' first.[/dim]"
            )
            conn.close()
            return

        console.print("[dim]Generating study plan for top gaps...[/dim]")
        analyzer = SkillGapAnalyzer()
        result = analyzer.generate_study_plan(conn, gaps)
        if not result:
            console.print("[red]Failed to generate study plan. Check logs.[/red]")
            conn.close()
            return
        plan = models.get_study_plan(conn)

    console.print()
    for p in plan:
        target = p.get("target_hours") or 0
        logged = p.get("study_hours_logged") or 0
        pct = int(logged / target * 100) if target > 0 else 0
        bar_filled = int(pct / 10)
        bar = "\u2588" * bar_filled + "\u2591" * (10 - bar_filled)

        rank = p.get("priority_rank", "?")
        console.print(
            f"[bold]{rank}. {p['skill_name']}[/bold] "
            f"(Priority: {'HIGH' if rank and rank <= 2 else 'MED'})"
        )
        console.print(f"   Progress: {logged}/{target} hrs  {bar}  {pct}%")

        # Show resources
        try:
            resources = _json.loads(p.get("resources") or "[]")
        except (ValueError, TypeError):
            resources = []
        for r in resources[:3]:
            console.print(f"   - {r.get('title', 'Resource')}: {r.get('url', '')}")

        console.print()

    conn.close()


@skills.command("rate")
@click.argument("skill_name")
@click.argument("level", type=int)
def skills_rate(skill_name, level):
    """Self-assess a skill (1-5). Updates skills table and recomputes match levels."""
    from src.db import models
    from src.skills.tracker import SkillTracker

    if not 1 <= level <= 5:
        console.print("[red]Level must be between 1 and 5.[/red]")
        return

    conn = models.get_connection()

    # Try to update existing skill
    existing = models.get_skill(conn, skill_name)
    if existing:
        models.update_skill(conn, skill_name, level, source="self_assessment")
    else:
        # Add as new skill
        models.add_skill(conn, skill_name, category="", current_level=level, target_level=max(level, 3))

    models.update_match_levels(conn)
    conn.close()

    console.print(f"[green]Rated '{skill_name}' at level {level}/5. Match levels updated.[/green]")


@skills.command("log")
@click.argument("skill_name")
@click.argument("hours", type=float)
@click.option("--note", default="", help="Study note.")
def skills_log_time(skill_name, hours, note):
    """Log study time for a skill."""
    from src.db import models

    conn = models.get_connection()
    if models.log_study_time(conn, skill_name, hours, note):
        plan = conn.execute(
            "SELECT study_hours_logged, target_hours FROM study_plan WHERE skill_name = ?",
            (skill_name,),
        ).fetchone()
        logged = plan["study_hours_logged"] if plan else hours
        target = plan["target_hours"] if plan and plan["target_hours"] else 0
        pct = int(logged / target * 100) if target > 0 else 0
        bar_filled = int(pct / 10)
        bar = "\u2588" * bar_filled + "\u2591" * (10 - bar_filled)
        console.print(
            f"[green]Logged {hours}h for {skill_name}.[/green] "
            f"Total: {logged}/{target} hrs  {bar}  {pct}%"
        )
    else:
        console.print(
            f"[red]'{skill_name}' not in study plan. "
            f"Run 'skills plan' first.[/red]"
        )
    conn.close()


@skills.command("focus")
def skills_focus():
    """Show this week's top 3 study priorities."""
    from src.db import models

    conn = models.get_connection()
    plan = models.get_study_plan(conn)
    conn.close()

    if not plan:
        console.print("[dim]No active study plan. Run 'skills plan' to generate one.[/dim]")
        return

    console.print("\n[bold]This Week's Focus:[/bold]")
    for p in plan[:3]:
        target = p.get("target_hours") or 0
        logged = p.get("study_hours_logged") or 0
        pct = int(logged / target * 100) if target > 0 else 0
        bar_filled = int(pct / 10)
        bar = "\u2588" * bar_filled + "\u2591" * (10 - bar_filled)

        demand = ""
        console.print(
            f"  {p.get('priority_rank', '?')}. [bold]{p['skill_name']}[/bold] "
            f"({logged}/{target} hrs)  {bar}  {pct}%"
        )
    console.print()


@skills.command("match")
@click.argument("application_id", type=int)
def skills_match(application_id):
    """Show skill match analysis for a specific application."""
    from src.db import models

    conn = models.get_connection()
    app = conn.execute(
        "SELECT title, company FROM applications WHERE id = ?", (application_id,)
    ).fetchone()
    if not app:
        console.print(f"[red]Application #{application_id} not found.[/red]")
        conn.close()
        return

    app_skills = models.get_skills_for_application(conn, application_id)
    conn.close()

    if not app_skills:
        console.print(
            f"[dim]No skill data for '{app['title']}' at {app['company']}. "
            f"Run 'skills scan' first.[/dim]"
        )
        return

    console.print(f"\n[bold]Skill Match: {app['title']} at {app['company']}[/bold]\n")

    table = Table()
    table.add_column("Skill", style="bold")
    table.add_column("Required?")
    table.add_column("Match")
    table.add_column("Demand")

    for s in app_skills:
        level = s.get("requirement_level", "mentioned")
        match = s.get("match_level") or "unknown"
        seen = s.get("times_seen") or 0

        if match == "strong":
            match_str = "[green]Strong[/green]"
        elif match == "partial":
            match_str = "[yellow]Partial[/yellow]"
        elif match == "gap":
            match_str = "[red]Gap[/red]"
        else:
            match_str = "[dim]?[/dim]"

        table.add_row(
            s["skill_name"],
            level,
            match_str,
            f"{seen} jobs" if seen else "",
        )

    console.print(table)


@skills.command("report")
def skills_report():
    """Full skill gap report with demand, gaps, and study progress."""
    from src.db import models

    conn = models.get_connection()
    demands = models.get_skill_demand(conn)
    plan = models.get_study_plan(conn)

    if not demands:
        console.print("[dim]No skill data. Run 'skills scan' first.[/dim]")
        conn.close()
        return

    gaps = [d for d in demands if d.get("match_level") == "gap"]
    partials = [d for d in demands if d.get("match_level") == "partial"]
    strongs = [d for d in demands if d.get("match_level") == "strong"]

    console.print(Panel(
        f"[bold]{len(demands)}[/bold] skills tracked across job applications\n"
        f"[red]{len(gaps)}[/red] gaps | "
        f"[yellow]{len(partials)}[/yellow] partial | "
        f"[green]{len(strongs)}[/green] strong",
        title="Skill Gap Report",
    ))

    # Top gaps
    if gaps:
        console.print("\n[bold]Top Gaps to Address:[/bold]")
        for d in gaps[:5]:
            console.print(
                f"  [red]{d['skill_name']}[/red] -- "
                f"{d['times_seen']} jobs, {d['required_count']} required"
            )

    # Study progress
    if plan:
        console.print("\n[bold]Study Progress:[/bold]")
        for p in plan[:5]:
            target = p.get("target_hours") or 0
            logged = p.get("study_hours_logged") or 0
            pct = int(logged / target * 100) if target > 0 else 0
            bar_filled = int(pct / 10)
            bar = "\u2588" * bar_filled + "\u2591" * (10 - bar_filled)
            console.print(
                f"  {p['skill_name']}: {logged}/{target} hrs  {bar}  {pct}%"
            )

    # Strengths
    if strongs:
        console.print("\n[bold]Competitive Advantages:[/bold]")
        names = ", ".join(d["skill_name"] for d in strongs[:8])
        console.print(f"  [green]{names}[/green]")

    console.print()
    conn.close()


@cli.command()
@click.option("--hours", default=15, help="Available study hours per week (default 15).")
def roadmap(hours):
    """Generate a study roadmap from skill gaps via Claude."""
    from src.skills.roadmap import RoadmapGenerator
    from src.skills.tracker import SkillTracker

    tracker = SkillTracker()
    tracker.seed_defaults()
    gaps = tracker.get_gaps()
    tracker.close()

    if not gaps:
        console.print("[green]No skill gaps! All skills at target level.[/green]")
        return

    console.print(f"[dim]Generating roadmap for {len(gaps)} skill gap(s) ({hours} hrs/week)...[/dim]")

    generator = RoadmapGenerator()
    roadmap_text = generator.generate_roadmap(gaps, available_hours_per_week=hours)

    if roadmap_text:
        console.print(Panel(roadmap_text, title="Study Roadmap", border_style="cyan"))
    else:
        console.print("[red]Failed to generate roadmap. Check logs.[/red]")


@cli.command()
@click.option("--profile", "profile_ids", multiple=True, help="Profile ID(s) to run. Omit for interactive selection.")
def search(profile_ids):
    """Run job search profiles across Indeed and Dice."""
    import webbrowser

    from config.search_profiles import SEARCH_PROFILES
    from src.jobs.analyzer import JobAnalyzer
    from src.jobs.searcher import JobSearcher
    from src.jobs.tracker import ApplicationTracker

    profiles = SEARCH_PROFILES

    # Interactive profile selection if none specified
    if not profile_ids:
        console.print("[bold]Available Search Profiles:[/bold]\n")
        profile_list = list(profiles.items())
        for i, (pid, p) in enumerate(profile_list, 1):
            console.print(f"  [bold]{i}.[/bold] {p.get('label', pid)} [{p.get('sources', 'both')}]")

        console.print(f"\n  [bold]a.[/bold] Run all profiles")
        choice = console.input("\nSelect profiles (comma-separated numbers, or 'a' for all): ").strip().lower()

        if choice == "a" or choice == "":
            profile_ids = [pid for pid, _ in profile_list]
        else:
            try:
                indices = [int(x.strip()) for x in choice.split(",")]
                profile_ids = [profile_list[i - 1][0] for i in indices if 1 <= i <= len(profile_list)]
            except (ValueError, IndexError):
                console.print("[red]Invalid selection.[/red]")
                return

        if not profile_ids:
            console.print("[yellow]No profiles selected.[/yellow]")
            return

    console.print(f"\n[dim]Running {len(profile_ids)} profile(s)...[/dim]")

    searcher = JobSearcher()
    results = []

    for pid in profile_ids:
        p = profiles.get(pid)
        if not p:
            console.print(f"[red]Unknown profile: {pid}[/red]")
            continue
        console.print(f"  [dim]Searching: {p.get('label', pid)}...[/dim]")
        profile_results = searcher.run_profiles([pid])
        results.extend(profile_results)
        console.print(f"    [dim]{len(profile_results)} results[/dim]")

    # Deduplicate across profiles
    results = searcher._deduplicate(results)

    if not results:
        console.print("[yellow]No results found.[/yellow]")
        return

    # Display results table
    table = Table(title=f"Job Search Results ({len(results)} jobs)")
    table.add_column("#", style="dim", width=4)
    table.add_column("Title", style="bold")
    table.add_column("Company")
    table.add_column("Location")
    table.add_column("Salary")
    table.add_column("Source", width=8)

    for i, r in enumerate(results, 1):
        table.add_row(
            str(i),
            str(r.get("title", ""))[:45],
            str(r.get("company", ""))[:25],
            str(r.get("location", ""))[:20],
            str(r.get("salary", ""))[:15] or "-",
            r.get("source", ""),
        )

    console.print(table)

    # Interactive per-job actions
    tracker = ApplicationTracker()
    analyzer = JobAnalyzer()

    console.print(
        "\n[dim]Per-job actions: [s]ave to tracker, [A]pply, "
        "[a]nalyze fit, [o]pen in browser, [n]ext, [q]uit[/dim]\n"
    )

    from src.jobs.applicant import JobApplicant
    applicant = JobApplicant()

    for i, r in enumerate(results, 1):
        console.rule(f"[bold]{i}. {r.get('title', '?')} at {r.get('company', '?')}[/bold]")
        console.print(f"  Location: {r.get('location', '?')}  |  Salary: {r.get('salary', '') or 'N/A'}")
        console.print(f"  Source: {r.get('source', '?')}  |  Posted: {r.get('posted_date', '') or 'N/A'}")
        easy_apply = r.get("easy_apply", False)
        console.print(f"  Easy Apply: {'[green]Yes[/green]' if easy_apply else '[dim]No[/dim]'}")
        if r.get("url"):
            console.print(f"  URL: {r['url'][:80]}")

        while True:
            action = console.input(
                "  [bold][s][/bold]ave  [bold][A][/bold]pply  [bold][a][/bold]nalyze  "
                "[bold][o][/bold]pen  [bold][n][/bold]ext  [bold][q][/bold]uit: "
            ).strip()

            if action.lower() == "s":
                job_id = tracker.save_job(r)
                console.print(f"  [green]Saved to tracker (id={job_id})[/green]")
            elif action == "A":
                _apply_job_flow(r, applicant, console)
            elif action.lower() == "a":
                console.print("  [dim]Analyzing fit...[/dim]")
                job_desc = (
                    f"Title: {r.get('title', '')}\n"
                    f"Company: {r.get('company', '')}\n"
                    f"Location: {r.get('location', '')}\n"
                    f"Salary: {r.get('salary', '')}\n"
                    f"Type: {r.get('job_type', '')}\n"
                )
                analysis = analyzer.analyze_fit(job_desc)
                if analysis:
                    _display_fit_analysis(analysis)
                else:
                    console.print("  [red]Analysis failed.[/red]")
            elif action.lower() == "o":
                url = r.get("url", "")
                if url:
                    webbrowser.open(url)
                    console.print("  [cyan]Opened in browser.[/cyan]")
                else:
                    console.print("  [yellow]No URL available.[/yellow]")
            elif action.lower() == "n":
                break
            elif action.lower() == "q":
                tracker.close()
                applicant.close()
                return
            else:
                console.print("  [red]Invalid choice. Enter s, A, a, o, n, or q.[/red]")

    tracker.close()
    applicant.close()
    console.print("\n[dim]Search complete.[/dim]")


def _apply_job_flow(job_data, applicant, con):
    """Interactive apply flow for a single job."""
    easy_apply = job_data.get("easy_apply", False)
    method = "easy_apply" if easy_apply else "browser"

    # Show job details panel
    details = (
        f"[bold]{job_data.get('title', '?')}[/bold] at {job_data.get('company', '?')}\n"
        f"Location: {job_data.get('location', '?')}  |  Salary: {job_data.get('salary', '') or 'N/A'}\n"
        f"Source: {job_data.get('source', '?')}  |  Easy Apply: {'Yes' if easy_apply else 'No'}"
    )
    if job_data.get("url"):
        details += f"\nURL: {job_data['url'][:80]}"
    con.print(Panel(details, title="Apply", border_style="green"))

    confirm = con.input("  Apply to this position? [y/n]: ").strip().lower()
    if confirm != "y":
        con.print("  [dim]Skipped.[/dim]")
        return

    # Offer document generation
    gen_docs = con.input("  Generate tailored resume and cover letter? [y/n]: ").strip().lower()
    if gen_docs == "y":
        with con.status("[cyan]Generating documents with Claude...[/cyan]"):
            doc_result = applicant.generate_application_docs(job_data)
        if doc_result.get("resume_path"):
            con.print(f"  [green]Resume:[/green] {doc_result['resume_path']}")
        if doc_result.get("cover_letter_path"):
            con.print(f"  [green]Cover letter:[/green] {doc_result['cover_letter_path']}")
        if not doc_result.get("resume_path") and not doc_result.get("cover_letter_path"):
            con.print("  [yellow]Document generation failed — continuing with apply.[/yellow]")

    # Execute apply flow
    if easy_apply:
        result = applicant.apply_dice_easy(job_data)
    else:
        result = applicant.apply_with_resume(job_data)

    if result.get("opened"):
        con.print("  [cyan]Opened in browser.[/cyan]")
    if result.get("clipboard"):
        con.print("  [green]Profile copied to clipboard — paste into application form.[/green]")

    # Clipboard field picker
    con.print("  [dim]Copy field: [n]ame [e]mail [p]hone [a]ddress [l]inkedIn [r]esume summary[/dim]")
    while True:
        field_choice = con.input("  Copy field (or Enter to continue): ").strip().lower()
        if not field_choice:
            break
        field_key = {
            "n": "name", "e": "email", "p": "phone",
            "a": "address", "l": "linkedin", "r": "summary",
        }.get(field_choice)
        if field_key:
            val = applicant.copy_field(field_key)
            if val:
                con.print(f"  [green]Copied {field_key}: {val[:60]}[/green]")
            else:
                con.print(f"  [yellow]{field_key} not set in profile.[/yellow]")
        else:
            con.print("  [red]Invalid field.[/red]")

    # Confirm actual submission
    tracker_id = result.get("tracker_id")
    did_apply = con.input("  Mark as applied? [y/n]: ").strip().lower()
    if did_apply == "y" and tracker_id:
        applicant.mark_applied(tracker_id, method=method)
        con.print(f"  [green]Marked as applied (id={tracker_id}).[/green]")
    else:
        con.print("  [dim]Not marked — you can update the tracker later.[/dim]")


def _display_fit_analysis(analysis):
    """Display job fit analysis with Rich formatting."""
    score = analysis.get("match_score", 0)
    if score >= 7:
        score_color = "green"
    elif score >= 5:
        score_color = "yellow"
    else:
        score_color = "red"

    console.print(f"  [bold]Match Score:[/bold] [{score_color}]{score}/10[/{score_color}]")

    matching = analysis.get("matching_skills", [])
    if matching:
        console.print(f"  [green]Matching:[/green] {', '.join(matching)}")

    gaps = analysis.get("gap_skills", [])
    if gaps:
        console.print(f"  [yellow]Gaps:[/yellow] {', '.join(gaps)}")

    tweaks = analysis.get("resume_tweaks", [])
    if tweaks:
        console.print("  [cyan]Resume tweaks:[/cyan]")
        for t in tweaks:
            console.print(f"    - {t}")

    flags = analysis.get("red_flags", [])
    if flags:
        console.print("  [red]Red flags:[/red]")
        for f in flags:
            console.print(f"    - {f}")


@cli.command()
def apply():
    """Batch apply to tracked jobs (status: found or interested)."""
    from src.jobs.applicant import JobApplicant

    applicant = JobApplicant()
    jobs = applicant.get_actionable_jobs()

    if not jobs:
        console.print("[yellow]No actionable jobs (found/interested) in tracker.[/yellow]")
        applicant.close()
        return

    # Display jobs table
    table = Table(title=f"Actionable Jobs ({len(jobs)})")
    table.add_column("#", style="dim", width=4)
    table.add_column("Title", style="bold")
    table.add_column("Company")
    table.add_column("Location")
    table.add_column("Status")
    table.add_column("Easy Apply", width=10)

    for i, j in enumerate(jobs, 1):
        table.add_row(
            str(i),
            str(j.get("title", ""))[:40],
            str(j.get("company", ""))[:25],
            str(j.get("location", ""))[:20],
            j.get("status", ""),
            "[green]Yes[/green]" if j.get("easy_apply") else "[dim]No[/dim]",
        )

    console.print(table)

    selection = console.input(
        "\nSelect jobs to apply (comma-separated numbers, ranges like 1-3, or 'all'): "
    ).strip()
    if not selection:
        console.print("[dim]No selection made.[/dim]")
        applicant.close()
        return

    selected = applicant.batch_select(jobs, selection)
    if not selected:
        console.print("[yellow]No valid jobs selected.[/yellow]")
        applicant.close()
        return

    # Offer batch doc generation
    batch_docs = console.input(
        "Auto-generate tailored docs for all selected jobs? [y/n]: "
    ).strip().lower()
    if batch_docs == "y":
        console.print("[cyan]Generating documents for selected jobs...[/cyan]")
        for j in selected:
            with console.status(f"[cyan]Generating docs for {j.get('title', '?')}...[/cyan]"):
                doc_result = applicant.generate_application_docs(j)
            if doc_result.get("resume_path"):
                console.print(f"  [green]Resume:[/green] {doc_result['resume_path']}")
            if doc_result.get("cover_letter_path"):
                console.print(f"  [green]Cover letter:[/green] {doc_result['cover_letter_path']}")
        console.print()

    console.print(f"\n[bold]Applying to {len(selected)} job(s)...[/bold]\n")

    applied_count = 0
    skipped_count = 0

    for j in selected:
        _apply_job_flow(j, applicant, console)
        # Check if it was marked applied
        updated = applicant._tracker.get_job(j["id"])
        if updated and updated.get("status") == "applied":
            applied_count += 1
        else:
            skipped_count += 1
        console.print()

    applicant.close()
    console.print(
        f"\n[bold]Batch complete:[/bold] "
        f"[green]{applied_count} applied[/green], "
        f"[dim]{skipped_count} skipped[/dim]"
    )


@cli.group(invoke_without_command=True)
@click.pass_context
def tracker(ctx):
    """Manage job application pipeline."""
    if ctx.invoked_subcommand is not None:
        return
    # Default: show pipeline
    ctx.invoke(tracker_show)


@tracker.command("show")
def tracker_show():
    """Show application pipeline (kanban view)."""
    from src.jobs.tracker import ApplicationTracker

    t = ApplicationTracker()
    pipeline = t.get_pipeline()
    t.close()

    total = sum(len(jobs) for jobs in pipeline.values())
    if total == 0:
        console.print("[yellow]No applications tracked. Run 'search' to find and save jobs.[/yellow]")
        return

    # Display order and colors
    status_display = [
        ("found", "dim"),
        ("interested", "blue"),
        ("applied", "yellow"),
        ("phone_screen", "bright_yellow"),
        ("interview", "green"),
        ("offer", "bright_green"),
        ("rejected", "red"),
        ("withdrawn", "magenta"),
        ("ghosted", "dim red"),
    ]

    console.print(Panel(f"[bold]Application Pipeline[/bold] ({total} jobs)", border_style="cyan"))

    for status, color in status_display:
        jobs = pipeline.get(status, [])
        if not jobs:
            continue

        table = Table(title=f"{status.replace('_', ' ').title()} ({len(jobs)})", border_style=color)
        table.add_column("ID", style="dim", width=5)
        table.add_column("Title", style="bold")
        table.add_column("Company")
        table.add_column("Location")
        table.add_column("Days", justify="center", width=6)

        for j in jobs:
            days = ""
            ref_date = j.get("date_applied") or j.get("date_found") or ""
            if ref_date:
                try:
                    from datetime import datetime
                    dt = datetime.fromisoformat(ref_date)
                    days = str((datetime.now() - dt).days)
                except (ValueError, TypeError):
                    pass

            table.add_row(
                str(j["id"]),
                str(j.get("title", ""))[:40],
                str(j.get("company", ""))[:25],
                str(j.get("location", ""))[:20],
                days,
            )

        console.print(table)
        console.print()


@tracker.command("stats")
def tracker_stats():
    """Show search and application analytics."""
    from src.jobs.tracker import ApplicationTracker

    t = ApplicationTracker()
    stats = t.get_stats()
    t.close()

    if stats["total"] == 0:
        console.print("[yellow]No applications tracked yet.[/yellow]")
        return

    # Summary panel
    summary = (
        f"Total jobs tracked: [bold]{stats['total']}[/bold]\n"
        f"Applications sent: [bold]{stats['applied_count']}[/bold]\n"
        f"Responses received: [bold]{stats['responded_count']}[/bold]\n"
        f"Response rate: [bold]{stats['response_rate']:.0f}%[/bold]\n"
        f"Avg days to response: [bold]{stats['avg_days_to_response']:.1f}[/bold]"
    )
    console.print(Panel(summary, title="Application Stats", border_style="cyan"))

    # Status breakdown
    table = Table(title="Status Breakdown")
    table.add_column("Status", style="bold")
    table.add_column("Count", justify="center")
    table.add_column("Bar")

    max_count = max(stats["by_status"].values()) if stats["by_status"] else 1
    status_colors = {
        "found": "dim", "interested": "blue", "applied": "yellow",
        "phone_screen": "bright_yellow", "interview": "green",
        "offer": "bright_green", "rejected": "red",
        "withdrawn": "magenta", "ghosted": "dim red",
    }

    for status in ["found", "interested", "applied", "phone_screen",
                    "interview", "offer", "rejected", "withdrawn", "ghosted"]:
        count = stats["by_status"].get(status, 0)
        if count == 0:
            continue
        bar_len = int(count / max_count * 20) if max_count > 0 else 0
        color = status_colors.get(status, "white")
        bar = f"[{color}]{'█' * bar_len}[/{color}]"
        table.add_row(status.replace("_", " ").title(), str(count), bar)

    console.print(table)


@tracker.command("update")
@click.argument("job_id", type=int)
def tracker_update(job_id):
    """Update application status for a job."""
    from src.jobs.tracker import ApplicationTracker, VALID_STATUSES

    t = ApplicationTracker()
    job = t.get_job(job_id)

    if not job:
        console.print(f"[red]Job id={job_id} not found.[/red]")
        t.close()
        return

    # Show current details
    console.print(f"\n[bold]{job['title']}[/bold] at {job['company']}")
    console.print(f"  Status: [bold]{job['status']}[/bold]  |  Location: {job.get('location', '')}")
    if job.get("notes"):
        console.print(f"  Notes: {job['notes'][:80]}")

    # Show status options
    statuses = sorted(VALID_STATUSES)
    console.print("\n[bold]Available statuses:[/bold]")
    for i, s in enumerate(statuses, 1):
        marker = " [green]<-- current[/green]" if s == job["status"] else ""
        console.print(f"  {i}. {s.replace('_', ' ').title()}{marker}")

    choice = console.input("\nSelect new status (number): ").strip()
    try:
        idx = int(choice) - 1
        new_status = statuses[idx]
    except (ValueError, IndexError):
        console.print("[red]Invalid selection.[/red]")
        t.close()
        return

    notes = console.input("Notes (optional, press Enter to skip): ").strip() or None

    if t.update_status(job_id, new_status, notes=notes):
        console.print(f"[green]Updated to '{new_status}'.[/green]")
    else:
        console.print("[red]Update failed.[/red]")

    t.close()


@tracker.command("applied-today")
def tracker_applied_today():
    """Show jobs applied to today."""
    from src.jobs.applicant import JobApplicant

    applicant = JobApplicant()
    jobs = applicant.get_applied_today()
    applicant.close()

    if not jobs:
        console.print("[dim]No applications submitted today.[/dim]")
        return

    table = Table(title=f"Applied Today ({len(jobs)})")
    table.add_column("ID", style="dim", width=5)
    table.add_column("Title", style="bold")
    table.add_column("Company")
    table.add_column("Notes")

    for j in jobs:
        table.add_row(
            str(j["id"]),
            str(j.get("title", ""))[:40],
            str(j.get("company", ""))[:25],
            str(j.get("notes", ""))[:30],
        )

    console.print(table)


@tracker.command("status")
@click.argument("job_id", type=int)
@click.argument("status", type=str)
def tracker_ext_status(job_id, status):
    """Set external ATS status on an application."""
    from src.jobs.tracker import ApplicationTracker

    t = ApplicationTracker()
    job = t.get_job(job_id)

    if not job:
        console.print(f"[red]Job id={job_id} not found.[/red]")
        t.close()
        return

    old_ext = job.get("external_status") or "(none)"
    if t.update_external_status(job_id, status):
        console.print(f"[bold]{job['title']}[/bold] at {job['company']}")
        console.print(f"  External status: {old_ext} → [green]{status}[/green]")
    else:
        console.print("[red]Update failed.[/red]")

    t.close()


@tracker.command("withdraw")
@click.argument("job_id", type=int)
def tracker_withdraw(job_id):
    """Withdraw an application."""
    from src.jobs.tracker import ApplicationTracker

    t = ApplicationTracker()
    job = t.get_job(job_id)

    if not job:
        console.print(f"[red]Job id={job_id} not found.[/red]")
        t.close()
        return

    if t.withdraw_application(job_id):
        console.print(
            f"[yellow]Withdrawn:[/yellow] [bold]{job['title']}[/bold] at {job['company']}"
        )
    else:
        console.print("[red]Withdraw failed.[/red]")

    t.close()


@tracker.command("stale")
def tracker_stale():
    """Show applications with no status update in 14+ days."""
    from datetime import datetime

    from src.jobs.tracker import ApplicationTracker

    t = ApplicationTracker()
    stale = t.get_stale_applications()
    t.close()

    if not stale:
        console.print("[green]No stale applications.[/green]")
        return

    table = Table(title=f"Stale Applications ({len(stale)})")
    table.add_column("ID", style="dim", width=5)
    table.add_column("Title", style="bold")
    table.add_column("Company")
    table.add_column("Status")
    table.add_column("External Status")
    table.add_column("Days Since Update", justify="center")

    for j in stale:
        days = ""
        ref = j.get("external_status_updated") or j.get("date_found") or ""
        if ref:
            try:
                dt = datetime.fromisoformat(ref)
                days = str((datetime.now() - dt).days)
            except (ValueError, TypeError):
                pass

        table.add_row(
            str(j["id"]),
            str(j.get("title", ""))[:40],
            str(j.get("company", ""))[:25],
            j.get("status", ""),
            j.get("external_status") or "(none)",
            days,
        )

    console.print(table)


def _run_tracker_add_wizard():
    """Interactive wizard for `tracker add`. Returns a dict of fields, or None if cancelled."""
    from rich.prompt import Confirm, Prompt

    from src.jobs.tracker import VALID_STATUSES

    console.print("\n[bold]Add a new application to the tracker.[/bold]")
    console.print("[dim]Press Ctrl-C at any time to cancel without saving.[/dim]\n")

    # Required fields — re-prompt on empty
    title = ""
    while not title.strip():
        title = Prompt.ask("  Title")
    company = ""
    while not company.strip():
        company = Prompt.ask("  Company")

    location = Prompt.ask("  Location [dim](optional)[/dim]", default="")
    url = Prompt.ask("  URL [dim](optional)[/dim]", default="")

    description = ""
    if Confirm.ask("  Open editor for job description?", default=False):
        description = click.edit() or ""

    status = Prompt.ask(
        "  Status",
        choices=sorted(VALID_STATUSES),
        default="interested",
    )
    notes = Prompt.ask("  Notes [dim](optional)[/dim]", default="")

    # Summary panel
    console.print("\n[bold]Summary:[/bold]")
    console.print(f"  Title:       {title}")
    console.print(f"  Company:     {company}")
    if location:
        console.print(f"  Location:    {location}")
    if url:
        console.print(f"  URL:         {url}")
    console.print(f"  Status:      {status}")
    if notes:
        console.print(f"  Notes:       {notes}")
    if description:
        console.print(f"  Description: {len(description)} chars")

    if not Confirm.ask("\nCreate this application?", default=True):
        return None

    return {
        "title": title.strip(),
        "company": company.strip(),
        "location": location.strip(),
        "url": url.strip(),
        "description": description,
        "status": status,
        "notes": notes.strip(),
    }


@tracker.command("add")
@click.option("--title", default=None, help="Job title.")
@click.option("--company", default=None, help="Company name.")
@click.option("--location", default="", help="Job location.")
@click.option("--url", default="", help="Job posting URL.")
@click.option("--description", default="", help="Job description text.")
@click.option(
    "--status",
    type=click.Choice(sorted([
        "found", "interested", "applied", "phone_screen",
        "interview", "offer", "rejected", "withdrawn", "ghosted",
    ])),
    default="interested",
    show_default=True,
    help="Initial application status.",
)
@click.option("--notes", default="", help="Free-form notes.")
def tracker_add(title, company, location, url, description, status, notes):
    """Add a job application manually via wizard or flags."""
    import sys

    from src.jobs.tracker import ApplicationTracker

    # Non-interactive path: both required flags present
    if title and company:
        fields = {
            "title": title.strip(),
            "company": company.strip(),
            "location": location,
            "url": url,
            "description": description,
            "status": status,
            "notes": notes,
        }
    else:
        try:
            wizard_fields = _run_tracker_add_wizard()
        except KeyboardInterrupt:
            console.print("\n[yellow]Cancelled — no application saved.[/yellow]")
            raise click.Abort()
        except EOFError:
            console.print(
                "[red]Error: --title and --company are required "
                "when not running interactively.[/red]"
            )
            sys.exit(2)
        if wizard_fields is None:
            console.print("[yellow]Cancelled — no application saved.[/yellow]")
            return
        fields = wizard_fields

    t = ApplicationTracker()
    try:
        if fields["url"]:
            dup = t.find_by_url(fields["url"])
            if dup:
                from rich.prompt import Confirm

                console.print(
                    f"[yellow]Possible duplicate: #{dup['id']} "
                    f"{dup.get('title') or '(untitled)'} @ "
                    f"{dup.get('company') or '(unknown)'}[/yellow]"
                )
                if not Confirm.ask("Create anyway?", default=False):
                    console.print("[yellow]Aborted — no application saved.[/yellow]")
                    return

        app_id = t.save_job(
            {
                "title": fields["title"],
                "company": fields["company"],
                "location": fields["location"],
                "url": fields["url"],
                "description": fields["description"] or None,
                "source": "manual",
                "notes": fields["notes"],
            },
            status=fields["status"],
        )
        console.print(
            f"[green]Created application #{app_id}: {fields['title']} @ "
            f"{fields['company']} [status={fields['status']}][/green]"
        )
    finally:
        t.close()


_IMPORT_STATUS_CHOICES = sorted({
    "found", "interested", "applied", "phone_screen",
    "interview", "offer", "rejected", "withdrawn", "ghosted",
})


@tracker.command("import-from-email")
@click.argument("message_id")
@click.option(
    "--status", "initial_status",
    type=click.Choice(_IMPORT_STATUS_CHOICES),
    default="found",
    show_default=True,
    help="Initial application status.",
)
@click.option("--dry-run", is_flag=True, help="Extract and preview without saving.")
def tracker_import_from_email(message_id, initial_status, dry_run):
    """Create an application from a Gmail message with a PDF/DOCX attachment."""
    from config import settings

    from src.gmail.attachments import extract_job_description_from_email
    from src.gmail.auth import get_gmail_service
    from src.jobs.tracker import ApplicationTracker

    t = ApplicationTracker()
    try:
        existing = t.find_application_by_message_id(message_id)
        if existing:
            console.print(
                f"[yellow]Application already imported from this email:[/yellow] "
                f"#{existing['id']} {existing.get('title') or '(untitled)'} @ "
                f"{existing.get('company') or '(unknown)'}"
            )
            return

        try:
            service = get_gmail_service(
                credentials_file=settings.GOOGLE_CREDENTIALS_FILE,
                token_path=settings.GMAIL_TOKEN_PATH,
                scopes=settings.GMAIL_SCOPES,
            )
        except FileNotFoundError as exc:
            console.print(f"[red]{exc}[/red]")
            raise click.Abort()

        result = extract_job_description_from_email(service, message_id)
        if not result:
            console.print(
                f"[red]No supported attachment (PDF or DOCX) found "
                f"on message {message_id}.[/red]"
            )
            raise click.Abort()

        console.print(f"\n[bold]Extracted from {result['filename']}[/bold]")
        console.print(f"  Title:   {result['title'] or '[dim](not detected)[/dim]'}")
        console.print(f"  Company: {result['company'] or '[dim](not detected)[/dim]'}")
        console.print(f"  Length:  {len(result['description'])} chars")

        if dry_run:
            console.print("\n[yellow]--dry-run set; not saving.[/yellow]")
            return

        app_id = t.save_job(
            {
                "title": result["title"] or "(untitled)",
                "company": result["company"] or "(unknown)",
                "description": result["description"],
                "source": "email_import",
                "message_id": message_id,
            },
            status=initial_status,
        )
        console.print(
            f"\n[green]Created application #{app_id}: "
            f"{result['title'] or '(untitled)'} @ "
            f"{result['company'] or '(unknown)'} "
            f"[status={initial_status}][/green]"
        )
    finally:
        t.close()


@cli.group(invoke_without_command=True)
@click.pass_context
def portals(ctx):
    """Manage ATS portal accounts."""
    if ctx.invoked_subcommand is not None:
        return
    ctx.invoke(portals_list)


@portals.command("list")
def portals_list():
    """List all ATS portal accounts."""
    from datetime import datetime

    from src.db import models

    conn = models.get_connection()
    all_portals = models.list_portals(conn)

    if not all_portals:
        console.print("[yellow]No portal accounts tracked. Run 'portals add' to add one.[/yellow]")
        conn.close()
        return

    table = Table(title=f"ATS Portal Accounts ({len(all_portals)})")
    table.add_column("ID", style="dim", width=5)
    table.add_column("Company", style="bold")
    table.add_column("ATS Type")
    table.add_column("Portal URL")
    table.add_column("Email")
    table.add_column("Last Checked")
    table.add_column("Apps", justify="center")

    for p in all_portals:
        # Count pending apps for this portal
        app_count = conn.execute(
            "SELECT COUNT(*) as cnt FROM applications "
            "WHERE portal_id = ? AND status NOT IN ('withdrawn', 'rejected', 'ghosted')",
            (p["id"],),
        ).fetchone()["cnt"]

        # Determine staleness color
        style = ""
        if app_count > 0 and p["last_checked"]:
            try:
                last = datetime.fromisoformat(p["last_checked"])
                days_ago = (datetime.now() - last).days
                if days_ago >= 14:
                    style = "red"
                elif days_ago >= 7:
                    style = "yellow"
            except (ValueError, TypeError):
                pass
        elif app_count > 0 and not p["last_checked"]:
            style = "red"

        last_checked_display = ""
        if p["last_checked"]:
            try:
                last = datetime.fromisoformat(p["last_checked"])
                days_ago = (datetime.now() - last).days
                if days_ago == 0:
                    last_checked_display = "Today"
                elif days_ago == 1:
                    last_checked_display = "Yesterday"
                else:
                    last_checked_display = f"{days_ago} days ago"
            except (ValueError, TypeError):
                last_checked_display = p["last_checked"][:10]
        else:
            last_checked_display = "Never"

        table.add_row(
            str(p["id"]),
            f"[{style}]{p['company']}[/{style}]" if style else p["company"],
            p["ats_type"],
            str(p["portal_url"])[:40],
            p["email_used"],
            last_checked_display,
            str(app_count),
        )

    console.print(table)
    conn.close()


@portals.command("add")
def portals_add():
    """Add a new ATS portal account."""
    from src.db import models

    company = click.prompt("Company")
    ats_type = click.prompt(
        "ATS type",
        type=click.Choice(["Workday", "Greenhouse", "Lever", "iCIMS", "Taleo", "Custom"]),
    )
    portal_url = click.prompt("Portal URL")
    email_used = click.prompt("Email", default="jlfowler1084@gmail.com")
    notes = click.prompt("Notes", default="", show_default=False)

    conn = models.get_connection()
    pid = models.add_portal(
        conn, company=company, ats_type=ats_type, portal_url=portal_url,
        email_used=email_used, notes=notes or None,
    )
    conn.close()

    console.print(f"[green]Portal added (id={pid}): {company} ({ats_type})[/green]")


@portals.command("check")
@click.argument("portal_id", type=int)
def portals_check(portal_id):
    """Open a portal in the browser and mark as checked."""
    import webbrowser
    from datetime import datetime

    from src.db import models

    conn = models.get_connection()
    portal_list = models.list_portals(conn, active_only=False)
    portal = None
    for p in portal_list:
        if p["id"] == portal_id:
            portal = p
            break

    if not portal:
        console.print(f"[red]Portal id={portal_id} not found.[/red]")
        conn.close()
        return

    console.print(f"Opening [bold]{portal['company']}[/bold] ({portal['ats_type']})")
    console.print(f"  URL: {portal['portal_url']}")

    webbrowser.open(portal["portal_url"])
    models.update_portal(conn, portal_id, last_checked=datetime.now().isoformat())
    conn.close()

    console.print("[green]Marked as checked.[/green]")


@portals.command("stale")
def portals_stale():
    """Show portals not checked in 7+ days with pending applications."""
    from datetime import datetime

    from src.db import models

    conn = models.get_connection()
    stale = models.get_stale_portals(conn)
    conn.close()

    if not stale:
        console.print("[green]All portals are up to date.[/green]")
        return

    console.print(Panel("[bold yellow]Stale Portals[/bold yellow] — not checked in 7+ days", border_style="yellow"))

    table = Table()
    table.add_column("ID", style="dim", width=5)
    table.add_column("Company", style="bold yellow")
    table.add_column("ATS Type")
    table.add_column("Portal URL")
    table.add_column("Last Checked")
    table.add_column("Pending Apps", justify="center")

    for p in stale:
        last = "Never"
        if p["last_checked"]:
            try:
                days_ago = (datetime.now() - datetime.fromisoformat(p["last_checked"])).days
                last = f"{days_ago} days ago"
            except (ValueError, TypeError):
                last = p["last_checked"][:10]
        table.add_row(
            str(p["id"]),
            p["company"],
            p["ats_type"],
            str(p["portal_url"])[:40],
            last,
            str(p["pending_app_count"]),
        )

    console.print(table)


@cli.command()
@click.argument("job_id", type=int)
@click.option("--resume", default=None, help="Path to resume text file.")
def analyze(job_id, resume):
    """Analyze job fit for a tracked application."""
    from src.jobs.analyzer import JobAnalyzer
    from src.jobs.tracker import ApplicationTracker

    t = ApplicationTracker()
    job = t.get_job(job_id)
    t.close()

    if not job:
        console.print(f"[red]Job id={job_id} not found in tracker.[/red]")
        return

    console.print(f"\n[bold]Analyzing fit:[/bold] {job['title']} at {job['company']}")

    job_desc = (
        f"Title: {job.get('title', '')}\n"
        f"Company: {job.get('company', '')}\n"
        f"Location: {job.get('location', '')}\n"
        f"Salary: {job.get('salary_range', '')}\n"
        f"Source: {job.get('source', '')}\n"
    )
    if job.get("notes"):
        job_desc += f"\nAdditional info: {job['notes']}\n"

    resume_text = None
    if resume:
        from pathlib import Path
        rp = Path(resume)
        if rp.exists():
            resume_text = rp.read_text(encoding="utf-8", errors="replace")
        else:
            console.print(f"[red]Resume file not found: {resume}[/red]")
            return

    analyzer = JobAnalyzer()
    console.print("[dim]Analyzing...[/dim]")
    analysis = analyzer.analyze_fit(job_desc, resume_text=resume_text)

    if not analysis:
        console.print("[red]Analysis failed. Check logs.[/red]")
        return

    _display_fit_analysis(analysis)


@cli.group()
def interview():
    """Interview transcript analysis and coaching."""
    pass


@interview.command("analyze")
@click.argument("source")
@click.option("--job-title", default=None, help="Job title for context.")
@click.option("--company", default=None, help="Company name for context.")
@click.option("--kind", type=click.Choice(CANONICAL_KINDS), default=None,
              help="Override the transcript's stored kind (ID path only).")
def interview_analyze(source, job_title, company, kind):
    """Analyze an interview transcript (file path or transcript ID)."""
    from src.interviews.coach import InterviewCoach
    from src.interviews.transcripts import TranscriptLoader

    transcript_id = None
    transcript_app_id = None
    is_id_path = source.isdigit()

    if is_id_path:
        # Load from transcripts table by ID
        from src.transcripts.transcript_store import get_transcript
        from src.transcripts.transcript_parser import to_coach_turns
        from src.db import models as _models

        record = get_transcript(int(source))
        if not record:
            console.print(f"[red]Transcript ID {source} not found.[/red]")
            return
        turns = to_coach_turns(record)
        transcript_id = int(source)
        transcript_kind = kind or record.kind  # --kind overrides stored kind

        # Get application_id (not on TranscriptRecord dataclass; query directly)
        _conn = _models.get_connection()
        _row = _conn.execute(
            "SELECT application_id FROM transcripts WHERE id = ?", (transcript_id,)
        ).fetchone()
        _conn.close()
        transcript_app_id = _row["application_id"] if _row else None

        console.print(
            f"[dim]Loaded transcript #{source} "
            f"(kind: {transcript_kind}, {len(turns)} turns, {record.source} source). "
            f"Analyzing...[/dim]"
        )
    else:
        # File path flow — kind forced to generic 'interview', no context aggregation
        loader = TranscriptLoader()
        turns = loader.load_transcript(source)
        if not turns:
            console.print("[red]Failed to load transcript. Check file path and format.[/red]")
            return
        transcript_kind = "interview"
        console.print(f"[dim]Loaded {len(turns)} speaker turns. Analyzing...[/dim]")

    coach = InterviewCoach()
    analysis = coach.analyze_interview(
        turns, job_title=job_title, company=company,
        kind=transcript_kind, application_id=transcript_app_id,
    )

    if not analysis:
        console.print("[red]Analysis failed. Check logs.[/red]")
        coach.close()
        return

    # Display results
    _display_analysis(analysis)

    if not is_id_path:
        console.print(
            "\n[dim]Tip: For kind-aware analysis stored to your transcript history, "
            "use 'interview import-otter', 'import-samsung', or 'transcribe' first.[/dim]"
        )

    # Offer to save
    save = console.input("\nSave analysis and create journal entry? [bold][y][/bold]/[bold][n][/bold]: ").strip().lower()
    if save == "y":
        if is_id_path:
            # Consolidated single write — coach.save_analysis now writes to transcripts.analysis_json
            coach.save_analysis(transcript_id=transcript_id, analysis=analysis)
            save_label = f"transcript #{transcript_id}"
        else:
            # File-path flow: no write to transcripts table
            save_label = None

        # Auto-create journal entry
        from src.journal.entries import JournalManager
        manager = JournalManager()
        content = (
            f"## Interview Analysis\n\n"
            f"**Source:** {source}\n"
            f"**Company:** {company or 'N/A'}\n"
            f"**Role:** {job_title or 'N/A'}\n"
            f"**Kind:** {transcript_kind}\n"
            f"**Overall Score:** {analysis.get('overall_score', '?')}/10\n\n"
            f"### Technical Gaps\n"
            + "\n".join(f"- {g}" for g in analysis.get("technical_gaps", [])) + "\n\n"
            f"### Top Improvements\n"
            + "\n".join(f"- {imp}" for imp in analysis.get("top_improvements", []))
        )
        journal_file = manager.create_entry("interview", content, tags=["interview", "analysis"])
        if save_label:
            console.print(f"[green]Analysis saved ({save_label}). Journal entry: {journal_file}[/green]")
        else:
            console.print(f"[green]Journal entry: {journal_file}[/green]")

        # Cross-reference gaps with skill tracker
        _check_skill_gaps(analysis.get("technical_gaps", []))
    else:
        console.print("[dim]Analysis not saved.[/dim]")
    coach.close()

    coach.close()


def _display_analysis(analysis):
    """Display interview analysis with Rich formatting."""
    from rich.markdown import Markdown

    # Questions & Responses
    if analysis.get("response_quality"):
        table = Table(title="Questions & Response Quality")
        table.add_column("#", style="dim", width=3)
        table.add_column("Question")
        table.add_column("Rating", justify="center", width=8)
        table.add_column("Summary")

        for i, rq in enumerate(analysis["response_quality"], 1):
            rating = rq.get("rating", 0)
            if rating >= 4:
                rating_str = f"[green]{rating}/5[/green]"
            elif rating == 3:
                rating_str = f"[yellow]{rating}/5[/yellow]"
            else:
                rating_str = f"[red]{rating}/5[/red]"

            table.add_row(
                str(i),
                str(rq.get("question", ""))[:60],
                rating_str,
                str(rq.get("summary", ""))[:50],
            )

        console.print(table)
        console.print()

        # Detailed strengths/weaknesses for each
        for i, rq in enumerate(analysis["response_quality"], 1):
            rating = rq.get("rating", 0)
            if rating >= 4:
                border = "green"
            elif rating == 3:
                border = "yellow"
            else:
                border = "red"
            detail = (
                f"**Strengths:** {rq.get('strengths', 'N/A')}\n\n"
                f"**Weaknesses:** {rq.get('weaknesses', 'N/A')}"
            )
            console.print(Panel(
                Markdown(detail),
                title=f"Q{i}: {str(rq.get('question', ''))[:50]}",
                border_style=border,
            ))

    # Overall score
    score = analysis.get("overall_score", 0)
    if score >= 7:
        score_color = "green"
    elif score >= 5:
        score_color = "yellow"
    else:
        score_color = "red"
    console.print(f"\n[bold]Overall Score:[/bold] [{score_color}]{score}/10[/{score_color}]")
    console.print(f"  {analysis.get('overall_justification', '')}")

    # Behavioral assessment
    behavioral = analysis.get("behavioral_assessment", {})
    if behavioral:
        console.print()
        console.print(Panel(
            "\n".join(f"**{k.replace('_', ' ').title()}:** {v}" for k, v in behavioral.items()),
            title="Behavioral Assessment",
            border_style="blue",
        ))

    # Technical gaps
    gaps = analysis.get("technical_gaps", [])
    if gaps:
        console.print()
        console.print(Panel(
            "\n".join(f"- {g}" for g in gaps),
            title="Technical Gaps",
            border_style="red",
        ))

    # Top improvements
    improvements = analysis.get("top_improvements", [])
    if improvements:
        console.print()
        console.print(Panel(
            "\n".join(f"{i+1}. {imp}" for i, imp in enumerate(improvements)),
            title="Top Improvements",
            border_style="yellow",
        ))

    # Practice questions
    questions = analysis.get("practice_questions", [])
    if questions:
        console.print()
        console.print(Panel(
            "\n".join(f"{i+1}. {q}" for i, q in enumerate(questions)),
            title="Practice Questions",
            border_style="cyan",
        ))


def _check_skill_gaps(technical_gaps):
    """Cross-reference technical gaps with skill tracker."""
    if not technical_gaps:
        return

    from src.skills.tracker import SkillTracker

    tracker = SkillTracker()
    tracker.seed_defaults()
    all_skills = tracker.get_all_skills()
    tracker.close()

    skill_names = {s["name"].lower(): s for s in all_skills}
    flagged = []
    for gap in technical_gaps:
        gap_lower = gap.lower()
        for name, skill in skill_names.items():
            if name in gap_lower or gap_lower in name:
                flagged.append((gap, skill))
                break

    if flagged:
        console.print()
        console.print("[bold]Skill Tracker Matches:[/bold]")
        for gap_text, skill in flagged:
            console.print(
                f"  [yellow]>[/yellow] '{gap_text}' matches tracked skill "
                f"'{skill['name']}' (level {skill['current_level']}/{skill['target_level']})"
            )


def _prompt_link_application():
    """Prompt user to link a transcript to an application. Returns app_id or None."""
    from src.db import models
    conn = models.get_connection()
    try:
        rows = conn.execute(
            "SELECT id, title, company FROM applications ORDER BY id DESC LIMIT 10"
        ).fetchall()
    finally:
        conn.close()

    if not rows:
        return None

    console.print("\n[bold]Link to an application?[/bold]")
    for i, row in enumerate(rows, 1):
        console.print(f"  {i}. {row['title']} at {row['company']}")
    console.print("  skip. Don't link")

    choice = console.input("\nChoice: ").strip().lower()
    if choice == "skip" or not choice:
        return None
    try:
        idx = int(choice) - 1
        if 0 <= idx < len(rows):
            return rows[idx]["id"]
    except ValueError:
        pass
    return None


@interview.command("import-samsung")
@click.argument("path")
@click.option("--kind", type=click.Choice(CANONICAL_KINDS), default="interview",
              help="Transcript kind (recruiter_prep, technical, etc.)")
def interview_import_samsung(path, kind):
    """Import a Samsung call recording transcript or directory."""
    from src.transcripts.samsung_importer import import_samsung
    from src.transcripts.whisper_transcriber import transcribe
    from src.transcripts.transcript_store import store_transcript

    record = import_samsung(path)

    # If audio-only, offer Whisper transcription
    if record.raw_metadata.get("needs_whisper"):
        if record.audio_path:
            do_whisper = console.input("No transcript found. Transcribe audio with Whisper? [bold][y][/bold]/[bold][n][/bold]: ").strip().lower()
            if do_whisper == "y":
                console.print("[dim]Transcribing with Whisper (base model)...[/dim]")
                try:
                    record = transcribe(record.audio_path)
                except RuntimeError as e:
                    console.print(f"[red]{e}[/red]")
                    return
            else:
                console.print("[dim]Skipped transcription.[/dim]")
                return
        else:
            console.print("[red]No transcript or audio file found.[/red]")
            return

    record.kind = kind
    # Show summary
    word_count = len(record.full_text.split())
    console.print(f"[green]Imported:[/green] {record.source} | {len(record.segments)} segments | {word_count} words | {record.duration_seconds:.0f}s")

    # Application linking
    app_id = _prompt_link_application()
    row_id = store_transcript(record, application_id=app_id)
    console.print(f"[green]Saved as transcript #{row_id} (kind: {kind})[/green]")


@interview.command("import-otter")
@click.argument("file")
@click.option("--kind", type=click.Choice(CANONICAL_KINDS), default="interview",
              help="Transcript kind (recruiter_prep, technical, etc.)")
def interview_import_otter(file, kind):
    """Import an Otter.ai transcript (.txt or .srt)."""
    from src.transcripts.otter_importer import import_otter
    from src.transcripts.transcript_store import store_transcript

    record = import_otter(file)
    record.kind = kind
    word_count = len(record.full_text.split())
    console.print(f"[green]Imported:[/green] {record.source} | {len(record.segments)} segments | {word_count} words | {record.duration_seconds:.0f}s")

    app_id = _prompt_link_application()
    row_id = store_transcript(record, application_id=app_id)
    console.print(f"[green]Saved as transcript #{row_id} (kind: {kind})[/green]")


@interview.command("transcribe")
@click.argument("audio_file")
@click.option("--model", default="base", help="Whisper model size (tiny/base/small/medium/large-v3/turbo).")
@click.option("--kind", type=click.Choice(CANONICAL_KINDS), default="interview",
              help="Transcript kind (recruiter_prep, technical, etc.)")
def interview_transcribe(audio_file, model, kind):
    """Transcribe an audio file with local Whisper."""
    from src.transcripts.whisper_transcriber import transcribe
    from src.transcripts.transcript_store import store_transcript

    console.print(f"[dim]Transcribing with Whisper ({model} model)...[/dim]")
    try:
        record = transcribe(audio_file, model_size=model)
    except RuntimeError as e:
        console.print(f"[red]{e}[/red]")
        return

    record.kind = kind
    word_count = len(record.full_text.split())
    console.print(f"[green]Done:[/green] {len(record.segments)} segments | {word_count} words | {record.duration_seconds:.0f}s | Language: {record.language}")

    app_id = _prompt_link_application()
    row_id = store_transcript(record, application_id=app_id)
    console.print(f"[green]Saved as transcript #{row_id} (kind: {kind})[/green]")


@interview.command("watch")
@click.option("--model", default="base", help="Whisper model for audio files.")
@click.option("--kind", type=click.Choice(CANONICAL_KINDS), default="interview",
              help="Transcript kind applied to every auto-imported file in this session.")
def interview_watch(model, kind):
    """Watch data/transcripts/ for new files and auto-import."""
    from src.transcripts.watch_folder import watch
    watch(model_size=model, kind=kind)


@interview.command("list")
def interview_list():
    """List all imported transcripts."""
    from src.transcripts.transcript_store import list_transcripts

    transcripts = list_transcripts()
    if not transcripts:
        console.print("[dim]No transcripts imported yet.[/dim]")
        return

    table = Table(title="Imported Transcripts")
    table.add_column("ID", style="dim", width=4)
    table.add_column("Date", width=12)
    table.add_column("Source", width=8)
    table.add_column("Duration", width=8)
    table.add_column("Application", width=20)
    table.add_column("Preview", max_width=50)

    for t in transcripts:
        date = t["imported_at"][:10] if t.get("imported_at") else ""
        duration = f"{t['duration_seconds']:.0f}s" if t.get("duration_seconds") else ""
        app = t.get("company") or ""
        if t.get("app_title"):
            app = f"{app} ({t['app_title']})" if app else t["app_title"]
        preview = (t.get("preview") or "")[:50]
        analyzed = " *" if t.get("analyzed_at") else ""

        table.add_row(str(t["id"]), date, t["source"] + analyzed, duration, app, preview)

    console.print(table)


@interview.command("mock")
@click.option("--questions", "num_questions", default=5, help="Number of questions (default 5).")
def interview_mock(num_questions):
    """Start an interactive mock interview session."""
    from rich.markdown import Markdown

    from src.interviews.coach import InterviewCoach

    console.print(Panel(
        "Describe the target role for your mock interview.\n"
        "Include: job title, technologies, responsibilities.",
        title="Mock Interview Setup",
        border_style="cyan",
    ))

    role_desc = console.input("\nRole description: ").strip()
    if not role_desc:
        console.print("[yellow]No role description provided, cancelled.[/yellow]")
        return

    console.print(f"\n[bold]Starting mock interview ({num_questions} questions)...[/bold]")
    console.print("[dim]Answer each question as you would in a real interview.[/dim]\n")

    coach = InterviewCoach()

    def rich_output(text):
        console.print(text)

    def rich_input(prompt):
        return console.input(f"[bold]{prompt}[/bold]")

    assessment = coach.mock_interview(
        role_description=role_desc,
        num_questions=num_questions,
        input_fn=rich_input,
        output_fn=rich_output,
    )
    coach.close()

    if not assessment:
        console.print("[red]Mock interview failed. Check logs.[/red]")
        return

    # Display final assessment
    console.print("\n")
    console.rule("[bold]Final Assessment[/bold]")

    score = assessment.get("overall_score", 0)
    if score >= 7:
        score_color = "green"
    elif score >= 5:
        score_color = "yellow"
    else:
        score_color = "red"

    console.print(f"\n[bold]Overall Score:[/bold] [{score_color}]{score}/10[/{score_color}]")
    console.print(f"  {assessment.get('overall_justification', '')}")

    gaps = assessment.get("technical_gaps", [])
    if gaps:
        console.print(Panel(
            "\n".join(f"- {g}" for g in gaps),
            title="Technical Gaps",
            border_style="red",
        ))

    improvements = assessment.get("top_improvements", [])
    if improvements:
        console.print(Panel(
            "\n".join(f"{i+1}. {imp}" for i, imp in enumerate(improvements)),
            title="Top Improvements",
            border_style="yellow",
        ))

    practice = assessment.get("practice_questions", [])
    if practice:
        console.print(Panel(
            "\n".join(f"{i+1}. {q}" for i, q in enumerate(practice)),
            title="Follow-up Practice Questions",
            border_style="cyan",
        ))

    # Offer to save
    save = console.input("\nSave mock interview as journal entry? [bold][y][/bold]/[bold][n][/bold]: ").strip().lower()
    if save == "y":
        from src.journal.entries import JournalManager

        qa_text = "\n\n".join(
            f"**Q{i+1}:** {qa['question']}\n"
            f"**A:** {qa['answer']}\n"
            f"**Rating:** {qa['evaluation'].get('rating', '?')}/5"
            for i, qa in enumerate(assessment.get("qa_pairs", []))
        )
        content = (
            f"## Mock Interview\n\n"
            f"**Role:** {role_desc}\n"
            f"**Score:** {score}/10\n\n"
            f"### Q&A\n\n{qa_text}\n\n"
            f"### Technical Gaps\n"
            + "\n".join(f"- {g}" for g in gaps) + "\n\n"
            f"### Improvements\n"
            + "\n".join(f"- {imp}" for imp in improvements)
        )
        manager = JournalManager()
        filename = manager.create_entry("interview", content, tags=["mock-interview"])
        console.print(f"[green]Saved journal entry: {filename}[/green]")

        _check_skill_gaps(gaps)
    else:
        console.print("[dim]Not saved.[/dim]")


@interview.command("history")
def interview_history():
    """Show past interview analyses with scores and gaps."""
    from src.interviews.coach import InterviewCoach

    coach = InterviewCoach()
    analyses = coach.get_all_analyses()
    coach.close()

    if not analyses:
        console.print("[yellow]No interview analyses found. Run 'interview analyze' first.[/yellow]")
        return

    table = Table(title=f"Interview History ({len(analyses)} analyses)")
    table.add_column("ID", style="dim", width=4)
    table.add_column("Date", width=12)
    table.add_column("Company")
    table.add_column("Role")
    table.add_column("Score", justify="center", width=7)
    table.add_column("Top Gaps")

    for a in analyses:
        analysis = a.get("analysis", {})
        score = analysis.get("overall_score", 0)
        if score >= 7:
            score_str = f"[green]{score}/10[/green]"
        elif score >= 5:
            score_str = f"[yellow]{score}/10[/yellow]"
        else:
            score_str = f"[red]{score}/10[/red]"

        gaps = ", ".join(analysis.get("technical_gaps", [])[:3])

        table.add_row(
            str(a.get("id", "")),
            str(a.get("analyzed_at", ""))[:10],
            a.get("company", "") or "-",
            a.get("role", "") or "-",
            score_str,
            gaps[:50] if gaps else "-",
        )

    console.print(table)


@interview.command("compare")
def interview_compare():
    """Compare all stored analyses to identify trends."""
    from src.interviews.coach import InterviewCoach

    coach = InterviewCoach()
    analyses = coach.get_all_analyses()

    if len(analyses) < 2:
        console.print(f"[yellow]Need at least 2 analyses to compare (found {len(analyses)}). "
                       "Run more 'interview analyze' sessions first.[/yellow]")
        coach.close()
        return

    console.print(f"[dim]Comparing {len(analyses)} analyses...[/dim]")
    comparison = coach.compare_interviews(analyses)
    coach.close()

    if not comparison:
        console.print("[red]Comparison failed. Check logs.[/red]")
        return

    # Trajectory
    trajectory = comparison.get("trajectory", "unknown")
    traj_colors = {"improving": "green", "plateauing": "yellow", "declining": "red"}
    traj_color = traj_colors.get(trajectory, "white")
    console.print(f"\n[bold]Trajectory:[/bold] [{traj_color}]{trajectory.upper()}[/{traj_color}]")
    console.print(f"  {comparison.get('trajectory_explanation', '')}")

    # Recurring weak topics
    weak = comparison.get("recurring_weak_topics", [])
    if weak:
        console.print(Panel(
            "\n".join(f"- {w}" for w in weak),
            title="Recurring Weak Topics",
            border_style="red",
        ))

    # Improved skills
    improved = comparison.get("improved_skills", [])
    if improved:
        console.print(Panel(
            "\n".join(f"- {s}" for s in improved),
            title="Improved Skills",
            border_style="green",
        ))

    # Persistent gaps
    persistent = comparison.get("persistent_gaps", [])
    if persistent:
        console.print(Panel(
            "\n".join(f"- {g}" for g in persistent),
            title="Persistent Gaps (Need Focused Study)",
            border_style="yellow",
        ))

    # Recommendations
    recs = comparison.get("recommendations", [])
    if recs:
        console.print(Panel(
            "\n".join(f"{i+1}. {r}" for i, r in enumerate(recs)),
            title="Recommendations",
            border_style="cyan",
        ))


# --- Phase 6: Dashboard, Morning Scan, Daily Summary, Quick Entry, Status ---


DASHBOARD_MENU = """\
[bold][1][/bold] Morning Scan \u2014 Gmail + Job Search
[bold][2][/bold] Journal \u2014 New entry, recent, insights
[bold][3][/bold] Skills \u2014 Inventory, update, roadmap
[bold][4][/bold] Applications \u2014 Pipeline, update, stats
[bold][5][/bold] Interviews \u2014 Analyze, history, mock
[bold][6][/bold] Calendar \u2014 View availability
[bold][7][/bold] Daily Summary \u2014 AI end-of-day recap
[bold][8][/bold] Inbox \u2014 Threaded email dashboard
[bold][q][/bold] Quit"""


@cli.command()
@click.pass_context
def dashboard(ctx):
    """Interactive dashboard \u2014 main menu for all CareerPilot features."""
    from datetime import datetime

    from src.db import models

    # Startup banner
    conn = models.get_connection()
    last_scan = models.get_kv(conn, "last_morning_scan") or "never"
    conn.close()

    console.print()
    console.print(Panel(
        f"[bold cyan]CareerPilot v1.0[/bold cyan]\n"
        f"Today: {datetime.now().strftime('%A, %B %d, %Y')}\n"
        f"Last morning scan: {last_scan}",
        border_style="cyan",
    ))

    while True:
        console.print()
        console.print(Panel(DASHBOARD_MENU, title="CareerPilot Dashboard", border_style="cyan"))

        choice = console.input("\nSelect: ").strip().lower()

        try:
            if choice == "1":
                ctx.invoke(morning)
            elif choice == "2":
                _journal_submenu(ctx)
            elif choice == "3":
                _skills_submenu(ctx)
            elif choice == "4":
                _tracker_submenu(ctx)
            elif choice == "5":
                _interviews_submenu(ctx)
            elif choice == "6":
                ctx.invoke(calendar, days=5)
            elif choice == "7":
                ctx.invoke(daily)
            elif choice == "8":
                ctx.invoke(inbox)
            elif choice == "q":
                console.print("[dim]Goodbye.[/dim]")
                return
            else:
                console.print("[red]Invalid choice.[/red]")
        except SystemExit:
            pass  # Click commands may call sys.exit on completion


def _journal_submenu(ctx):
    """Journal submenu."""
    console.print("\n[bold]Journal:[/bold]  [1] New  [2] List  [3] Insights  [4] Search  [b] Back")
    c = console.input("Select: ").strip().lower()
    if c == "1":
        ctx.invoke(journal_new, entry_type="daily", mood=None, time_spent=None)
    elif c == "2":
        ctx.invoke(journal_list, days=30, entry_type=None)
    elif c == "3":
        ctx.invoke(journal_insights)
    elif c == "4":
        kw = console.input("Search keyword: ").strip()
        if kw:
            ctx.invoke(journal_search, keyword=kw)


def _skills_submenu(ctx):
    """Skills submenu."""
    console.print("\n[bold]Skills:[/bold]  [1] Inventory  [2] Update  [3] Roadmap  [b] Back")
    c = console.input("Select: ").strip().lower()
    if c == "1":
        ctx.invoke(skills)
    elif c == "2":
        name = console.input("Skill name: ").strip()
        level = console.input("New level (1-5): ").strip()
        if name and level.isdigit():
            ctx.invoke(skills_update, name=name, level=int(level))
    elif c == "3":
        ctx.invoke(roadmap, hours=15)


def _tracker_submenu(ctx):
    """Application tracker submenu."""
    console.print("\n[bold]Applications:[/bold]  [1] Pipeline  [2] Update  [3] Stats  [b] Back")
    c = console.input("Select: ").strip().lower()
    if c == "1":
        ctx.invoke(tracker_show)
    elif c == "2":
        jid = console.input("Job ID: ").strip()
        if jid.isdigit():
            ctx.invoke(tracker_update, job_id=int(jid))
    elif c == "3":
        ctx.invoke(tracker_stats)


def _interviews_submenu(ctx):
    """Interviews submenu."""
    console.print("\n[bold]Interviews:[/bold]  [1] Analyze  [2] History  [3] Compare  [4] Mock  [b] Back")
    c = console.input("Select: ").strip().lower()
    if c == "1":
        fp = console.input("Transcript file path: ").strip()
        if fp:
            ctx.invoke(interview_analyze, filepath=fp, job_title=None, company=None)
    elif c == "2":
        ctx.invoke(interview_history)
    elif c == "3":
        ctx.invoke(interview_compare)
    elif c == "4":
        ctx.invoke(interview_mock, num_questions=5)


@cli.command()
def morning():
    """Morning scan \u2014 Gmail (24h) + quick job search + pending apps."""
    from datetime import datetime

    from src.db import models

    console.rule("[bold cyan]Morning Scan[/bold cyan]")
    console.print()

    # --- Gmail scan (last 24 hours) ---
    email_count = 0
    try:
        from src.gmail.scanner import GmailScanner

        scanner = GmailScanner()
        scanner.authenticate()
        results = scanner.scan_inbox(days_back=1)
        email_count = len(results) if results else 0

        if results:
            table = Table(title=f"Recruiter Emails ({email_count})")
            table.add_column("#", style="dim")
            table.add_column("Category", style="bold")
            table.add_column("From")
            table.add_column("Subject")
            table.add_column("Company")

            for i, r in enumerate(results, 1):
                color = CATEGORY_COLORS.get(r["category"], "white")
                table.add_row(
                    str(i),
                    f"[{color}]{r['category']}[/{color}]",
                    r["sender"][:30],
                    r["subject"][:40],
                    r["company"],
                )
            console.print(table)
        else:
            console.print("[dim]No new recruiter emails in the last 24 hours.[/dim]")
    except Exception:
        console.print("[yellow]Gmail scan skipped (auth failed or unavailable).[/yellow]")

    console.print()

    # --- Quick job search (3 default profiles) ---
    job_count = 0
    try:
        from src.jobs.searcher import JobSearcher

        searcher = JobSearcher()
        default_profiles = ["sysadmin_local", "syseng_local", "contract_infra"]
        console.print("[dim]Searching default profiles...[/dim]")
        jobs = searcher.run_profiles(default_profiles)
        job_count = len(jobs)

        if jobs:
            table = Table(title=f"New Job Listings ({job_count})")
            table.add_column("#", style="dim", width=4)
            table.add_column("Title", style="bold")
            table.add_column("Company")
            table.add_column("Location")
            table.add_column("Source", width=8)

            for i, j in enumerate(jobs[:15], 1):
                table.add_row(
                    str(i),
                    str(j.get("title", ""))[:40],
                    str(j.get("company", ""))[:25],
                    str(j.get("location", ""))[:20],
                    j.get("source", ""),
                )
            console.print(table)
            if job_count > 15:
                console.print(f"[dim]... and {job_count - 15} more. Run 'search' for full results.[/dim]")
        else:
            console.print("[dim]No new job listings found.[/dim]")
    except Exception:
        console.print("[yellow]Job search skipped (MCP unavailable).[/yellow]")

    console.print()

    # --- Pending applications ---
    try:
        from src.jobs.tracker import ApplicationTracker

        t = ApplicationTracker()
        stats = t.get_stats()
        t.close()

        applied = stats["by_status"].get("applied", 0)
        interviewing = stats["by_status"].get("interview", 0) + stats["by_status"].get("phone_screen", 0)
        console.print(
            f"[bold]Pipeline:[/bold] {applied} awaiting response, "
            f"{interviewing} interviewing, {stats['total']} total tracked"
        )
    except Exception:
        pass

    # --- Inbox digest ---
    inbox_digest_text = ""
    try:
        from src.gmail.auth import get_gmail_service
        from src.gmail.dashboard import EmailDashboard
        from src.gmail.thread_actions import ThreadActions

        svc = get_gmail_service()
        dash = EmailDashboard(svc)
        thread_actions = ThreadActions(svc)

        inbox_threads = dash.fetch_threads(max_results=50)
        awaiting = 0
        stale_threads = []
        for t in inbox_threads:
            snoozed, _ = thread_actions.is_snoozed(t["thread_id"])
            if snoozed:
                continue
            si = dash.classify_thread_status(t)
            if si["status"] == "awaiting_reply":
                awaiting += 1
                if si["is_stale"]:
                    stale_threads.append(t)

        parts = [f"{awaiting} threads awaiting reply"]
        if stale_threads:
            parts.append(f"[red]{len(stale_threads)} stale[/red]")
        inbox_digest_text = f"Inbox: {', '.join(parts)}"
        console.print(f"[bold]{inbox_digest_text}[/bold]")

        if stale_threads:
            for st in stale_threads[:5]:
                si = dash.classify_thread_status(st)
                days = int(si["hours_since_last"] / 24)
                console.print(
                    f"  [red]Stale:[/red] {st['subject'][:60]} ({days} day{'s' if days != 1 else ''})"
                )
    except Exception:
        pass

    # --- Portal check reminders ---
    # Note: `datetime` and `models` are already imported at the top of morning()
    try:
        _conn = models.get_connection()
        # Get all active portals with pending app counts
        _portals = _conn.execute(
            "SELECT p.*, COUNT(a.id) AS pending_app_count "
            "FROM ats_portals p "
            "LEFT JOIN applications a ON a.portal_id = p.id "
            "  AND a.status NOT IN ('withdrawn', 'rejected', 'ghosted') "
            "WHERE p.active = 1 "
            "GROUP BY p.id "
            "HAVING pending_app_count > 0 "
            "ORDER BY p.last_checked ASC",
        ).fetchall()

        if _portals:
            console.print()
            console.print("[bold]📋 Portal Check Reminders:[/bold]")
            for _p in _portals:
                _p = dict(_p)
                if _p["last_checked"]:
                    try:
                        _last = datetime.fromisoformat(_p["last_checked"])
                        _days_ago = (datetime.now() - _last).days
                        if _days_ago == 0:
                            _time_str = "checked today"
                        elif _days_ago == 1:
                            _time_str = "last checked yesterday"
                        else:
                            _time_str = f"last checked {_days_ago} days ago"
                    except (ValueError, TypeError):
                        _time_str = "unknown last check"
                        _days_ago = 999
                else:
                    _time_str = "never checked"
                    _days_ago = 999

                _app_label = "application" if _p["pending_app_count"] == 1 else "applications"

                if _days_ago >= 7:
                    console.print(
                        f"  [yellow]⚠ {_p['company']} ({_p['ats_type']})[/yellow] — "
                        f"{_time_str}, {_p['pending_app_count']} pending {_app_label}"
                    )
                else:
                    console.print(
                        f"  [green]✅ {_p['company']}[/green] — {_time_str}"
                    )

        _conn.close()
    except Exception:
        pass

    # --- Contact follow-ups ---
    try:
        conn = models.get_connection()
        followups = models.get_followup_due(conn)
        stale = models.get_stale_contacts(conn)
        active_warm = models.list_contacts(conn, status="active") + models.list_contacts(conn, status="warm")
        conn.close()

        if followups or active_warm:
            console.print()
            console.print("[bold]Contact Follow-ups:[/bold]")

            # Due follow-ups first
            for r in followups:
                company_str = f" ({r['company']})" if r.get("company") else ""
                console.print(
                    f"  [cyan]📅[/cyan] {r['name']}{company_str} "
                    f"— scheduled follow-up"
                )

            # Then stale contacts
            stale_ids = {r["id"] for r in stale}
            shown_ids = {r["id"] for r in followups}
            for r in sorted(active_warm, key=lambda x: x.get("last_contact") or ""):
                if r["id"] in shown_ids:
                    continue
                shown_ids.add(r["id"])
                company_str = f" ({r['company']})" if r.get("company") else ""
                if r["id"] in stale_ids:
                    days_ago = (datetime.now() - datetime.fromisoformat(r["last_contact"])).days
                    console.print(
                        f"  [yellow]⚠[/yellow] {r['name']}{company_str} "
                        f"— last contact {days_ago} days ago"
                    )
                elif r.get("last_contact"):
                    days_ago = (datetime.now() - datetime.fromisoformat(r["last_contact"])).days
                    console.print(
                        f"  [green]✅[/green] {r['name']}{company_str} "
                        f"— contacted {days_ago} days ago"
                    )
                else:
                    console.print(
                        f"  [dim]●[/dim] {r['name']}{company_str} — no contact logged"
                    )
    except Exception:
        pass

    # --- Skill study focus ---
    try:
        conn = models.get_connection()
        plan = models.get_study_plan(conn)
        conn.close()

        if plan:
            console.print()
            console.print("[bold]Skill Focus This Week:[/bold]")
            for p in plan[:3]:
                target = p.get("target_hours") or 0
                logged = p.get("study_hours_logged") or 0
                pct = int(logged / target * 100) if target > 0 else 0
                bar_filled = int(pct / 10)
                bar = "\u2588" * bar_filled + "\u2591" * (10 - bar_filled)
                console.print(
                    f"  {p.get('priority_rank', '?')}. {p['skill_name']} "
                    f"({logged}/{target} hrs)  {bar}  {pct}%"
                )
    except Exception:
        pass

    # Record scan timestamp
    try:
        conn = models.get_connection()
        models.set_kv(conn, "last_morning_scan", datetime.now().strftime("%Y-%m-%d %H:%M"))
        conn.close()
    except Exception:
        pass

    console.print()
    console.print(
        f"[bold]Summary:[/bold] {email_count} emails, {job_count} jobs found."
    )


# ═══════════════════════════════════════════════════════════════
# Contacts
# ═══════════════════════════════════════════════════════════════


CONTACT_TYPES = [
    "recruiter", "hiring_manager", "networking", "reference",
    "colleague", "mentor", "school_contact", "other",
]

CONTACT_SOURCES = [
    "staffing_agency", "linkedin", "meetup", "referral",
    "conference", "cold_outreach", "job_application", "email_import", "other",
]

CONTACT_METHODS = ["email", "phone", "linkedin", "in_person", "text"]


def _get_contact_manager():
    """Return a configured ContactManager (Supabase-backed)."""
    from src.db.contacts import ContactManager, ContactManagerNotConfiguredError
    try:
        return ContactManager()
    except ContactManagerNotConfiguredError as exc:
        console.print(f"[red]ContactManager not configured:[/red] {exc}")
        raise click.Abort() from exc


def _contacts_table(all_contacts, title="Contacts"):
    """Build a Rich table for contacts list."""
    from datetime import datetime

    table = Table(title=title)
    table.add_column("ID", style="dim", width=20)
    table.add_column("Name", style="bold")
    table.add_column("Company")
    table.add_column("Type")
    table.add_column("Specialization")
    table.add_column("Last Contact")
    table.add_column("Status")
    table.add_column("Tags")

    now = datetime.now()
    for r in all_contacts:
        status = r.get("relationship_status", "new")
        style = "dim"
        lcd = r.get("last_contact_date") or r.get("last_contact")
        if status in ("active", "warm"):
            if lcd:
                try:
                    dt = datetime.fromisoformat(str(lcd).replace("Z", "+00:00"))
                    dt = dt.replace(tzinfo=None) if dt.tzinfo else dt
                    days = (now - dt).days
                    style = "red" if days >= 14 else ("yellow" if days >= 7 else "green")
                except (ValueError, TypeError):
                    style = "green"
            else:
                style = "green"
        elif status == "do_not_contact":
            style = "dim red"

        last_contact_disp = str(lcd)[:10] if lcd else ""
        tags_raw = r.get("tags") or ""
        tags_disp = ", ".join(tags_raw) if isinstance(tags_raw, list) else tags_raw

        table.add_row(
            str(r["id"]),
            f"[{style}]{r['name']}[/{style}]",
            r.get("company", "") or "",
            r.get("contact_type", "") or "",
            r.get("specialization", "") or "",
            last_contact_disp,
            f"[{style}]{status}[/{style}]",
            tags_disp,
        )

    return table


@cli.group(invoke_without_command=True)
@click.pass_context
def contacts(ctx):
    """Professional contacts manager."""
    if ctx.invoked_subcommand is not None:
        return

    mgr = _get_contact_manager()
    all_contacts = mgr.list_contacts()

    if not all_contacts:
        console.print("[dim]No contacts tracked yet. Use 'contacts add' to add one.[/dim]")
        return

    console.print(_contacts_table(all_contacts))


@contacts.command("add")
def contacts_add():
    """Add a new contact (interactive wizard)."""
    console.print("[bold]Add Contact[/bold]")
    console.print()

    name = click.prompt("  Name")

    console.print("  Contact type:")
    for i, t in enumerate(CONTACT_TYPES, 1):
        console.print(f"    {i}. {t}")
    type_choice = click.prompt("  Type", type=int, default=1)
    contact_type = CONTACT_TYPES[min(type_choice, len(CONTACT_TYPES)) - 1]

    company = click.prompt("  Company", default="", show_default=False) or None
    title = click.prompt("  Their title", default="", show_default=False) or None
    email = click.prompt("  Email", default="", show_default=False) or None
    phone = click.prompt("  Phone", default="", show_default=False) or None
    linkedin = click.prompt("  LinkedIn URL", default="", show_default=False) or None

    console.print("  Source — where you met them:")
    for i, s in enumerate(CONTACT_SOURCES, 1):
        console.print(f"    {i}. {s}")
    source_choice = click.prompt("  Source", type=int, default=1)
    source = CONTACT_SOURCES[min(source_choice, len(CONTACT_SOURCES)) - 1]

    default_spec = "Infrastructure" if contact_type == "recruiter" else ""
    spec = click.prompt("  Specialization", default=default_spec, show_default=bool(default_spec)) or None
    tags_raw = click.prompt("  Tags (comma-separated)", default="", show_default=False) or None
    tags = [t.strip() for t in tags_raw.split(",") if t.strip()] if tags_raw else None
    notes = click.prompt("  Notes", default="", show_default=False) or None

    mgr = _get_contact_manager()
    cid = mgr.add_contact(
        name, contact_type,
        company=company, title=title, email=email, phone=phone,
        linkedin_url=linkedin, source=source, specialization=spec,
        tags=tags, notes=notes,
    )

    company_str = f" ({company})" if company else ""
    console.print(f"\n  [green]Added contact {cid}: {name}{company_str} [{contact_type}][/green]")


@contacts.command("show")
@click.argument("contact_id", type=str)
def contacts_show(contact_id):
    """Show detailed contact info with interaction history."""
    from src.db import models

    mgr = _get_contact_manager()
    c = mgr.get_contact(contact_id)
    if not c:
        console.print(f"[red]Contact {contact_id} not found.[/red]")
        return

    # Build detail panel
    lines = [f"[bold]{c['name']}[/bold]"]
    if c.get("company"):
        lines.append(f"Company: {c['company']}")
    if c.get("title"):
        lines.append(f"Title: {c['title']}")
    lines.append(f"Type: {c.get('contact_type', 'recruiter')}")
    lines.append(f"Status: {c.get('relationship_status', 'new')}")
    if c.get("email"):
        lines.append(f"Email: {c['email']}")
    if c.get("phone"):
        lines.append(f"Phone: {c['phone']}")
    if c.get("linkedin_url"):
        lines.append(f"LinkedIn: {c['linkedin_url']}")
    if c.get("specialization"):
        lines.append(f"Specialization: {c['specialization']}")
    if c.get("source"):
        lines.append(f"Source: {c['source']}")
    tags_raw = c.get("tags")
    if tags_raw:
        tags_disp = ", ".join(tags_raw) if isinstance(tags_raw, list) else tags_raw
        lines.append(f"Tags: {tags_disp}")
    lcd = c.get("last_contact_date")
    if lcd:
        lines.append(f"Last Contact: {str(lcd)[:10]} via {c.get('contact_method', 'N/A')}")
    if c.get("next_followup"):
        lines.append(f"Next Follow-up: {c['next_followup']}")
    if c.get("notes"):
        lines.append(f"\nNotes:\n{c['notes']}")

    console.print(Panel("\n".join(lines), title=f"Contact {contact_id}"))

    # Interaction history (local SQLite log)
    conn = models.get_connection()
    interactions = models.get_contact_interactions(conn, contact_id)
    conn.close()
    if interactions:
        console.print(f"\n[bold]Recent Interactions ({len(interactions)}):[/bold]")
        for i in interactions[:10]:
            direction = "->" if i.get("direction") == "outbound" else "<-"
            console.print(
                f"  {i['created_at'][:10]} {direction} {i['interaction_type']}: "
                f"{i.get('subject', 'N/A')}"
            )
            if i.get("summary"):
                console.print(f"    {i['summary'][:80]}")

    # Submitted roles (local SQLite log)
    conn = models.get_connection()
    roles = models.get_submitted_roles(conn, contact_uuid=contact_id)
    conn.close()
    if roles:
        console.print(f"\n[bold]Submitted Roles ({len(roles)}):[/bold]")
        for role in roles:
            console.print(
                f"  {role['role_title']} at {role['company']} [{role['status']}]"
            )
            if role.get("pay_rate"):
                console.print(f"    Pay: {role['pay_rate']}")


@contacts.command("edit")
@click.argument("contact_id", type=str)
def contacts_edit(contact_id):
    """Update contact fields interactively."""
    mgr = _get_contact_manager()
    c = mgr.get_contact(contact_id)
    if not c:
        console.print(f"[red]Contact {contact_id} not found.[/red]")
        return

    console.print(f"[bold]Edit {c['name']}[/bold] (press Enter to keep current value)")

    updates = {}
    for field in ["name", "company", "title", "email", "phone", "linkedin_url",
                   "specialization", "notes"]:
        current = c.get(field, "") or ""
        val = click.prompt(f"  {field}", default=current, show_default=True)
        if val != current:
            updates[field] = val if val else None

    # Tags (display as comma-separated, store as list)
    tags_current = c.get("tags") or []
    tags_str = ", ".join(tags_current) if isinstance(tags_current, list) else (tags_current or "")
    tags_val = click.prompt("  tags", default=tags_str, show_default=True)
    if tags_val != tags_str:
        updates["tags"] = [t.strip() for t in tags_val.split(",") if t.strip()] if tags_val else []

    # Status
    current_status = c.get("relationship_status", "new")
    statuses = ["new", "active", "warm", "cold", "do_not_contact"]
    console.print(f"  Status (current: {current_status}):")
    for i, s in enumerate(statuses, 1):
        console.print(f"    {i}. {s}")
    status_choice = click.prompt("  Status", default="", show_default=False)
    if status_choice and status_choice.isdigit():
        idx = int(status_choice) - 1
        if 0 <= idx < len(statuses) and statuses[idx] != current_status:
            updates["relationship_status"] = statuses[idx]

    if updates:
        mgr.update_contact(contact_id, **updates)
        console.print(f"  [green]Updated {', '.join(updates.keys())}.[/green]")
    else:
        console.print("  [dim]No changes made.[/dim]")


@contacts.command("log")
@click.argument("contact_id", type=str)
def contacts_log(contact_id):
    """Log an interaction with a contact."""
    from datetime import datetime

    from src.db import models

    mgr = _get_contact_manager()
    c = mgr.get_contact(contact_id)
    if not c:
        console.print(f"[red]Contact {contact_id} not found.[/red]")
        return

    company_str = f" ({c['company']})" if c.get("company") else ""
    console.print(f"[bold]Log contact with {c['name']}{company_str}[/bold]")
    method = click.prompt(
        "  Contact method",
        type=click.Choice(CONTACT_METHODS),
    )
    subject = click.prompt("  Subject / note (optional)", default="", show_default=False) or None
    followup = click.prompt(
        "  Next follow-up date (YYYY-MM-DD, optional)",
        default="", show_default=False,
    ) or None

    now_iso = datetime.now().isoformat()

    # Write interaction log to local SQLite
    conn = models.get_connection()
    models.add_contact_interaction(conn, contact_id, method, subject=subject)
    conn.close()

    # Update Supabase contact record
    updates: dict = {"last_contact_date": now_iso, "contact_method": method}
    if followup:
        updates["next_followup"] = followup
    mgr.update_contact(contact_id, **updates)

    console.print(f"  [green]Contact logged for {c['name']}.[/green]")


@contacts.command("search")
@click.argument("query")
def contacts_search(query):
    """Search contacts by name, company, email, or notes."""
    mgr = _get_contact_manager()
    results = mgr.search_contacts(query)

    if not results:
        console.print(f"[dim]No contacts matching '{query}'.[/dim]")
        return

    console.print(_contacts_table(results, title=f"Search: '{query}'"))


@contacts.command("stale")
def contacts_stale():
    """Show contacts not contacted in 14+ days (active/warm only)."""
    from datetime import datetime

    mgr = _get_contact_manager()
    stale = mgr.get_stale_contacts()

    if not stale:
        console.print("[green]No stale contacts. All follow-ups are current.[/green]")
        return

    table = Table(title="Stale Contacts (14+ days)")
    table.add_column("ID", style="dim", width=20)
    table.add_column("Name", style="bold red")
    table.add_column("Company")
    table.add_column("Type")
    table.add_column("Last Contact", style="yellow")
    table.add_column("Days Ago", style="red")
    table.add_column("Status")

    now = datetime.now()
    for r in stale:
        lcd = r.get("last_contact_date") or ""
        try:
            dt = datetime.fromisoformat(str(lcd).replace("Z", "+00:00"))
            dt = dt.replace(tzinfo=None) if dt.tzinfo else dt
            days_str = str((now - dt).days)
        except (ValueError, TypeError):
            days_str = "?"
        table.add_row(
            str(r["id"]),
            r["name"],
            r.get("company", "") or "",
            r.get("contact_type", "") or "",
            str(lcd)[:10],
            days_str,
            r.get("relationship_status", ""),
        )

    console.print(table)


@contacts.command("followups")
def contacts_followups():
    """Show contacts with follow-ups due today or overdue."""
    mgr = _get_contact_manager()
    due = mgr.get_followup_due()

    if not due:
        console.print("[green]No follow-ups due.[/green]")
        return

    table = Table(title="Follow-ups Due")
    table.add_column("ID", style="dim", width=20)
    table.add_column("Name", style="bold")
    table.add_column("Company")
    table.add_column("Due Date", style="yellow")
    table.add_column("Status")

    for r in due:
        table.add_row(
            str(r["id"]),
            r["name"],
            r.get("company", "") or "",
            r.get("next_followup", "") or "",
            r.get("relationship_status", ""),
        )

    console.print(table)


@contacts.command("tag")
@click.argument("contact_id", type=str)
@click.argument("tag")
def contacts_tag(contact_id, tag):
    """Add a tag to a contact."""
    mgr = _get_contact_manager()
    if not mgr.add_tag(contact_id, tag):
        console.print(f"[red]Contact {contact_id} not found.[/red]")
        return
    console.print(f"  [green]Tag '{tag}' added to contact {contact_id}.[/green]")


@contacts.command("untag")
@click.argument("contact_id", type=str)
@click.argument("tag")
def contacts_untag(contact_id, tag):
    """Remove a tag from a contact."""
    mgr = _get_contact_manager()
    if not mgr.remove_tag(contact_id, tag):
        console.print(f"[red]Contact {contact_id} not found.[/red]")
        return
    console.print(f"  [green]Tag '{tag}' removed from contact {contact_id}.[/green]")


@contacts.command("by-type")
@click.argument("contact_type")
def contacts_by_type(contact_type):
    """Filter contacts by type (recruiter, hiring_manager, etc.)."""
    mgr = _get_contact_manager()
    results = mgr.list_contacts(contact_type=contact_type)

    if not results:
        console.print(f"[dim]No contacts of type '{contact_type}'.[/dim]")
        return

    console.print(_contacts_table(results, title=f"Contacts: {contact_type}"))


@contacts.command("create-from-email")
@click.argument("email")
@click.option("--name", default=None, help="Contact name (prompted if omitted).")
def contacts_create_from_email(email, name):
    """Create a contact from an email address (quick capture)."""
    email = (email or "").strip()
    if not email:
        console.print("[red]Email is required.[/red]")
        raise click.Abort()

    mgr = _get_contact_manager()
    existing = mgr.find_by_email(email)
    if existing:
        company_str = f" ({existing['company']})" if existing.get("company") else ""
        console.print(
            f"[yellow]Contact already exists:[/yellow] "
            f"#{existing['id']} {existing['name']}{company_str} "
            f"[{existing.get('contact_type', 'recruiter')}]"
        )
        return

    if not name:
        name = click.prompt("  Name").strip()
    if not name:
        console.print("[red]Name is required to create a new contact.[/red]")
        raise click.Abort()

    cid = mgr.add_contact(name, contact_type="recruiter", email=email, source="email_import")
    console.print(
        f"[green]Added contact {cid}: {name} <{email}> "
        f"[recruiter, source=email_import][/green]"
    )


# --- Backward-compat: 'recruiters' alias ---


@cli.group(invoke_without_command=True)
@click.pass_context
def recruiters(ctx):
    """Recruiter contacts (alias for 'contacts' filtered by type=recruiter)."""
    if ctx.invoked_subcommand is not None:
        return

    mgr = _get_contact_manager()
    result = mgr.list_contacts(contact_type="recruiter")

    if not result:
        console.print("[dim]No recruiters tracked yet. Use 'contacts add' to add one.[/dim]")
        return

    console.print(_contacts_table(result, title="Recruiter Contacts"))


@cli.command()
def daily():
    """End-of-day AI summary of today's career activity."""
    from datetime import datetime

    from src.db import models
    from src.journal.entries import JournalManager

    today = datetime.now().strftime("%Y-%m-%d")

    # Gather today's activity
    activity_parts = []

    # Journal entries today
    manager = JournalManager()
    entries = manager.list_entries(days_back=1)
    today_entries = [e for e in entries if e["date"] == today]
    if today_entries:
        activity_parts.append(
            f"Journal entries today: {len(today_entries)} "
            f"({', '.join(e['type'] for e in today_entries)})"
        )

    # Application changes today
    try:
        from src.jobs.tracker import ApplicationTracker

        t = ApplicationTracker()
        all_jobs = t.get_all_jobs()
        today_jobs = [
            j for j in all_jobs
            if (j.get("date_found", "") or "").startswith(today)
            or (j.get("date_applied", "") or "").startswith(today)
            or (j.get("date_response", "") or "").startswith(today)
        ]
        t.close()
        if today_jobs:
            activity_parts.append(f"Application activity: {len(today_jobs)} job(s) updated")
    except Exception:
        pass

    # Skill updates today
    try:
        conn = models.get_connection()
        logs = models.get_skill_log(conn)
        conn.close()
        today_logs = [l for l in logs if (l.get("changed_at", "") or "").startswith(today)]
        if today_logs:
            activity_parts.append(
                f"Skills updated: {', '.join(l['skill_name'] for l in today_logs)}"
            )
    except Exception:
        pass

    # Interview analyses today
    try:
        from src.interviews.coach import InterviewCoach

        coach = InterviewCoach()
        analyses = coach.get_all_analyses()
        coach.close()
        today_analyses = [
            a for a in analyses
            if (a.get("analyzed_at", "") or "").startswith(today)
        ]
        if today_analyses:
            activity_parts.append(f"Interview analyses: {len(today_analyses)}")
    except Exception:
        pass

    if not activity_parts:
        activity_text = "No career-related activity recorded today."
    else:
        activity_text = "\n".join(f"- {p}" for p in activity_parts)

    # Send to router
    prompt = (
        f"Today's date: {today}\n\n"
        f"Today's career activity:\n{activity_text}\n\n"
        "Based on today's career-related activity, give me a brief end-of-day summary "
        "(3-5 sentences) and suggest 3 specific priorities for tomorrow. "
        "Be direct and practical \u2014 no motivational fluff. "
        "If there was no activity today, say so honestly and suggest what to focus on."
    )

    try:
        from src.llm.router import router
        summary = router.complete(task="daily_summary", prompt=prompt)
    except Exception:
        summary = (
            "Could not generate AI summary (API unavailable).\n\n"
            f"Raw activity:\n{activity_text}"
        )

    console.print()
    console.print(Panel(summary, title="Daily Summary", border_style="cyan"))


@cli.command()
@click.option("--type", "entry_type",
              type=click.Choice(["daily", "interview", "study", "project", "reflection"]),
              default="daily", help="Entry type (default: daily).")
def quick(entry_type):
    """Rapid journal entry \u2014 type and save, no menus."""
    from src.journal.entries import JournalManager

    console.print(f"[dim]Quick {entry_type} entry (press Enter twice to finish):[/dim]")
    lines = []
    while True:
        line = console.input("")
        if line == "" and lines and lines[-1] == "":
            lines.pop()
            break
        lines.append(line)

    content = "\n".join(lines).strip()
    if not content:
        console.print("[yellow]Empty entry, cancelled.[/yellow]")
        return

    manager = JournalManager()
    filename = manager.create_entry(entry_type, content)
    console.print(f"[green]Saved: {filename}[/green]")


@cli.command()
def status():
    """One-shot overview \u2014 today's activity at a glance."""
    from datetime import datetime

    from src.db import models

    today = datetime.now().strftime("%Y-%m-%d")
    parts = []

    # Journal entries today
    try:
        from src.journal.entries import JournalManager

        manager = JournalManager()
        entries = manager.list_entries(days_back=1)
        today_count = sum(1 for e in entries if e["date"] == today)
        parts.append(f"Journal entries today: [bold]{today_count}[/bold]")
    except Exception:
        parts.append("Journal: [dim]unavailable[/dim]")

    # Application pipeline
    try:
        from src.jobs.tracker import ApplicationTracker

        t = ApplicationTracker()
        stats = t.get_stats()
        t.close()

        applied = stats["by_status"].get("applied", 0)
        interviewing = stats["by_status"].get("interview", 0) + stats["by_status"].get("phone_screen", 0)
        offers = stats["by_status"].get("offer", 0)
        parts.append(
            f"Applications: [yellow]{applied}[/yellow] applied, "
            f"[green]{interviewing}[/green] interviewing, "
            f"[bright_green]{offers}[/bright_green] offers"
        )
    except Exception:
        parts.append("Applications: [dim]unavailable[/dim]")

    # Skill gaps
    try:
        conn = models.get_connection()
        gaps = models.get_gaps(conn)
        conn.close()
        parts.append(f"Skill gaps remaining: [bold]{len(gaps)}[/bold]")
    except Exception:
        parts.append("Skills: [dim]unavailable[/dim]")

    # Next calendar event
    try:
        from src.calendar.scheduler import CalendarScheduler

        sched = CalendarScheduler()
        sched.authenticate()
        events = sched.get_events(days_ahead=3)
        if events:
            next_evt = events[0]
            parts.append(f"Next event: [cyan]{next_evt['title']}[/cyan] at {next_evt['start']}")
        else:
            parts.append("Next event: [dim]none in next 3 days[/dim]")
    except Exception:
        parts.append("Calendar: [dim]not connected[/dim]")

    console.print()
    console.print(Panel("\n".join(parts), title="CareerPilot Status", border_style="cyan"))


# ─── Profile Commands ───────────────────────────────────────────────


@cli.group(invoke_without_command=True)
@click.pass_context
def profile(ctx):
    """Manage candidate profile for auto-fill and exports."""
    if ctx.invoked_subcommand is not None:
        return
    ctx.invoke(profile_show)


@profile.command("show")
def profile_show():
    """Display full profile in Rich panels, organized by section."""
    from src.profile.manager import ProfileManager

    mgr = ProfileManager()
    p = mgr.get_profile()
    mgr.close()

    personal = p.get("personal", {})
    if not personal:
        console.print("[yellow]No profile data. Run 'profile setup' or 'profile import' first.[/yellow]")
        return

    # Personal info panel
    personal_lines = []
    if personal.get("full_name"):
        personal_lines.append(f"[bold]{personal['full_name']}[/bold]")
    for field, label in [("email", "Email"), ("phone", "Phone")]:
        if personal.get(field):
            personal_lines.append(f"{label}: {personal[field]}")
    addr_parts = [personal.get(f, "") for f in ("street", "city", "state", "zip")]
    addr = ", ".join(part for part in addr_parts if part)
    if addr:
        personal_lines.append(f"Address: {addr}")
    for field, label in [("linkedin_url", "LinkedIn"), ("github_url", "GitHub"),
                         ("website", "Website")]:
        if personal.get(field):
            personal_lines.append(f"{label}: {personal[field]}")
    if personal.get("work_authorization"):
        personal_lines.append(f"Work Authorization: {personal['work_authorization'].replace('_', ' ').title()}")
    if personal.get("remote_preference"):
        personal_lines.append(f"Remote Preference: {personal['remote_preference'].replace('_', ' ').title()}")
    if personal.get("desired_salary_min") or personal.get("desired_salary_max"):
        sal = f"${personal.get('desired_salary_min', '?')} - ${personal.get('desired_salary_max', '?')}"
        personal_lines.append(f"Desired Salary: {sal}")
    if personal.get("available_start_date"):
        personal_lines.append(f"Available: {personal['available_start_date']}")
    console.print(Panel("\n".join(personal_lines), title="Personal Information", border_style="cyan"))

    # Work history
    work = p.get("work_history", [])
    if work:
        work_table = Table(title="Work History", border_style="green")
        work_table.add_column("ID", style="dim", width=4)
        work_table.add_column("Title", style="bold")
        work_table.add_column("Company")
        work_table.add_column("Location")
        work_table.add_column("Dates")
        work_table.add_column("Current", justify="center", width=7)
        for w in work:
            end = w.get("end_date") or "Present"
            current = "[green]Yes[/green]" if w.get("is_current") else ""
            work_table.add_row(
                str(w["id"]), w["title"], w["company"],
                w.get("location", ""),
                f"{w.get('start_date', '')} - {end}",
                current,
            )
        console.print(work_table)

    # Education
    edu = p.get("education", [])
    if edu:
        edu_table = Table(title="Education", border_style="blue")
        edu_table.add_column("ID", style="dim", width=4)
        edu_table.add_column("School", style="bold")
        edu_table.add_column("Degree")
        edu_table.add_column("Field")
        edu_table.add_column("Date")
        edu_table.add_column("GPA", width=5)
        for e in edu:
            edu_table.add_row(
                str(e["id"]), e["school"], e.get("degree", ""),
                e.get("field_of_study", ""), e.get("graduation_date", ""),
                e.get("gpa") or "",
            )
        console.print(edu_table)

    # Certifications
    certs = p.get("certifications", [])
    if certs:
        cert_table = Table(title="Certifications", border_style="yellow")
        cert_table.add_column("ID", style="dim", width=4)
        cert_table.add_column("Name", style="bold")
        cert_table.add_column("Issuer")
        cert_table.add_column("Date Obtained")
        cert_table.add_column("Status")
        for c in certs:
            status = "[yellow]In Progress[/yellow]" if c.get("in_progress") else "[green]Complete[/green]"
            cert_table.add_row(
                str(c["id"]), c["name"], c.get("issuer", ""),
                c.get("date_obtained", ""), status,
            )
        console.print(cert_table)

    # References
    refs = p.get("references", [])
    if refs:
        ref_table = Table(title="References", border_style="magenta")
        ref_table.add_column("ID", style="dim", width=4)
        ref_table.add_column("Name", style="bold")
        ref_table.add_column("Title")
        ref_table.add_column("Company")
        ref_table.add_column("Phone")
        ref_table.add_column("Email")
        ref_table.add_column("Relationship")
        for r in refs:
            ref_table.add_row(
                str(r["id"]), r["name"], r.get("title", ""),
                r.get("company", ""), r.get("phone", ""),
                r.get("email", ""), r.get("relationship", ""),
            )
        console.print(ref_table)

    # EEO (if populated)
    eeo = p.get("eeo", {})
    if any(eeo.get(f) for f in ("gender", "race_ethnicity", "veteran_status", "disability_status")):
        eeo_lines = []
        for field, label in [("gender", "Gender"), ("race_ethnicity", "Race/Ethnicity"),
                             ("veteran_status", "Veteran Status"),
                             ("disability_status", "Disability Status")]:
            if eeo.get(field):
                eeo_lines.append(f"{label}: {eeo[field]}")
        console.print(Panel("\n".join(eeo_lines), title="EEO (Private)", border_style="red"))


@profile.command("setup")
def profile_setup():
    """Interactive wizard to walk through all profile sections."""
    from rich.prompt import Confirm, Prompt

    from src.profile.manager import ProfileManager

    mgr = ProfileManager()
    console.print(Panel("[bold]Profile Setup Wizard[/bold]\nPress Enter to skip any field.",
                        border_style="cyan"))

    # --- Personal Info ---
    console.print("\n[bold cyan]== Personal Information ==[/bold cyan]")
    current = mgr.get_personal() or {}

    personal_fields = [
        ("full_name", "Full Name"),
        ("email", "Email"),
        ("phone", "Phone"),
        ("street", "Street Address"),
        ("city", "City"),
        ("state", "State"),
        ("zip", "ZIP Code"),
        ("linkedin_url", "LinkedIn URL"),
        ("github_url", "GitHub URL"),
        ("website", "Website"),
    ]
    updates = {}
    for field, label in personal_fields:
        cur = current.get(field, "")
        display = f" [dim]({cur})[/dim]" if cur else ""
        val = Prompt.ask(f"  {label}{display}", default=cur or "")
        if val:
            updates[field] = val

    # Work authorization
    cur_auth = current.get("work_authorization", "")
    console.print(f"\n  Work Authorization [dim]({cur_auth or 'not set'})[/dim]")
    console.print("    1) US Citizen  2) Permanent Resident  3) Require Sponsorship")
    auth_map = {"1": "us_citizen", "2": "permanent_resident", "3": "require_sponsorship"}
    auth_choice = Prompt.ask("  Choice", default="")
    if auth_choice in auth_map:
        updates["work_authorization"] = auth_map[auth_choice]

    # Remote preference
    cur_remote = current.get("remote_preference", "")
    console.print(f"\n  Remote Preference [dim]({cur_remote or 'not set'})[/dim]")
    console.print("    1) Remote Only  2) Hybrid  3) Onsite  4) Flexible")
    remote_map = {"1": "remote_only", "2": "hybrid", "3": "onsite", "4": "flexible"}
    remote_choice = Prompt.ask("  Choice", default="")
    if remote_choice in remote_map:
        updates["remote_preference"] = remote_map[remote_choice]

    # Salary
    cur_min = current.get("desired_salary_min") or ""
    cur_max = current.get("desired_salary_max") or ""
    sal_min = Prompt.ask(f"  Desired Salary Min [dim]({cur_min})[/dim]", default=str(cur_min) if cur_min else "")
    sal_max = Prompt.ask(f"  Desired Salary Max [dim]({cur_max})[/dim]", default=str(cur_max) if cur_max else "")
    if sal_min:
        updates["desired_salary_min"] = int(sal_min)
    if sal_max:
        updates["desired_salary_max"] = int(sal_max)

    # Relocate
    cur_relocate = current.get("willing_to_relocate", False)
    updates["willing_to_relocate"] = Confirm.ask(
        f"  Willing to relocate?", default=bool(cur_relocate))

    # Available start
    cur_avail = current.get("available_start_date", "")
    avail = Prompt.ask(f"  Available Start Date [dim]({cur_avail or 'not set'})[/dim]",
                       default=cur_avail or "")
    if avail:
        updates["available_start_date"] = avail

    if updates:
        mgr.update_personal(**updates)
        console.print("[green]Personal info saved.[/green]")

    # --- Work History ---
    console.print("\n[bold cyan]== Work History ==[/bold cyan]")
    existing_work = mgr.get_all_work_history()
    if existing_work:
        console.print(f"  {len(existing_work)} entries on file.")
    if Confirm.ask("  Add a new work history entry?", default=False):
        company = Prompt.ask("  Company")
        title = Prompt.ask("  Title")
        location = Prompt.ask("  Location", default="")
        start = Prompt.ask("  Start Date (YYYY-MM)", default="")
        end = Prompt.ask("  End Date (YYYY-MM, blank=current)", default="")
        desc = Prompt.ask("  Description", default="")
        is_current = end == ""
        mgr.add_work_history(company, title, location=location,
                             start_date=start, end_date=end or None,
                             description=desc, is_current=is_current)
        console.print("[green]Work history entry added.[/green]")

    # --- Education ---
    console.print("\n[bold cyan]== Education ==[/bold cyan]")
    existing_edu = mgr.get_all_education()
    if existing_edu:
        console.print(f"  {len(existing_edu)} entries on file.")
    if Confirm.ask("  Add a new education entry?", default=False):
        school = Prompt.ask("  School")
        degree = Prompt.ask("  Degree", default="")
        field = Prompt.ask("  Field of Study", default="")
        grad = Prompt.ask("  Graduation Date", default="")
        gpa = Prompt.ask("  GPA (optional)", default="")
        mgr.add_education(school, degree=degree, field_of_study=field,
                          graduation_date=grad, gpa=gpa or None)
        console.print("[green]Education entry added.[/green]")

    # --- Certifications ---
    console.print("\n[bold cyan]== Certifications ==[/bold cyan]")
    existing_certs = mgr.get_all_certifications()
    if existing_certs:
        console.print(f"  {len(existing_certs)} entries on file.")
    if Confirm.ask("  Add a new certification?", default=False):
        cert_name = Prompt.ask("  Certification Name")
        issuer = Prompt.ask("  Issuer", default="")
        obtained = Prompt.ask("  Date Obtained", default="")
        expiry = Prompt.ask("  Expiry Date (optional)", default="")
        in_prog = Confirm.ask("  In progress?", default=False)
        mgr.add_certification(cert_name, issuer=issuer, date_obtained=obtained,
                              expiry_date=expiry or None, in_progress=in_prog)
        console.print("[green]Certification added.[/green]")

    # --- References ---
    console.print("\n[bold cyan]== References ==[/bold cyan]")
    existing_refs = mgr.get_all_references()
    if existing_refs:
        console.print(f"  {len(existing_refs)} entries on file.")
    if Confirm.ask("  Add a new reference?", default=False):
        ref_name = Prompt.ask("  Name")
        ref_title = Prompt.ask("  Title", default="")
        ref_company = Prompt.ask("  Company", default="")
        ref_phone = Prompt.ask("  Phone", default="")
        ref_email = Prompt.ask("  Email", default="")
        ref_rel = Prompt.ask("  Relationship", default="")
        mgr.add_reference(ref_name, title=ref_title, company=ref_company,
                          phone=ref_phone, email=ref_email, relationship=ref_rel)
        console.print("[green]Reference added.[/green]")

    # --- EEO ---
    console.print("\n[bold cyan]== EEO (Optional, Private) ==[/bold cyan]")
    if Confirm.ask("  Set EEO fields?", default=False):
        eeo_updates = {}
        gender = Prompt.ask("  Gender", default="")
        if gender:
            eeo_updates["gender"] = gender
        race = Prompt.ask("  Race/Ethnicity", default="")
        if race:
            eeo_updates["race_ethnicity"] = race
        veteran = Prompt.ask("  Veteran Status", default="")
        if veteran:
            eeo_updates["veteran_status"] = veteran
        disability = Prompt.ask("  Disability Status", default="")
        if disability:
            eeo_updates["disability_status"] = disability
        if eeo_updates:
            mgr.update_eeo(**eeo_updates)
            console.print("[green]EEO data saved.[/green]")

    mgr.close()
    console.print(Panel("[bold green]Profile setup complete![/bold green]", border_style="green"))


@profile.command("edit")
@click.argument("section", type=click.Choice(
    ["personal", "work", "education", "certs", "references", "eeo", "preferences"]))
def profile_edit(section):
    """Edit a specific section of the profile."""
    from rich.prompt import Confirm, Prompt

    from src.profile.manager import ProfileManager

    mgr = ProfileManager()

    if section == "personal":
        current = mgr.get_personal() or {}
        updates = {}
        for field, label in [("full_name", "Full Name"), ("email", "Email"),
                             ("phone", "Phone"), ("street", "Street"),
                             ("city", "City"), ("state", "State"), ("zip", "ZIP"),
                             ("linkedin_url", "LinkedIn"), ("github_url", "GitHub"),
                             ("website", "Website")]:
            cur = current.get(field, "")
            val = Prompt.ask(f"  {label} [dim]({cur})[/dim]", default=cur or "")
            if val:
                updates[field] = val
        if updates:
            mgr.update_personal(**updates)
            console.print("[green]Personal info updated.[/green]")

    elif section == "work":
        entries = mgr.get_all_work_history()
        if entries:
            for w in entries:
                end = w.get("end_date") or "Present"
                console.print(f"  [dim]{w['id']}[/dim] {w['title']} at {w['company']} ({w.get('start_date', '')} - {end})")
        action = Prompt.ask("  [a]dd / [d]elete by ID / [s]kip", default="s")
        if action.lower() == "a":
            company = Prompt.ask("  Company")
            title = Prompt.ask("  Title")
            location = Prompt.ask("  Location", default="")
            start = Prompt.ask("  Start Date (YYYY-MM)", default="")
            end = Prompt.ask("  End Date (blank=current)", default="")
            desc = Prompt.ask("  Description", default="")
            mgr.add_work_history(company, title, location=location,
                                 start_date=start, end_date=end or None,
                                 description=desc, is_current=end == "")
            console.print("[green]Added.[/green]")
        elif action.lower() == "d":
            del_id = Prompt.ask("  ID to delete")
            if mgr.remove_work_history(int(del_id)):
                console.print("[green]Deleted.[/green]")
            else:
                console.print("[red]Not found.[/red]")

    elif section == "education":
        entries = mgr.get_all_education()
        if entries:
            for e in entries:
                console.print(f"  [dim]{e['id']}[/dim] {e.get('degree', '')} — {e['school']}")
        action = Prompt.ask("  [a]dd / [d]elete by ID / [s]kip", default="s")
        if action.lower() == "a":
            school = Prompt.ask("  School")
            degree = Prompt.ask("  Degree", default="")
            field = Prompt.ask("  Field of Study", default="")
            grad = Prompt.ask("  Graduation Date", default="")
            gpa = Prompt.ask("  GPA (optional)", default="")
            mgr.add_education(school, degree=degree, field_of_study=field,
                              graduation_date=grad, gpa=gpa or None)
            console.print("[green]Added.[/green]")
        elif action.lower() == "d":
            del_id = Prompt.ask("  ID to delete")
            if mgr.remove_education(int(del_id)):
                console.print("[green]Deleted.[/green]")
            else:
                console.print("[red]Not found.[/red]")

    elif section == "certs":
        entries = mgr.get_all_certifications()
        if entries:
            for c in entries:
                status = "(In Progress)" if c.get("in_progress") else ""
                console.print(f"  [dim]{c['id']}[/dim] {c['name']} {status}")
        action = Prompt.ask("  [a]dd / [d]elete by ID / [s]kip", default="s")
        if action.lower() == "a":
            cert_name = Prompt.ask("  Name")
            issuer = Prompt.ask("  Issuer", default="")
            obtained = Prompt.ask("  Date Obtained", default="")
            expiry = Prompt.ask("  Expiry Date (optional)", default="")
            in_prog = Confirm.ask("  In progress?", default=False)
            mgr.add_certification(cert_name, issuer=issuer, date_obtained=obtained,
                                  expiry_date=expiry or None, in_progress=in_prog)
            console.print("[green]Added.[/green]")
        elif action.lower() == "d":
            del_id = Prompt.ask("  ID to delete")
            if mgr.remove_certification(int(del_id)):
                console.print("[green]Deleted.[/green]")
            else:
                console.print("[red]Not found.[/red]")

    elif section == "references":
        entries = mgr.get_all_references()
        if entries:
            for r in entries:
                console.print(f"  [dim]{r['id']}[/dim] {r['name']} — {r.get('company', '')}")
        action = Prompt.ask("  [a]dd / [d]elete by ID / [s]kip", default="s")
        if action.lower() == "a":
            ref_name = Prompt.ask("  Name")
            ref_title = Prompt.ask("  Title", default="")
            ref_company = Prompt.ask("  Company", default="")
            ref_phone = Prompt.ask("  Phone", default="")
            ref_email = Prompt.ask("  Email", default="")
            ref_rel = Prompt.ask("  Relationship", default="")
            mgr.add_reference(ref_name, title=ref_title, company=ref_company,
                              phone=ref_phone, email=ref_email, relationship=ref_rel)
            console.print("[green]Added.[/green]")
        elif action.lower() == "d":
            del_id = Prompt.ask("  ID to delete")
            if mgr.remove_reference(int(del_id)):
                console.print("[green]Deleted.[/green]")
            else:
                console.print("[red]Not found.[/red]")

    elif section == "eeo":
        current_eeo = mgr.get_eeo() or {}
        updates = {}
        for field, label in [("gender", "Gender"), ("race_ethnicity", "Race/Ethnicity"),
                             ("veteran_status", "Veteran Status"),
                             ("disability_status", "Disability Status")]:
            cur = current_eeo.get(field, "")
            val = Prompt.ask(f"  {label} [dim]({cur})[/dim]", default=cur or "")
            if val:
                updates[field] = val
        if updates:
            mgr.update_eeo(**updates)
            console.print("[green]EEO data updated.[/green]")

    elif section == "preferences":
        current = mgr.get_personal() or {}
        updates = {}
        # Work auth
        console.print(f"  Work Authorization: {current.get('work_authorization', 'not set')}")
        console.print("    1) US Citizen  2) Permanent Resident  3) Require Sponsorship")
        auth_map = {"1": "us_citizen", "2": "permanent_resident", "3": "require_sponsorship"}
        choice = Prompt.ask("  Choice (blank=skip)", default="")
        if choice in auth_map:
            updates["work_authorization"] = auth_map[choice]
        # Remote
        console.print(f"  Remote Preference: {current.get('remote_preference', 'not set')}")
        console.print("    1) Remote Only  2) Hybrid  3) Onsite  4) Flexible")
        remote_map = {"1": "remote_only", "2": "hybrid", "3": "onsite", "4": "flexible"}
        choice = Prompt.ask("  Choice (blank=skip)", default="")
        if choice in remote_map:
            updates["remote_preference"] = remote_map[choice]
        # Salary
        sal_min = Prompt.ask(f"  Salary Min [dim]({current.get('desired_salary_min', '')})[/dim]", default="")
        sal_max = Prompt.ask(f"  Salary Max [dim]({current.get('desired_salary_max', '')})[/dim]", default="")
        if sal_min:
            updates["desired_salary_min"] = int(sal_min)
        if sal_max:
            updates["desired_salary_max"] = int(sal_max)
        # Relocate
        updates["willing_to_relocate"] = Confirm.ask("  Willing to relocate?",
                                                      default=bool(current.get("willing_to_relocate", False)))
        # Available
        avail = Prompt.ask(f"  Available [dim]({current.get('available_start_date', '')})[/dim]", default="")
        if avail:
            updates["available_start_date"] = avail
        if updates:
            mgr.update_personal(**updates)
            console.print("[green]Preferences updated.[/green]")

    mgr.close()


@profile.command("export")
@click.option("--format", "fmt", type=click.Choice(["json", "text", "ats"]),
              default="json", help="Export format.")
@click.option("--output", "-o", "outfile", default=None, help="Save to file instead of stdout.")
def profile_export(fmt, outfile):
    """Export profile to specified format."""
    from src.profile.manager import ProfileManager

    mgr = ProfileManager()

    if fmt == "json":
        result = mgr.export_json()
    elif fmt == "text":
        result = mgr.export_text()
    elif fmt == "ats":
        import json
        result = json.dumps(mgr.export_ats_fields(), indent=2)

    mgr.close()

    if outfile:
        from pathlib import Path
        Path(outfile).write_text(result, encoding="utf-8")
        console.print(f"[green]Exported to {outfile}[/green]")
    else:
        click.echo(result)


@profile.command("import")
def profile_import():
    """Import profile from resume via Claude API."""
    from src.profile.manager import ProfileManager

    console.print("[cyan]Importing profile from built-in resume data via Claude API...[/cyan]")

    mgr = ProfileManager()
    try:
        data = mgr.import_from_resume()
        console.print("[green]Profile imported successfully![/green]")
        sections = {k: v for k, v in data.items() if v}
        for section, content in sections.items():
            if isinstance(content, list):
                console.print(f"  {section}: {len(content)} entries")
            elif isinstance(content, dict):
                filled = sum(1 for v in content.values() if v)
                console.print(f"  {section}: {filled} fields populated")
    except Exception as e:
        console.print(f"[red]Import failed: {e}[/red]")
    finally:
        mgr.close()


@profile.command("seed")
def profile_seed():
    """Pre-populate profile with Joseph's data (no API call)."""
    from src.profile.manager import ProfileManager

    mgr = ProfileManager()
    mgr.seed_joseph_data()
    mgr.close()
    console.print("[green]Profile seeded with Joseph Fowler's data.[/green]")


# ─── Document Generation Commands ──────────────────────────────────


@cli.group()
def docs():
    """Generate tailored resumes and cover letters."""


@docs.command("resume")
@click.argument("job_id", type=int)
def docs_resume(job_id):
    """Generate a tailored resume for a tracked job."""
    from src.documents.resume_generator import ResumeGenerator
    from src.jobs.tracker import ApplicationTracker

    tracker = ApplicationTracker()
    job = tracker.get_job(job_id)
    tracker.close()

    if not job:
        console.print(f"[red]Job id={job_id} not found in tracker.[/red]")
        return

    console.print(Panel(
        f"[bold]{job.get('title', '?')}[/bold] at {job.get('company', '?')}\n"
        f"Location: {job.get('location', '?')}",
        title="Generating Resume", border_style="cyan",
    ))

    gen = ResumeGenerator()
    job_data = {
        "description": job.get("notes", "") or job.get("title", ""),
        "company": job.get("company", "Unknown"),
        "title": job.get("title", "Unknown"),
    }

    with console.status("[cyan]Tailoring resume with Claude...[/cyan]"):
        path = gen.generate_for_application(job_data)

    if not path:
        console.print("[red]Failed to generate resume. Check logs.[/red]")
        return

    console.print(f"[green]Resume saved:[/green] {path}")

    # Show preview
    try:
        from docx import Document as DocxDocument
        doc = DocxDocument(path)
        preview_text = "\n".join(p.text for p in doc.paragraphs[:20] if p.text.strip())
        console.print(Panel(preview_text, title="Resume Preview", border_style="green"))
    except Exception:
        pass

    _docs_post_action(path, gen, job_data)


@docs.command("cover-letter")
@click.argument("job_id", type=int)
def docs_cover_letter(job_id):
    """Generate a tailored cover letter for a tracked job."""
    from src.documents.cover_letter_generator import CoverLetterGenerator
    from src.jobs.tracker import ApplicationTracker
    from src.profile.manager import ProfileManager

    tracker = ApplicationTracker()
    job = tracker.get_job(job_id)
    tracker.close()

    if not job:
        console.print(f"[red]Job id={job_id} not found in tracker.[/red]")
        return

    console.print(Panel(
        f"[bold]{job.get('title', '?')}[/bold] at {job.get('company', '?')}\n"
        f"Location: {job.get('location', '?')}",
        title="Generating Cover Letter", border_style="cyan",
    ))

    mgr = ProfileManager()
    profile = mgr.get_profile()
    mgr.close()

    gen = CoverLetterGenerator(profile=profile)
    job_data = {
        "description": job.get("notes", "") or job.get("title", ""),
        "company": job.get("company", "Unknown"),
        "title": job.get("title", "Unknown"),
    }

    with console.status("[cyan]Generating cover letter with Claude...[/cyan]"):
        path = gen.generate_for_application(job_data)

    if not path:
        console.print("[red]Failed to generate cover letter. Check logs.[/red]")
        return

    console.print(f"[green]Cover letter saved:[/green] {path}")

    try:
        from docx import Document as DocxDocument
        doc = DocxDocument(path)
        preview_text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        console.print(Panel(preview_text, title="Cover Letter Preview", border_style="green"))
    except Exception:
        pass

    _docs_post_action(path, gen, job_data)


@docs.command("both")
@click.argument("job_id", type=int)
def docs_both(job_id):
    """Generate both tailored resume and cover letter for a tracked job."""
    from src.jobs.applicant import JobApplicant
    from src.jobs.tracker import ApplicationTracker

    tracker = ApplicationTracker()
    job = tracker.get_job(job_id)
    tracker.close()

    if not job:
        console.print(f"[red]Job id={job_id} not found in tracker.[/red]")
        return

    console.print(Panel(
        f"[bold]{job.get('title', '?')}[/bold] at {job.get('company', '?')}\n"
        f"Location: {job.get('location', '?')}",
        title="Generating Application Documents", border_style="cyan",
    ))

    job_data = {
        "description": job.get("notes", "") or job.get("title", ""),
        "company": job.get("company", "Unknown"),
        "title": job.get("title", "Unknown"),
    }

    applicant = JobApplicant()
    with console.status("[cyan]Generating documents with Claude...[/cyan]"):
        result = applicant.generate_application_docs(job_data)
    applicant.close()

    if result.get("resume_path"):
        console.print(f"[green]Resume saved:[/green] {result['resume_path']}")
        try:
            from docx import Document as DocxDocument
            doc = DocxDocument(result["resume_path"])
            preview = "\n".join(p.text for p in doc.paragraphs[:15] if p.text.strip())
            console.print(Panel(preview, title="Resume Preview", border_style="green"))
        except Exception:
            pass
    else:
        console.print("[red]Resume generation failed.[/red]")

    if result.get("cover_letter_path"):
        console.print(f"[green]Cover letter saved:[/green] {result['cover_letter_path']}")
        try:
            from docx import Document as DocxDocument
            doc = DocxDocument(result["cover_letter_path"])
            preview = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
            console.print(Panel(preview, title="Cover Letter Preview", border_style="green"))
        except Exception:
            pass
    else:
        console.print("[red]Cover letter generation failed.[/red]")


@docs.command("list")
def docs_list():
    """Show all generated documents with associated job info."""
    from pathlib import Path as P

    from config import settings as s

    resume_dir = s.DATA_DIR / "resumes"
    cl_dir = s.DATA_DIR / "cover_letters"

    table = Table(title="Generated Documents")
    table.add_column("Type", style="cyan", width=12)
    table.add_column("Filename", style="bold")
    table.add_column("Size", justify="right")
    table.add_column("Created")

    found = False
    for doc_dir, doc_type in [(resume_dir, "Resume"), (cl_dir, "Cover Letter")]:
        if doc_dir.exists():
            for f in sorted(doc_dir.glob("*.docx"), key=lambda p: p.stat().st_mtime, reverse=True):
                found = True
                size = f.stat().st_size
                size_str = f"{size / 1024:.1f} KB" if size >= 1024 else f"{size} B"
                from datetime import datetime as dt
                created = dt.fromtimestamp(f.stat().st_mtime).strftime("%Y-%m-%d %H:%M")
                table.add_row(doc_type, f.name, size_str, created)

    if not found:
        console.print("[yellow]No generated documents found.[/yellow]")
        console.print("Generate with: [bold]python cli.py docs resume <job_id>[/bold]")
        return

    console.print(table)


def _docs_post_action(path, generator, job_data):
    """Post-generation actions: open, regenerate, or accept."""
    while True:
        choice = console.input("\n[o]pen file, [r]egenerate, [a]ccept: ").strip().lower()
        if choice == "o":
            import webbrowser as wb
            wb.open(path)
        elif choice == "r":
            console.print("[cyan]Regenerating...[/cyan]")
            with console.status("[cyan]Regenerating with Claude...[/cyan]"):
                new_path = generator.generate_for_application(job_data)
            if new_path:
                path = new_path
                console.print(f"[green]New document saved:[/green] {path}")
                try:
                    from docx import Document as DocxDocument
                    doc = DocxDocument(path)
                    preview = "\n".join(p.text for p in doc.paragraphs[:15] if p.text.strip())
                    console.print(Panel(preview, title="Preview", border_style="green"))
                except Exception:
                    pass
            else:
                console.print("[red]Regeneration failed.[/red]")
        elif choice == "a" or not choice:
            console.print("[green]Accepted.[/green]")
            break
        else:
            console.print("[dim]Invalid choice.[/dim]")


# ─── LinkedIn Commands ─────────────────────────────────────────────


@cli.group()
def linkedin():
    """LinkedIn job search integration — scan emails, search, manage alerts."""


@linkedin.command("scan")
@click.option("--days", default=14, help="Number of days to look back (default 14).")
def linkedin_scan(days):
    """Scan Gmail for LinkedIn job alert emails, parse job listings."""
    from src.jobs.linkedin_cli import cmd_scan

    cmd_scan(days=days)


@linkedin.command("search")
def linkedin_search():
    """Open LinkedIn job search URLs in browser."""
    from src.jobs.linkedin_cli import cmd_search

    cmd_search()


@linkedin.command("alerts")
def linkedin_alerts():
    """Show setup guide for LinkedIn job alerts."""
    from src.jobs.linkedin_cli import cmd_alerts

    cmd_alerts()


@linkedin.command("profiles")
def linkedin_profiles():
    """List configured LinkedIn search profiles."""
    from src.jobs.linkedin_cli import cmd_profiles

    cmd_profiles()


# ─── Gmail Filter Commands ─────────────────────────────────────────


@cli.group()
def filters():
    """Manage Gmail smart filters for job search emails."""


@filters.command("setup")
def filters_setup():
    """Create labels + filter rules + retroactively tag existing mail."""
    import copy

    from src.gmail.filter_config import FILTER_RULES, LABELS, USER_RECRUITER_DOMAINS_FILE, build_gmail_query
    from src.gmail.filters import GmailFilterManager

    console.print(Panel("Gmail Filter Setup", style="bold cyan"))

    # Load user-added recruiter domains
    user_domains = []
    if USER_RECRUITER_DOMAINS_FILE.exists():
        user_domains = [
            line.strip()
            for line in USER_RECRUITER_DOMAINS_FILE.read_text().splitlines()
            if line.strip()
        ]

    rules = copy.deepcopy(FILTER_RULES)
    if user_domains and "recruiters" in rules:
        existing = rules["recruiters"].get("from_domains", [])
        rules["recruiters"]["from_domains"] = list(set(existing + user_domains))

    mgr = GmailFilterManager()

    # Step 1: Create labels
    console.print("\n[bold]Creating label hierarchy...[/bold]")
    label_map = mgr.ensure_label_hierarchy(LABELS)
    for name, lid in label_map.items():
        console.print(f"  [green]OK[/green]  {name}")

    # Step 2: Remove old CareerPilot filters (idempotent re-run)
    console.print("\n[bold]Clearing old CareerPilot filters...[/bold]")
    old_filters = mgr.get_careerpilot_filters()
    for f in old_filters:
        mgr.delete_filter(f["id"])
    console.print(f"  Removed {len(old_filters)} old filter(s)")

    # Step 3: Create new filters
    console.print("\n[bold]Creating filter rules...[/bold]")
    for rule_name, rule in rules.items():
        query = build_gmail_query(rule)
        if not query:
            console.print(f"  [yellow]SKIP[/yellow]  {rule_name} (empty query)")
            continue
        label_name = rule["label"]
        label_id = label_map.get(label_name) or mgr.get_label_id(label_name)
        if not label_id:
            console.print(f"  [red]ERR[/red]   Label '{label_name}' not found")
            continue
        mgr.create_filter(query, label_id, archive=rule.get("archive", False))
        console.print(f"  [green]OK[/green]  {rule_name} -> {label_name}")

    # Step 4: Retroactively label existing messages
    console.print("\n[bold]Retroactively tagging existing messages...[/bold]")
    total_tagged = 0
    for rule_name, rule in rules.items():
        query = build_gmail_query(rule)
        if not query:
            continue
        label_name = rule["label"]
        label_id = label_map.get(label_name) or mgr.get_label_id(label_name)
        if not label_id:
            continue
        count = mgr.apply_label_to_matching(query, label_id)
        if count > 0:
            console.print(f"  {rule_name}: tagged {count} message(s)")
            total_tagged += count
    console.print(f"\n  Total: {total_tagged} messages retroactively tagged")
    console.print("\n[green]Setup complete! Check your Gmail sidebar for CareerPilot labels.[/green]")


@filters.command("list")
def filters_list():
    """Show current CareerPilot filter rules."""
    import copy

    from src.gmail.filter_config import FILTER_RULES, USER_RECRUITER_DOMAINS_FILE, build_gmail_query
    from src.gmail.filters import GmailFilterManager

    # Merge user domains
    user_domains = []
    if USER_RECRUITER_DOMAINS_FILE.exists():
        user_domains = [
            line.strip()
            for line in USER_RECRUITER_DOMAINS_FILE.read_text().splitlines()
            if line.strip()
        ]
    rules = copy.deepcopy(FILTER_RULES)
    if user_domains and "recruiters" in rules:
        existing = rules["recruiters"].get("from_domains", [])
        rules["recruiters"]["from_domains"] = list(set(existing + user_domains))

    table = Table(title="CareerPilot Filter Rules")
    table.add_column("Rule", style="cyan")
    table.add_column("Label", style="green")
    table.add_column("Query (truncated)")
    for rule_name, rule in rules.items():
        query = build_gmail_query(rule)
        table.add_row(rule_name, rule["label"], query[:80] + ("..." if len(query) > 80 else ""))
    console.print(table)

    if user_domains:
        console.print(f"\n[bold]User-added recruiter domains:[/bold] {', '.join(user_domains)}")

    try:
        mgr = GmailFilterManager()
        live = mgr.get_careerpilot_filters()
        console.print(f"\nLive Gmail filters targeting CareerPilot labels: {len(live)}")
    except Exception:
        console.print("\n[yellow]Could not connect to Gmail to check live filters[/yellow]")


@filters.command("add")
@click.argument("domain")
def filters_add(domain):
    """Add a recruiter domain to the filter list."""
    from src.gmail.filter_config import USER_RECRUITER_DOMAINS_FILE

    domain = domain.strip().lower().lstrip("@")

    domains = []
    if USER_RECRUITER_DOMAINS_FILE.exists():
        domains = [
            line.strip()
            for line in USER_RECRUITER_DOMAINS_FILE.read_text().splitlines()
            if line.strip()
        ]

    if domain in domains:
        console.print(f"[yellow]'{domain}' is already in the recruiter domain list.[/yellow]")
        return

    domains.append(domain)
    USER_RECRUITER_DOMAINS_FILE.parent.mkdir(parents=True, exist_ok=True)
    USER_RECRUITER_DOMAINS_FILE.write_text("\n".join(sorted(set(domains))) + "\n")
    console.print(f"[green]Added '{domain}' to recruiter filter domains.[/green]")
    console.print("Run [bold]python cli.py filters setup[/bold] to apply changes to Gmail.")


@filters.command("remove")
@click.argument("domain")
def filters_remove(domain):
    """Remove a recruiter domain from the user list."""
    from src.gmail.filter_config import USER_RECRUITER_DOMAINS_FILE

    domain = domain.strip().lower().lstrip("@")

    domains = []
    if USER_RECRUITER_DOMAINS_FILE.exists():
        domains = [
            line.strip()
            for line in USER_RECRUITER_DOMAINS_FILE.read_text().splitlines()
            if line.strip()
        ]

    if domain not in domains:
        console.print(f"[yellow]'{domain}' is not in the user-added domain list.[/yellow]")
        return

    domains.remove(domain)
    USER_RECRUITER_DOMAINS_FILE.write_text("\n".join(sorted(set(domains))) + "\n")
    console.print(f"[green]Removed '{domain}' from recruiter filter domains.[/green]")
    console.print("Run [bold]python cli.py filters setup[/bold] to apply changes to Gmail.")


@filters.command("test")
def filters_test():
    """Dry-run: show what queries would be created (no API calls)."""
    import copy

    from src.gmail.filter_config import FILTER_RULES, USER_RECRUITER_DOMAINS_FILE, build_gmail_query

    user_domains = []
    if USER_RECRUITER_DOMAINS_FILE.exists():
        user_domains = [
            line.strip()
            for line in USER_RECRUITER_DOMAINS_FILE.read_text().splitlines()
            if line.strip()
        ]
    rules = copy.deepcopy(FILTER_RULES)
    if user_domains and "recruiters" in rules:
        existing = rules["recruiters"].get("from_domains", [])
        rules["recruiters"]["from_domains"] = list(set(existing + user_domains))

    console.print(Panel("Dry Run -- Filter Queries (no API calls)", style="bold yellow"))
    for rule_name, rule in rules.items():
        query = build_gmail_query(rule)
        console.print(f"\n[cyan]{rule_name}[/cyan] -> {rule['label']}")
        console.print(f"  {query}")


@filters.command("nuke")
def filters_nuke():
    """Remove all CareerPilot filters from Gmail (labels preserved)."""
    from src.gmail.filters import GmailFilterManager

    if not click.confirm("This will remove all CareerPilot Gmail filters. Continue?"):
        console.print("[dim]Cancelled.[/dim]")
        return

    mgr = GmailFilterManager()
    cp_filters = mgr.get_careerpilot_filters()
    for f in cp_filters:
        mgr.delete_filter(f["id"])
    console.print(f"[green]Removed {len(cp_filters)} filter(s). Labels are preserved.[/green]")


@cli.command()
def inbox():
    """Threaded email dashboard with one-click actions."""
    from src.gmail.auth import get_gmail_service
    from src.gmail.dashboard import EmailDashboard
    from src.gmail.thread_actions import ThreadActions

    # Authenticate
    try:
        service = get_gmail_service()
    except FileNotFoundError as e:
        console.print(f"[red]{e}[/red]")
        return
    except Exception:
        console.print("[red]Gmail authentication failed. Check logs for details.[/red]")
        return

    dashboard_obj = EmailDashboard(service)
    responder = None
    cal_scheduler = None

    # Set up responder
    try:
        from src.gmail.responder import RecruiterResponder
        responder = RecruiterResponder(service)
    except Exception:
        pass

    # Set up calendar (non-blocking)
    try:
        from src.calendar.scheduler import CalendarScheduler
        cal_scheduler = CalendarScheduler()
        cal_scheduler.authenticate()
    except Exception:
        cal_scheduler = None

    actions = ThreadActions(service, responder=responder, cal_scheduler=cal_scheduler)

    # Fetch threads
    console.print("[dim]Fetching inbox threads...[/dim]")
    threads = dashboard_obj.fetch_threads(max_results=50)

    if not threads:
        console.print("[yellow]No threads found in CareerPilot labels.[/yellow]")
        return

    # Classify all threads and check snooze status
    enriched = []
    for t in threads:
        snoozed, snooze_info = actions.is_snoozed(t["thread_id"])
        if snoozed:
            continue  # Hide snoozed threads

        status_info = dashboard_obj.classify_thread_status(t)
        t["status"] = status_info["status"]
        t["hours_since_last"] = status_info["hours_since_last"]
        t["is_stale"] = status_info["is_stale"]

        # Check if snooze just expired
        if snooze_info and not snoozed:
            t["snooze_expired"] = True
        else:
            t["snooze_expired"] = False

        # Check linked job
        linked_job = actions.get_linked_job(t["thread_id"])
        t["linked_job_id"] = linked_job

        enriched.append(t)

    # Digest header
    digest = {
        "awaiting_reply": sum(1 for t in enriched if t["status"] == "awaiting_reply"),
        "stale_count": sum(1 for t in enriched if t["is_stale"]),
        "interview_count": sum(1 for t in enriched if t["category"] == "Interviews"),
        "new_24h": sum(
            1 for t in enriched
            if t["last_message_date"].tzinfo
            and (
                __import__("datetime").datetime.now(__import__("datetime").timezone.utc)
                - t["last_message_date"]
            ).total_seconds() < 86400
        ),
    }

    digest_parts = []
    digest_parts.append(f"{digest['awaiting_reply']} awaiting reply")
    if digest["stale_count"]:
        digest_parts.append(f"[red]{digest['stale_count']} stale![/red]")
    if digest["new_24h"]:
        digest_parts.append(f"{digest['new_24h']} new today")
    if digest["interview_count"]:
        digest_parts.append(f"{digest['interview_count']} interview(s)")

    console.print(Panel(
        f"[bold]Inbox:[/bold] {', '.join(digest_parts)}",
        border_style="cyan",
    ))
    console.print()

    # Thread table
    table = Table(title=f"Email Threads ({len(enriched)})")
    table.add_column("#", style="dim", width=4)
    table.add_column("Status", width=12)
    table.add_column("Category", width=16)
    table.add_column("From/Company")
    table.add_column("Subject")
    table.add_column("Last Activity", width=14)
    table.add_column("Age", width=8)

    STATUS_COLORS = {
        "awaiting_reply": "yellow",
        "awaiting_response": "green",
        "scheduled": "cyan",
        "unknown": "dim",
    }

    for i, t in enumerate(enriched, 1):
        status = t["status"]
        color = STATUS_COLORS.get(status, "white")
        if t["is_stale"]:
            color = "red"

        status_display = status.replace("_", " ")
        if t["snooze_expired"]:
            status_display = "follow up!"
            color = "bright_magenta"

        # Participant display (first non-self sender)
        from_display = ""
        for p in t["participants"]:
            user_email = dashboard_obj._get_user_email()
            if user_email and user_email.lower() not in p.lower():
                from_display = p[:30]
                break
        if not from_display and t["participants"]:
            from_display = t["participants"][0][:30]

        # Age display
        hours = t["hours_since_last"]
        if hours < 24:
            age = f"{int(hours)}h"
        else:
            age = f"{int(hours / 24)}d"

        last_date = t["last_message_date"].strftime("%m/%d %H:%M")

        linked = ""
        if t["linked_job_id"]:
            linked = f" [dim](#{t['linked_job_id']})[/dim]"

        table.add_row(
            str(i),
            f"[{color}]{status_display}[/{color}]",
            t["category"],
            from_display,
            t["subject"][:40] + linked,
            last_date,
            age,
        )

    console.print(table)
    console.print()

    # Interactive thread selection
    while True:
        choice = console.input("Select thread # or [bold][q][/bold]uit: ").strip().lower()
        if choice == "q":
            return

        try:
            idx = int(choice) - 1
            if not 0 <= idx < len(enriched):
                console.print("[red]Invalid thread number.[/red]")
                continue
        except ValueError:
            console.print("[red]Enter a number or 'q'.[/red]")
            continue

        thread = enriched[idx]
        _inbox_thread_actions(thread, actions, dashboard_obj, console)


def _inbox_thread_actions(thread, actions, dashboard_obj, console):
    """Handle per-thread actions after selecting from inbox."""
    from rich.panel import Panel

    # Show preview (last 2 messages)
    messages = dashboard_obj.get_thread_messages(thread["thread_id"])
    if messages:
        preview_msgs = messages[-2:] if len(messages) > 1 else messages
        for msg in preview_msgs:
            if msg["is_from_me"]:
                border = "green"
                label = "You"
            else:
                border = "cyan"
                label = msg["sender"][:50]
            body_preview = (msg["body"] or "")[:500]
            console.print(Panel(
                body_preview,
                title=f"{label} ({msg['date']})",
                border_style=border,
            ))

    # Show linked application if tracked
    if thread.get("linked_job_id"):
        try:
            from src.jobs.tracker import ApplicationTracker
            t = ApplicationTracker()
            job = t.get_job(thread["linked_job_id"])
            t.close()
            if job:
                console.print(
                    f"  [dim]Linked:[/dim] {job.get('title', '?')} at "
                    f"{job.get('company', '?')} ({job.get('status', '?')})"
                )
        except Exception:
            pass

    console.print()
    console.print(
        "[bold][r][/bold]eply  [bold][b][/bold]ook  [bold][s][/bold]nooze  "
        "[bold][a][/bold]rchive  [bold][t][/bold]rack  [bold][v][/bold]iew full  "
        "[bold][q][/bold]back"
    )

    while True:
        action = console.input("Action: ").strip().lower()

        if action == "q":
            return

        elif action == "r":
            _inbox_reply_flow(thread, actions, console)

        elif action == "b":
            _inbox_book_flow(thread, actions, console)

        elif action == "s":
            days_input = console.input("Snooze for how many days? [3]: ").strip()
            days = int(days_input) if days_input.isdigit() else 3
            if actions.snooze(thread["thread_id"], days=days, subject=thread["subject"]):
                console.print(f"[green]Snoozed for {days} days.[/green]")
            else:
                console.print("[red]Failed to snooze.[/red]")
            return

        elif action == "a":
            confirm = console.input("Archive this thread? [y/n]: ").strip().lower()
            if confirm == "y":
                if actions.archive(thread["thread_id"]):
                    console.print("[green]Thread archived.[/green]")
                else:
                    console.print("[red]Failed to archive.[/red]")
            return

        elif action == "t":
            _inbox_track_flow(thread, actions, console)

        elif action == "v":
            actions.view(thread["thread_id"], console)

        else:
            console.print("[red]Invalid action. Enter r, b, s, a, t, v, or q.[/red]")


def _inbox_reply_flow(thread, actions, console):
    """Reply flow: choose mode, generate draft, review."""
    from rich.panel import Panel

    console.print(
        "  Mode: [bold][i][/bold]nterested  [bold][d][/bold]ecline  "
        "[bold][m][/bold]ore info"
    )
    mode_choice = console.input("  Mode: ").strip().lower()
    mode_map = {"i": "interested", "d": "not_interested", "m": "more_info"}
    mode = mode_map.get(mode_choice)
    if not mode:
        console.print("[red]Invalid mode.[/red]")
        return

    console.print("[dim]Generating reply...[/dim]")
    draft_text = actions.reply(thread["thread_id"], mode=mode)

    if not draft_text:
        console.print("[red]Failed to generate reply. Check logs.[/red]")
        return

    while True:
        console.print(Panel(draft_text, title=f"Draft Reply ({mode})", border_style="cyan"))

        choice = console.input(
            "  [bold][a][/bold]pprove (save draft)  "
            "[bold][e][/bold]dit (re-generate)  "
            "[bold][c][/bold]ancel: "
        ).strip().lower()

        if choice == "a":
            draft_id = actions.save_reply_draft(thread["thread_id"], draft_text)
            if draft_id:
                console.print(f"[green]Draft saved to Gmail (draft_id={draft_id}). NOT sent.[/green]")
            else:
                console.print("[red]Failed to save draft. Check logs.[/red]")
            return

        elif choice == "e":
            console.print("[dim]Re-generating...[/dim]")
            draft_text = actions.reply(thread["thread_id"], mode=mode)
            if not draft_text:
                console.print("[red]Failed to re-generate. Check logs.[/red]")
                return

        elif choice == "c":
            console.print("[yellow]Cancelled.[/yellow]")
            return

        else:
            console.print("[red]Invalid choice. Enter a, e, or c.[/red]")


def _inbox_book_flow(thread, actions, console):
    """Booking flow: show availability, generate scheduling reply."""
    from rich.panel import Panel

    console.print("[dim]Generating scheduling response...[/dim]")
    draft_text, slots = actions.book(thread["thread_id"])

    if not draft_text:
        console.print("[red]Failed to generate booking reply. Check logs.[/red]")
        return

    console.print(Panel(draft_text, title="Scheduling Reply", border_style="cyan"))

    choice = console.input(
        "  [bold][a][/bold]pprove (save draft)  [bold][c][/bold]ancel: "
    ).strip().lower()

    if choice == "a":
        draft_id = actions.save_reply_draft(thread["thread_id"], draft_text)
        if draft_id:
            console.print(f"[green]Draft saved (draft_id={draft_id}). NOT sent.[/green]")
        else:
            console.print("[red]Failed to save draft.[/red]")

        # Offer calendar holds
        if slots and actions._cal_scheduler:
            hold = console.input("Create calendar holds for suggested times? [y/n]: ").strip().lower()
            if hold == "y":
                company = ""
                for p in thread.get("participants", []):
                    if "@" in p:
                        company = p.split("@")[-1].split(".")[0].title()
                        break
                title = f"Interview \u2014 {company}" if company else "Interview Hold"
                for slot in slots[:3]:
                    event_id = actions._cal_scheduler.create_hold(title, slot)
                    if event_id:
                        console.print(
                            f"  [green]Hold: {slot.strftime('%A %B %#d at %#I:%M %p %Z')}[/green]"
                        )
    elif choice == "c":
        console.print("[yellow]Cancelled.[/yellow]")


def _inbox_track_flow(thread, actions, console):
    """Link a thread to an application in the tracker."""
    try:
        from src.jobs.tracker import ApplicationTracker
        t = ApplicationTracker()
        pipeline = t.get_pipeline()
    except Exception:
        console.print("[red]Could not access application tracker.[/red]")
        return

    # Collect all jobs
    all_jobs = []
    for status, jobs in pipeline.items():
        all_jobs.extend(jobs)

    if not all_jobs:
        console.print("[yellow]No applications in tracker to link.[/yellow]")
        t.close()
        return

    table = Table(title="Applications")
    table.add_column("#", style="dim", width=4)
    table.add_column("Title", style="bold")
    table.add_column("Company")
    table.add_column("Status")

    for i, j in enumerate(all_jobs, 1):
        table.add_row(
            str(i),
            str(j.get("title", ""))[:35],
            str(j.get("company", ""))[:25],
            j.get("status", ""),
        )

    console.print(table)

    choice = console.input("Link to application #: ").strip()
    try:
        idx = int(choice) - 1
        if 0 <= idx < len(all_jobs):
            job_id = all_jobs[idx]["id"]
            if actions.track(thread["thread_id"], job_id):
                console.print(f"[green]Linked to: {all_jobs[idx].get('title', '?')}[/green]")
            else:
                console.print("[red]Failed to link.[/red]")
        else:
            console.print("[red]Invalid number.[/red]")
    except ValueError:
        console.print("[red]Invalid input.[/red]")

    t.close()


# ─── ATS Form Auto-Fill Commands ──────────────────────────────────


@cli.group()
def fill():
    """ATS form auto-fill — generate prompts, cheatsheets, and detect ATS type."""


@fill.command("url")
@click.argument("job_url")
@click.option("--ats", "ats_type", default=None, help="Force ATS type (workday, greenhouse, lever, icims).")
@click.option("--execute", is_flag=True, help="Launch Claude Code + Chrome to fill the form.")
@click.option("--clipboard", is_flag=True, help="Copy cheatsheet to clipboard instead of prompt.")
@click.option("--resume", "resume_path", default=None, help="Path to resume file.")
@click.option("--cover-letter", "cover_letter_path", default=None, help="Path to cover letter file.")
def fill_url(job_url, ats_type, execute, clipboard, resume_path, cover_letter_path):
    """Auto-detect ATS and generate a fill prompt for JOB_URL."""
    from src.browser.fill_cli import cmd_fill

    cmd_fill(
        job_url=job_url,
        ats_type=ats_type,
        execute=execute,
        clipboard=clipboard,
        resume_path=resume_path,
        cover_letter_path=cover_letter_path,
    )


@fill.command("detect")
@click.argument("url")
def fill_detect(url):
    """Identify which ATS system a URL belongs to."""
    from src.browser.fill_cli import cmd_detect

    cmd_detect(url)


@fill.command("list-ats")
def fill_list_ats():
    """Show all supported ATS systems."""
    from src.browser.fill_cli import cmd_list_ats

    cmd_list_ats()


@fill.command("cheatsheet")
def fill_cheatsheet():
    """Print quick-fill values from your profile."""
    from src.browser.fill_cli import cmd_cheatsheet

    cmd_cheatsheet()


# ─── IT Staffing Agency Commands ─────────────────────────────────


@cli.group()
def agencies():
    """IT staffing agency integration — manage recruiters, roles, and outreach."""


@agencies.command("list")
def agencies_list():
    """List all configured staffing agencies."""
    from src.agencies.agencies_cli import cmd_list_agencies

    cmd_list_agencies()


@agencies.command("search")
@click.argument("keyword", required=False)
@click.option("--all", "open_all", is_flag=True, help="Open all agencies with default keywords.")
def agencies_search(keyword, open_all):
    """Open agency job boards in browser for KEYWORD."""
    from src.agencies.agencies_cli import cmd_search

    cmd_search(keyword=keyword, open_all=open_all)


@agencies.group("recruiter")
def agencies_recruiter():
    """Manage recruiter contacts."""


@agencies_recruiter.command("add")
def agencies_recruiter_add():
    """Add a recruiter contact (interactive)."""
    from src.agencies.agencies_cli import cmd_recruiter_add

    cmd_recruiter_add()


@agencies_recruiter.command("list")
def agencies_recruiter_list():
    """List all recruiter contacts."""
    from src.agencies.agencies_cli import cmd_recruiter_list

    cmd_recruiter_list()


@agencies_recruiter.command("show")
@click.argument("recruiter_id", type=int)
def agencies_recruiter_show(recruiter_id):
    """Show recruiter details and history."""
    from src.agencies.agencies_cli import cmd_recruiter_show

    cmd_recruiter_show(recruiter_id)


@agencies.group("interaction")
def agencies_interaction():
    """Manage recruiter interactions."""


@agencies_interaction.command("log")
@click.argument("recruiter_id", type=int)
def agencies_interaction_log(recruiter_id):
    """Log an interaction with a recruiter."""
    from src.agencies.agencies_cli import cmd_interaction_log

    cmd_interaction_log(recruiter_id)


@agencies.group("role")
def agencies_role():
    """Manage submitted roles."""


@agencies_role.command("add")
@click.argument("recruiter_id", type=int)
def agencies_role_add(recruiter_id):
    """Add a submitted role for a recruiter."""
    from src.agencies.agencies_cli import cmd_role_add

    cmd_role_add(recruiter_id)


@agencies_role.command("list")
def agencies_role_list():
    """List all submitted roles."""
    from src.agencies.agencies_cli import cmd_role_list

    cmd_role_list()


@agencies_role.command("update")
@click.argument("role_id", type=int)
@click.argument("status")
def agencies_role_update(role_id, status):
    """Update a submitted role's status."""
    from src.agencies.agencies_cli import cmd_role_update

    cmd_role_update(role_id, status)


@agencies.command("outreach")
@click.argument("template", required=False)
@click.argument("agency", required=False)
@click.option("--list", "list_templates", is_flag=True, help="List available templates.")
def agencies_outreach(template, agency, list_templates):
    """Generate an outreach email from a template."""
    from src.agencies.agencies_cli import cmd_outreach

    if list_templates:
        cmd_outreach(template_key="--list")
    else:
        cmd_outreach(template_key=template, agency_key=agency)


@agencies.command("summary")
def agencies_summary():
    """Show recruiter relationship dashboard."""
    from src.agencies.agencies_cli import cmd_summary

    cmd_summary()


@agencies.command("seed")
def agencies_seed():
    """Seed tracker with David Perez / TEKsystems data (idempotent)."""
    from src.db import models

    conn = models.get_connection()

    existing = models.find_contact_by_email(conn, "dperez@teksystems.com")
    if existing:
        console.print(f"[yellow]Contact already exists (#{existing['id']}). Skipping seed.[/yellow]")
        conn.close()
        return

    rid = models.add_contact(
        conn, "David Perez", "recruiter",
        company="TEKsystems",
        email="dperez@teksystems.com",
        phone="317-810-7562",
        title="Sr. Information Technology Recruiter (Risk & Security)",
        notes="Active relationship. Indy office: 9265 Counselors Row.",
        source="staffing_agency",
    )
    console.print(f"[green]Added contact #{rid}: David Perez at TEKsystems[/green]")

    r1 = models.add_submitted_role(
        conn, rid, "MISO Energy", "Systems Administrator",
        location="Indianapolis, IN", role_type="contract",
    )
    r2 = models.add_submitted_role(
        conn, rid, "Corteva", "Domain Migration Support Specialist",
        location="Indianapolis, IN", role_type="contract",
    )
    r3 = models.add_submitted_role(
        conn, rid, "Delta", "Desktop Support Technician",
        location="Indianapolis, IN", role_type="contract",
    )
    console.print(f"[green]Added 3 submitted roles (#{r1}, #{r2}, #{r3})[/green]")

    iid = models.add_contact_interaction(
        conn, rid, "email", "inbound",
        subject="Miso - Systems Admin",
        summary="Presented to MISO Energy for sys admin role. Asked for consent to represent.",
        roles_discussed="MISO - Systems Admin, Corteva - Domain Migration, Delta - Desktop Support",
    )
    console.print(f"[green]Logged interaction #{iid}[/green]")

    conn.close()
    console.print("[bold green]Seed complete.[/bold green]")


# ── Intel ─────────────────────────────────────────────────────────────────────


def _display_brief(brief, company):
    """Render a company intel brief as Rich panels."""
    from rich.markdown import Markdown
    from rich.text import Text

    console.print()
    console.print(
        Panel(
            f"[bold white]{company}[/bold white]",
            title="Company Intelligence",
            border_style="bright_cyan",
            expand=True,
        )
    )

    # --- Overview ---
    overview = brief.get("company_overview", {})
    lines = []
    if overview.get("description"):
        lines.append(overview["description"])
    if overview.get("headquarters"):
        lines.append(f"HQ: {overview['headquarters']}")
    parts = []
    if overview.get("size"):
        parts.append(overview["size"])
    if overview.get("revenue_or_funding"):
        parts.append(overview["revenue_or_funding"])
    if parts:
        lines.append(" | ".join(parts))
    if overview.get("key_products"):
        lines.append(f"Key products: {', '.join(overview['key_products'])}")
    if overview.get("recent_news"):
        lines.append("")
        lines.append("[bold]Recent News:[/bold]")
        for item in overview["recent_news"][:5]:
            date = item.get("date", "")
            headline = item.get("headline", "")
            lines.append(f"  [{date}] {headline}")
    console.print(Panel("\n".join(lines), title="OVERVIEW", border_style="blue"))

    # --- Culture ---
    culture = brief.get("culture", {})
    lines = []
    if culture.get("glassdoor_rating"):
        lines.append(f"Glassdoor: {culture['glassdoor_rating']}")
    if culture.get("sentiment_summary"):
        lines.append(culture["sentiment_summary"])
    if culture.get("remote_policy"):
        lines.append(f"Remote policy: {culture['remote_policy']}")
    if culture.get("work_life_balance"):
        lines.append(f"Work-life balance: {culture['work_life_balance']}")
    if culture.get("pros"):
        lines.append(f"[green]Pros:[/green] {', '.join(culture['pros'])}")
    if culture.get("cons"):
        lines.append(f"[red]Cons:[/red] {', '.join(culture['cons'])}")
    console.print(Panel("\n".join(lines), title="CULTURE & ENVIRONMENT", border_style="yellow"))

    # --- IT Intelligence ---
    it_intel = brief.get("it_intelligence", {})
    lines = []
    if it_intel.get("tech_stack"):
        lines.append(f"Stack: {', '.join(it_intel['tech_stack'])}")
    if it_intel.get("cloud_provider"):
        lines.append(f"Cloud: {it_intel['cloud_provider']}")
    if it_intel.get("infrastructure_scale"):
        lines.append(f"Scale: {it_intel['infrastructure_scale']}")
    if it_intel.get("recent_it_postings"):
        lines.append("")
        lines.append("[bold]Recent IT Hiring:[/bold]")
        for p in it_intel["recent_it_postings"][:5]:
            lines.append(f"  {p.get('title', '')} — {p.get('signal', '')}")
    if it_intel.get("it_challenges"):
        lines.append(f"Challenges: {', '.join(it_intel['it_challenges'])}")
    console.print(Panel("\n".join(lines), title="IT & TECHNOLOGY", border_style="green"))

    # --- Role Analysis (conditional) ---
    role = brief.get("role_analysis")
    if role:
        lines = []
        if role.get("org_fit"):
            lines.append(f"Org fit: {role['org_fit']}")
        if role.get("day_to_day"):
            lines.append(f"Day-to-day: {role['day_to_day']}")
        if role.get("growth_potential"):
            lines.append(f"Growth: {role['growth_potential']}")
        if role.get("red_flags"):
            lines.append(f"[red]Red flags:[/red] {', '.join(role['red_flags'])}")
        if role.get("questions_to_ask"):
            lines.append("")
            lines.append("[bold]Questions to ask:[/bold]")
            for q in role["questions_to_ask"]:
                lines.append(f"  - {q}")
        console.print(Panel("\n".join(lines), title="ROLE ANALYSIS", border_style="magenta"))

    # --- Interviewer Prep (conditional) ---
    prep = brief.get("interviewer_prep")
    if prep:
        lines = []
        if prep.get("linkedin_summary"):
            lines.append(prep["linkedin_summary"])
        if prep.get("likely_interview_style"):
            lines.append(f"Interview style: {prep['likely_interview_style']}")
        if prep.get("rapport_topics"):
            lines.append(f"Rapport topics: {', '.join(prep['rapport_topics'])}")
        console.print(Panel("\n".join(lines), title="INTERVIEWER PREP", border_style="cyan"))

    # --- Sources ---
    sources = brief.get("sources", [])
    if sources:
        console.print(Panel("\n".join(sources), title="SOURCES", border_style="dim"))

    generated = brief.get("generated_at", "")
    if generated:
        console.print(f"[dim]Generated: {generated}[/dim]")
    console.print()


@cli.group()
def intel():
    """Company intelligence — research briefs for target companies."""


@intel.command("research")
@click.argument("company")
@click.option("--role", default=None, help="Role title for role analysis section.")
@click.option("--contact", default=None, help="Interviewer name for prep section.")
@click.option("--url", default=None, help="Job posting URL for context.")
def intel_research(company, role, contact, url):
    """Generate a fresh company intelligence brief."""
    from src.db import models
    from src.intel.company_intel import CompanyIntelEngine

    console.print(f"[bold cyan]Researching {company}...[/bold cyan]")
    console.print("[dim]This may take 30-60 seconds (web search + analysis).[/dim]")

    engine = CompanyIntelEngine()
    brief = engine.generate_brief(
        company, role_title=role, contact_name=contact, job_url=url,
    )

    if brief is None:
        console.print("[red]Failed to generate intel brief. Check logs.[/red]")
        return

    # Cache the brief
    conn = models.get_connection()
    try:
        models.cache_brief(conn, company, role, brief)
    finally:
        conn.close()

    _display_brief(brief, company)


@intel.command("show")
@click.argument("company")
def intel_show(company):
    """Show cached company intel brief (no API call)."""
    from src.db import models

    conn = models.get_connection()
    try:
        brief, row = models.get_cached_brief(conn, company)
    finally:
        conn.close()

    if brief is None:
        console.print(f"[yellow]No cached brief for '{company}'. Run: python cli.py intel research {company}[/yellow]")
        return

    _display_brief(brief, company)


@intel.command("refresh")
@click.argument("company")
@click.option("--role", default=None, help="Role title for role analysis section.")
@click.option("--contact", default=None, help="Interviewer name for prep section.")
@click.option("--url", default=None, help="Job posting URL for context.")
def intel_refresh(company, role, contact, url):
    """Force regenerate a company intel brief (ignores cache)."""
    from src.db import models
    from src.intel.company_intel import CompanyIntelEngine

    console.print(f"[bold cyan]Refreshing intel for {company}...[/bold cyan]")

    engine = CompanyIntelEngine()
    brief = engine.generate_brief(
        company, role_title=role, contact_name=contact, job_url=url,
    )

    if brief is None:
        console.print("[red]Failed to generate intel brief. Check logs.[/red]")
        return

    conn = models.get_connection()
    try:
        models.cache_brief(conn, company, role, brief)
    finally:
        conn.close()

    _display_brief(brief, company)


@intel.command("prep")
@click.argument("application_id", type=int)
def intel_prep(application_id):
    """Generate/show intel brief for a tracked application."""
    from src.db import models

    conn = models.get_connection()
    try:
        app = conn.execute(
            "SELECT * FROM applications WHERE id = ?", (application_id,)
        ).fetchone()
        if not app:
            console.print(f"[red]Application #{application_id} not found.[/red]")
            return
        app = dict(app)

        company = app["company"]
        role_title = app["title"]

        # Check for existing linked brief
        existing = models.get_brief_for_application(conn, application_id)
        if existing:
            console.print(f"[dim]Using cached brief for application #{application_id}[/dim]")
            _display_brief(existing, company)
            return

        # Check for any cached brief for this company
        cached, row = models.get_cached_brief(conn, company)
        if cached:
            # Link it to the application
            models.link_brief_to_application(conn, row["id"], application_id)
            console.print(f"[dim]Found existing brief for {company}, linked to application #{application_id}[/dim]")
            _display_brief(cached, company)
            return

        # Generate fresh
        from src.intel.company_intel import CompanyIntelEngine

        console.print(f"[bold cyan]Researching {company} for application #{application_id}...[/bold cyan]")
        console.print("[dim]This may take 30-60 seconds.[/dim]")

        engine = CompanyIntelEngine()
        brief = engine.generate_brief(company, role_title=role_title)

        if brief is None:
            console.print("[red]Failed to generate intel brief. Check logs.[/red]")
            return

        brief_id = models.cache_brief(conn, company, role_title, brief, application_id=application_id)
        _display_brief(brief, company)
    finally:
        conn.close()


# ─────────────────────────────────────────────────────────────────────────────
# llm — LLM call log management
# ─────────────────────────────────────────────────────────────────────────────

@cli.group()
def llm():
    """View and manage LLM call logs and budget."""


@llm.command("summary")
@click.option("--days", default=7, show_default=True, help="Look-back window in days.")
@click.option("--task", "filter_task", default=None, help="Filter by task ID.")
def llm_summary(days, filter_task):
    """Show a summary of recent LLM calls grouped by provider and task."""
    from src.db import models

    conn = models.get_connection()
    try:
        where_clauses = ["created_at >= datetime('now', ?)", "latency_ms IS NOT NULL"]
        params = [f"-{days} days"]

        if filter_task:
            where_clauses.append("task = ?")
            params.append(filter_task)

        where_sql = " AND ".join(where_clauses)

        # Per-provider totals
        provider_rows = conn.execute(
            f"""
            SELECT provider_used,
                   COUNT(*) AS calls,
                   SUM(tokens_in) AS total_in,
                   SUM(tokens_out) AS total_out,
                   ROUND(AVG(latency_ms) / 1000.0, 1) AS avg_latency_s,
                   SUM(CASE WHEN schema_invalid = 1 THEN 1 ELSE 0 END) AS schema_fails,
                   SUM(CASE WHEN fallback_reason IS NOT NULL
                               AND fallback_reason NOT IN (
                                   'kill_switch','env_override',
                                   'pii_fallback_blocked','fallback_budget_exhausted'
                               ) THEN 1 ELSE 0 END) AS infra_fallbacks
            FROM llm_calls
            WHERE {where_sql}
            GROUP BY provider_used
            ORDER BY calls DESC
            """,
            params,
        ).fetchall()

        # Per-task breakdown
        task_rows = conn.execute(
            f"""
            SELECT task, provider_used,
                   COUNT(*) AS calls,
                   SUM(tokens_in) AS total_in,
                   SUM(tokens_out) AS total_out,
                   ROUND(AVG(latency_ms) / 1000.0, 1) AS avg_latency_s
            FROM llm_calls
            WHERE {where_sql}
            GROUP BY task, provider_used
            ORDER BY calls DESC
            """,
            params,
        ).fetchall()

        # Budget status
        from config import settings
        budget_row = conn.execute(
            "SELECT fallback_count_since_reset, last_reset_at FROM llm_budget_resets ORDER BY id DESC LIMIT 1"
        ).fetchone()

        title = f"LLM Summary — last {days}d"
        if filter_task:
            title += f" / task={filter_task}"
        console.print(Panel(f"[bold]{title}[/bold]", expand=False))

        # Provider table
        ptable = Table(show_header=True, header_style="bold cyan", box=None, padding=(0, 1))
        ptable.add_column("Provider", style="bold")
        ptable.add_column("Calls", justify="right")
        ptable.add_column("Tokens In", justify="right")
        ptable.add_column("Tokens Out", justify="right")
        ptable.add_column("Avg Latency", justify="right")
        ptable.add_column("Schema Fails", justify="right")
        ptable.add_column("Infra Fallbacks", justify="right")
        for r in provider_rows:
            ptable.add_row(
                r["provider_used"],
                str(r["calls"]),
                str(r["total_in"] or 0),
                str(r["total_out"] or 0),
                f"{r['avg_latency_s']}s",
                str(r["schema_fails"] or 0),
                str(r["infra_fallbacks"] or 0),
            )
        console.print(ptable)

        if not filter_task:
            # Per-task table
            ttable = Table(show_header=True, header_style="bold cyan", box=None, padding=(0, 1))
            ttable.add_column("Task", style="dim")
            ttable.add_column("Provider")
            ttable.add_column("Calls", justify="right")
            ttable.add_column("Tokens In", justify="right")
            ttable.add_column("Tokens Out", justify="right")
            ttable.add_column("Avg Latency", justify="right")
            for r in task_rows:
                ttable.add_row(
                    r["task"],
                    r["provider_used"],
                    str(r["calls"]),
                    str(r["total_in"] or 0),
                    str(r["total_out"] or 0),
                    f"{r['avg_latency_s']}s",
                )
            console.print("\n[bold]Per-task breakdown[/bold]")
            console.print(ttable)

        if budget_row:
            used = budget_row["fallback_count_since_reset"]
            limit = settings.LLM_FALLBACK_BUDGET_PER_DAY
            console.print(
                f"\n[dim]Fallback budget: {used}/{limit} used since {budget_row['last_reset_at']}[/dim]"
            )
    finally:
        conn.close()


@llm.command("prune")
@click.option("--days", default=30, show_default=True, help="Delete rows older than N days.")
@click.option("--yes", "confirmed", is_flag=True, help="Skip confirmation prompt.")
def llm_prune(days, confirmed):
    """Delete LLM call log rows older than N days."""
    from src.db import models

    conn = models.get_connection()
    try:
        count_row = conn.execute(
            "SELECT COUNT(*) AS n FROM llm_calls WHERE created_at < datetime('now', ?)",
            (f"-{days} days",),
        ).fetchone()
        count = count_row["n"]

        if count == 0:
            console.print(f"[dim]No rows older than {days} days to prune.[/dim]")
            return

        if not confirmed:
            click.confirm(f"Delete {count} rows older than {days} days?", abort=True)

        conn.execute(
            "DELETE FROM llm_calls WHERE created_at < datetime('now', ?)",
            (f"-{days} days",),
        )
        conn.commit()
        console.print(f"[green]Pruned {count} rows older than {days} days.[/green]")
    finally:
        conn.close()


@llm.command("reset-budget")
@click.option("--yes", "confirmed", is_flag=True, help="Skip confirmation prompt.")
def llm_reset_budget(confirmed):
    """Reset the rolling LLM fallback budget counter to zero."""
    from src.db import models

    conn = models.get_connection()
    try:
        row = conn.execute(
            "SELECT fallback_count_since_reset, last_reset_at FROM llm_budget_resets ORDER BY id DESC LIMIT 1"
        ).fetchone()
        if row:
            console.print(
                f"[dim]Current budget: {row['fallback_count_since_reset']} fallbacks since {row['last_reset_at']}[/dim]"
            )

        if not confirmed:
            click.confirm("Reset fallback budget counter to zero?", abort=True)

        conn.execute(
            "UPDATE llm_budget_resets SET fallback_count_since_reset = 0, last_reset_at = datetime('now') "
            "WHERE id = (SELECT id FROM llm_budget_resets ORDER BY id DESC LIMIT 1)"
        )
        conn.commit()
        console.print("[green]Fallback budget reset to 0.[/green]")
    finally:
        conn.close()


@llm.command("embed-smoke")
@click.argument("text", default="CareerPilot embedding smoke test")
def llm_embed_smoke(text):
    """Send TEXT to the local embed endpoint and print the vector dimension."""
    from src.llm.router import router

    try:
        vector = router.embed(task="embed_default", text=text)
        console.print(f"[green]OK[/green] — dim={len(vector)}, first5={[round(v, 4) for v in vector[:5]]}")
    except Exception as exc:
        console.print(f"[red]FAIL[/red] — {exc}")
        raise SystemExit(1)


if __name__ == "__main__":
    try:
        cli()
    except KeyboardInterrupt:
        console.print("\n[dim]Goodbye.[/dim]")
