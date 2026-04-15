"""Cover letter generation — Claude-powered, tailored per application."""

from __future__ import annotations

import json
import logging
import os
import re
from datetime import datetime
from pathlib import Path
from typing import Dict, Optional

from config import settings

logger = logging.getLogger(__name__)


def _sanitize_filename(text: str) -> str:
    """Remove characters that aren't safe for filenames."""
    text = re.sub(r'[<>:"/\\|?*]', "", text)
    text = re.sub(r"\s+", "_", text.strip())
    return text[:50]


class CoverLetterGenerator:
    """Generates tailored cover letters using Claude + python-docx."""

    def __init__(self, profile: Dict = None):
        """Initialize with candidate profile data.

        Args:
            profile: Profile dict from ProfileManager.get_profile().
                     If None, profile is loaded lazily when needed.
        """
        self._profile = profile

    def _get_profile(self) -> Dict:
        """Get or load the candidate profile."""
        if self._profile is None:
            from src.profile.manager import ProfileManager
            mgr = ProfileManager()
            self._profile = mgr.get_profile()
            mgr.close()
        return self._profile

    def generate_cover_letter(
        self,
        job_description: str,
        company: str,
        role: str,
        fit_analysis: Dict = None,
    ) -> Optional[str]:
        """Generate a tailored cover letter via Claude.

        Args:
            job_description: Full job description text.
            company: Company name.
            role: Role title.
            fit_analysis: Optional fit analysis dict from JobAnalyzer.

        Returns:
            Cover letter text, or None on failure.
        """
        profile = self._get_profile()

        # Build profile summary for Claude
        personal = profile.get("personal", {})
        work_history = profile.get("work_history", [])
        certs = profile.get("certifications", [])

        profile_text = f"Name: {personal.get('full_name', 'Joseph Fowler')}\n"
        profile_text += f"Email: {personal.get('email', '')}\n"
        profile_text += f"Phone: {personal.get('phone', '')}\n"
        profile_text += f"Location: {personal.get('city', '')}, {personal.get('state', '')}\n\n"

        profile_text += "Work Experience:\n"
        for w in work_history:
            end = w.get("end_date") or "Present"
            profile_text += f"- {w.get('title', '')} at {w.get('company', '')} ({w.get('start_date', '')} - {end})\n"
            if w.get("description"):
                profile_text += f"  {w['description']}\n"

        profile_text += "\nCertifications:\n"
        for c in certs:
            status = " (In Progress)" if c.get("in_progress") else ""
            profile_text += f"- {c.get('name', '')}{status}\n"

        user_msg = (
            f"## Target Position\nCompany: {company}\nRole: {role}\n\n"
            f"## Job Description\n{job_description}\n\n"
            f"## Candidate Profile\n{profile_text}"
        )

        if fit_analysis:
            user_msg += (
                f"\n## Fit Analysis\n"
                f"Match Score: {fit_analysis.get('match_score', 'N/A')}/10\n"
                f"Matching Skills: {', '.join(fit_analysis.get('matching_skills', []))}\n"
                f"Gap Skills: {', '.join(fit_analysis.get('gap_skills', []))}\n"
            )

        try:
            from src.llm.router import router
            return router.complete(task="cover_letter", prompt=user_msg[:20000])
        except Exception:
            logger.error("Cover letter generation failed", exc_info=True)
            return None

    def generate_docx(
        self,
        cover_letter_text: str,
        company: str,
        role: str,
        output_path: str,
    ) -> str:
        """Create a .docx file with professional letter formatting.

        Returns the output path.
        """
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        profile = self._get_profile()
        personal = profile.get("personal", {})

        doc = Document()

        # Page margins
        for section in doc.sections:
            section.top_margin = Inches(1.0)
            section.bottom_margin = Inches(1.0)
            section.left_margin = Inches(1.0)
            section.right_margin = Inches(1.0)

        style = doc.styles["Normal"]
        style.font.name = "Calibri"
        style.font.size = Pt(11)

        # Candidate contact info header
        name = personal.get("full_name", "Joseph Fowler")
        header = doc.add_paragraph()
        header.alignment = WD_ALIGN_PARAGRAPH.LEFT
        name_run = header.add_run(name)
        name_run.bold = True
        name_run.font.size = Pt(14)

        contact_parts = []
        if personal.get("email"):
            contact_parts.append(personal["email"])
        if personal.get("phone"):
            contact_parts.append(personal["phone"])
        city = personal.get("city", "")
        state = personal.get("state", "")
        if city or state:
            contact_parts.append(f"{city}, {state}".strip(", "))
        if personal.get("linkedin_url"):
            contact_parts.append(personal["linkedin_url"])

        if contact_parts:
            contact = doc.add_paragraph(" | ".join(contact_parts))
            contact.paragraph_format.space_after = Pt(4)

        # Date
        date_str = datetime.now().strftime("%B %d, %Y")
        date_para = doc.add_paragraph(date_str)
        date_para.paragraph_format.space_before = Pt(12)
        date_para.paragraph_format.space_after = Pt(12)

        # Salutation
        doc.add_paragraph("Dear Hiring Manager,")

        # Cover letter body paragraphs
        for paragraph in cover_letter_text.split("\n\n"):
            paragraph = paragraph.strip()
            if paragraph:
                p = doc.add_paragraph(paragraph)
                p.paragraph_format.space_after = Pt(6)

        # Closing
        closing = doc.add_paragraph()
        closing.paragraph_format.space_before = Pt(12)
        closing.add_run("Sincerely,")
        sig = doc.add_paragraph()
        sig.add_run(name)

        # Ensure output directory exists
        Path(output_path).parent.mkdir(parents=True, exist_ok=True)
        doc.save(output_path)
        logger.info("Cover letter saved to %s", output_path)
        return output_path

    def generate_for_application(
        self,
        job_data: Dict,
        output_dir: str = None,
    ) -> Optional[str]:
        """Full pipeline: generate cover letter and save as .docx.

        Args:
            job_data: Dict with 'description' (or 'snippet'), 'company', 'title'.
            output_dir: Directory for output files. Defaults to data/cover_letters/.

        Returns:
            Path to generated .docx file, or None on failure.
        """
        if output_dir is None:
            output_dir = str(settings.DATA_DIR / "cover_letters")

        description = job_data.get("description", job_data.get("snippet", ""))
        company = job_data.get("company", "Unknown")
        role = job_data.get("title", "Unknown")

        # Optionally run fit analysis
        fit_analysis = None
        try:
            from src.jobs.analyzer import JobAnalyzer
            analyzer = JobAnalyzer()
            fit_analysis = analyzer.analyze_fit(description)
        except Exception:
            logger.debug("Fit analysis unavailable, generating without it")

        text = self.generate_cover_letter(
            description, company, role, fit_analysis=fit_analysis,
        )
        if not text:
            logger.error("Failed to generate cover letter for %s at %s", role, company)
            return None

        date_str = datetime.now().strftime("%Y%m%d")
        filename = f"{_sanitize_filename(company)}_{_sanitize_filename(role)}_{date_str}_CL.docx"
        output_path = os.path.join(output_dir, filename)

        return self.generate_docx(text, company, role, output_path)
