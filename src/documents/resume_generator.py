"""Resume tailoring — Claude-powered resume optimization per job description."""

from __future__ import annotations

import json
import logging
import os
import re
from datetime import datetime
from pathlib import Path
from typing import Dict, Optional

from config import settings

logger = logging.getLogger(__name__)

# Joseph Fowler's base resume as a structured dict
BASE_RESUME = {
    "professional_summary": (
        "Senior Systems Engineer with 20 years of progressive IT infrastructure "
        "experience. Proven track record leading enterprise operations for 900+ users "
        "across multi-site environments. Deep expertise in Windows Server, Active "
        "Directory, VMware vSphere, PowerShell automation, and Microsoft 365/Azure "
        "administration. Skilled in monitoring platforms (Splunk, SolarWinds), disaster "
        "recovery, and infrastructure modernization. Currently expanding into cloud-native "
        "technologies including Azure, Docker, Kubernetes, and Terraform."
    ),
    "core_skills": [
        "PowerShell Automation",
        "Windows Server Administration",
        "Active Directory / Group Policy",
        "VMware vSphere",
        "Microsoft 365 / Azure AD",
        "Splunk / SolarWinds Monitoring",
        "Python Scripting",
        "Docker / Kubernetes",
        "Terraform / IaC",
        "Git / GitHub",
        "Networking (DNS, DHCP, TCP/IP)",
        "Citrix XenApp / XenDesktop",
        "SCCM / Patch Management",
        "Veeam Backup & Recovery",
        "CI/CD Pipelines",
    ],
    "experience": [
        {
            "company": "Venable LLP",
            "title": "Senior Systems Engineer",
            "dates": "January 2020 - March 2025",
            "location": "Washington, DC / Remote",
            "bullets": [
                "Led enterprise infrastructure operations supporting 900+ users across 9 offices",
                "Managed VMware vSphere environment (20+ ESXi hosts, 200+ VMs)",
                "Administered Microsoft 365, Azure AD, and hybrid identity infrastructure",
                "Automated routine tasks with PowerShell, reducing manual effort by 40%",
                "Maintained Splunk and SolarWinds monitoring platforms",
                "Coordinated disaster recovery testing and documentation",
            ],
        },
        {
            "company": "Venable LLP",
            "title": "Systems Engineer II",
            "dates": "June 2015 - December 2019",
            "location": "Washington, DC",
            "bullets": [
                "Managed Windows Server infrastructure (2012R2/2016/2019) across data centers",
                "Administered Active Directory, Group Policy, DNS, DHCP for 900+ users",
                "Supported Citrix XenApp/XenDesktop virtual desktop environment",
                "Implemented SCCM for OS deployment and patch management",
                "Maintained network infrastructure including Cisco switches and firewalls",
            ],
        },
        {
            "company": "Venable LLP",
            "title": "Systems Engineer I",
            "dates": "March 2010 - May 2015",
            "location": "Washington, DC",
            "bullets": [
                "Provided Tier 2/3 support for server and network infrastructure",
                "Managed backup and recovery operations using Veeam",
                "Supported Exchange 2010/2013 messaging environment",
                "Assisted with VMware vSphere administration and migrations",
            ],
        },
        {
            "company": "Venable LLP",
            "title": "IT Support Specialist",
            "dates": "August 2005 - February 2010",
            "location": "Washington, DC",
            "bullets": [
                "Provided desktop and application support for attorneys and staff",
                "Managed hardware lifecycle including imaging, deployment, and retirement",
                "Supported document management and legal-specific applications",
                "Created technical documentation and knowledge base articles",
            ],
        },
    ],
    "education": [
        {
            "degree": "Network Information Systems Certificate",
            "school": "Tesst College of Technology",
            "year": "2005",
        },
    ],
    "certifications": [
        "Microsoft Azure Fundamentals (AZ-900) — In Progress",
        "ITIL V4 Foundation",
        "CompTIA Security+",
    ],
    "technical_knowledge": {
        "Operating Systems": "Windows Server 2012R2/2016/2019/2022, Windows 10/11",
        "Virtualization": "VMware vSphere/ESXi, Citrix XenApp/XenDesktop",
        "Cloud": "Microsoft Azure, Azure AD, Microsoft 365",
        "Automation": "PowerShell, Python, Terraform, CI/CD",
        "Monitoring": "Splunk, SolarWinds",
        "Backup & DR": "Veeam Backup & Replication",
        "Networking": "DNS, DHCP, TCP/IP, Cisco switches/firewalls",
        "Containers": "Docker, Kubernetes",
        "Version Control": "Git, GitHub",
    },
}

def _sanitize_filename(text: str) -> str:
    """Remove characters that aren't safe for filenames."""
    text = re.sub(r'[<>:"/\\|?*]', "", text)
    text = re.sub(r"\s+", "_", text.strip())
    return text[:50]


class ResumeGenerator:
    """Generates tailored resumes using Claude + python-docx."""

    def __init__(self, base_resume: Dict = None):
        self._base = base_resume or BASE_RESUME

    def tailor_resume(
        self,
        job_description: str,
        company: str = None,
        role: str = None,
    ) -> Optional[Dict]:
        """Send base resume + job description to Claude for tailoring.

        Returns tailored resume structure or None on failure.
        """
        context_parts = []
        if company:
            context_parts.append(f"Company: {company}")
        if role:
            context_parts.append(f"Role: {role}")
        context_line = "\n".join(context_parts)

        user_msg = (
            f"## Target Job\n{context_line}\n\n"
            f"## Job Description\n{job_description}\n\n"
            f"## Base Resume (JSON)\n{json.dumps(self._base, indent=2)}"
        )

        try:
            from src.llm.router import router
            tailored = router.complete(task="resume_generate", prompt=user_msg[:30000])
            if tailored:
                for key in self._base:
                    tailored.setdefault(key, self._base[key])
            return tailored
        except Exception:
            logger.error("Resume tailoring failed", exc_info=True)
            return None

    def generate_docx(self, tailored_data: Dict, output_path: str) -> str:
        """Create a .docx file from the tailored resume structure.

        Returns the output path.
        """
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()

        # Page margins
        for section in doc.sections:
            section.top_margin = Inches(0.6)
            section.bottom_margin = Inches(0.6)
            section.left_margin = Inches(0.75)
            section.right_margin = Inches(0.75)

        style = doc.styles["Normal"]
        style.font.name = "Calibri"
        style.font.size = Pt(10.5)

        # Name header
        name_para = doc.add_paragraph()
        name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = name_para.add_run("Joseph Fowler")
        run.bold = True
        run.font.size = Pt(18)
        run.font.name = "Calibri"

        # Contact line
        contact = doc.add_paragraph()
        contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact_run = contact.add_run(
            "jlfowler1084@gmail.com | 443-787-6528 | Sheridan, IN | github.com/jlfowler1084"
        )
        contact_run.font.size = Pt(9)
        contact_run.font.name = "Calibri"

        def add_section_heading(text):
            p = doc.add_paragraph()
            p.space_before = Pt(8)
            p.space_after = Pt(2)
            run = p.add_run(text.upper())
            run.bold = True
            run.font.size = Pt(11)
            run.font.name = "Calibri"
            # Add bottom border via paragraph formatting
            pf = p.paragraph_format
            pf.space_after = Pt(2)

        # Professional Summary
        add_section_heading("Professional Summary")
        summary = tailored_data.get("professional_summary", "")
        p = doc.add_paragraph(summary)
        p.paragraph_format.space_after = Pt(4)

        # Core Skills
        add_section_heading("Core Skills")
        skills = tailored_data.get("core_skills", [])
        if skills:
            p = doc.add_paragraph(" | ".join(skills))
            p.paragraph_format.space_after = Pt(4)

        # Professional Experience
        add_section_heading("Professional Experience")
        for role in tailored_data.get("experience", []):
            # Role title line
            p = doc.add_paragraph()
            p.space_before = Pt(4)
            title_run = p.add_run(f"{role['title']}, {role['company']}")
            title_run.bold = True
            title_run.font.size = Pt(10.5)

            # Date/location line
            p2 = doc.add_paragraph()
            p2.space_before = Pt(0)
            loc_run = p2.add_run(f"{role.get('location', '')} | {role.get('dates', '')}")
            loc_run.italic = True
            loc_run.font.size = Pt(9.5)

            # Bullets
            for bullet in role.get("bullets", []):
                bp = doc.add_paragraph(bullet, style="List Bullet")
                bp.paragraph_format.space_before = Pt(1)
                bp.paragraph_format.space_after = Pt(1)

        # Education
        add_section_heading("Education")
        for edu in tailored_data.get("education", []):
            p = doc.add_paragraph()
            run = p.add_run(edu.get("degree", ""))
            run.bold = True
            p.add_run(f"\n{edu.get('school', '')} | {edu.get('year', '')}")

        # Certifications
        add_section_heading("Certifications")
        for cert in tailored_data.get("certifications", []):
            doc.add_paragraph(f"- {cert}")

        # Technical Knowledge
        tech = tailored_data.get("technical_knowledge", {})
        if tech:
            add_section_heading("Technical Knowledge")
            for category, details in tech.items():
                p = doc.add_paragraph()
                cat_run = p.add_run(f"{category}: ")
                cat_run.bold = True
                p.add_run(details)

        # Ensure output directory exists
        Path(output_path).parent.mkdir(parents=True, exist_ok=True)
        doc.save(output_path)
        logger.info("Resume saved to %s", output_path)
        return output_path

    def generate_for_application(
        self,
        job_data: Dict,
        output_dir: str = None,
    ) -> Optional[str]:
        """Full pipeline: tailor resume and generate .docx.

        Args:
            job_data: Dict with 'description' (or 'snippet'), 'company', 'title'.
            output_dir: Directory for output files. Defaults to data/resumes/.

        Returns:
            Path to generated .docx file, or None on failure.
        """
        if output_dir is None:
            output_dir = str(settings.DATA_DIR / "resumes")

        description = job_data.get("description", job_data.get("snippet", ""))
        company = job_data.get("company", "Unknown")
        role = job_data.get("title", "Unknown")

        tailored = self.tailor_resume(description, company=company, role=role)
        if not tailored:
            logger.error("Failed to tailor resume for %s at %s", role, company)
            return None

        date_str = datetime.now().strftime("%Y%m%d")
        filename = f"{_sanitize_filename(company)}_{_sanitize_filename(role)}_{date_str}.docx"
        output_path = os.path.join(output_dir, filename)

        return self.generate_docx(tailored, output_path)
